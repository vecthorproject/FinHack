import io
import uuid
from pydoc import doc
import zipfile
import copy
import pandas as pd
import numpy as np
import docx
import os
import re 
from docxtpl import DocxTemplate
import matplotlib.pyplot as plt
from docxtpl import InlineImage
from docx.shared import Mm
import warnings
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn as _qn
from lxml import etree
import tempfile
import matplotlib.image as mpimg
import matplotlib.patches as patches

warnings.filterwarnings("ignore", category=UserWarning, module="matplotlib")


PERCORSO_LOGO_WATERMARK = "logofv.png" # WM

def aggiungi_watermark_fig(fig, applica):
    """Sovrappone logo e testo opachi, MANTENENDO intatte le dimensioni originali"""
    if not applica:
        return
        
    # Peschiamo l'asse originale del grafico per ancorarci a lui (evita sfasamenti e resize!)
    ax = fig.gca()
    
    # 1. 🛡️ VETRO SATINATO (alpha=0.96) ancorato all'area interna del grafico
    rect = patches.Rectangle((-0.1, -0.1), 1.2, 1.2, transform=ax.transAxes,
                             linewidth=0, edgecolor='none', facecolor='white', 
                             alpha=0.96, zorder=998, clip_on=False)
    ax.add_patch(rect)
    
    # 2. Creiamo un asse "fantasma" tenendolo rigorosamente DENTRO i margini (non allarga il grafico)
    ax_bg = fig.add_axes([0.1, 0.1, 0.8, 0.8], zorder=999)
    ax_bg.axis('off')
    
    try:
        # Inserisce il Logo a sinistra (Opacità 100%, Colori Vividi)
        logo = mpimg.imread(PERCORSO_LOGO_WATERMARK)
        ax_bg.imshow(logo, extent=[0.0, 0.40, 0.10, 0.90], alpha=1.0, aspect='auto')
        
        # Inserisce il Testo a destra (Opacità 100%, Colori Vividi)
        ax_bg.text(0.75, 0.5, "Finance & Value\nREPORT PREMIUM", fontsize=24, color='#2B3A67', 
                   alpha=1.0, ha='center', va='center', rotation=15, weight='bold')
    except:
        ax_bg.text(0.5, 0.5, "Finance & Value\nDATI BLOCCATI", fontsize=32, color='#2B3A67', 
                   alpha=1.0, ha='center', va='center', rotation=15, weight='bold')
        

def format_euro(numero, decimali=2):
    if pd.isna(numero): return "n.d."
    formato = f"{{:,.{decimali}f}}".format(numero)
    return formato.replace(',', 'X').replace('.', ',').replace('X', '.')


# =================================================================
# 🟢 INDICATORI ECONOMICI (100% SEPARATI)
# =================================================================

def get_intro_margini(descr_settore):
    return f"Le risultanze relative al tessuto competitivo del mercato ({descr_settore}) evidenziano che:"

def get_analisi_ebitda(az_ebitda, set_ebitda):
    if az_ebitda < set_ebitda:
        return f"• Il Margine EBITDA ({format_euro(az_ebitda)}%) risulta inferiore alla mediana settoriale ({format_euro(set_ebitda)}%). Valori più contenuti segnalano una minore capacità di trasformare i ricavi in margine operativo lordo, denotando una minore efficienza della gestione caratteristica prima degli ammortamenti e delle svalutazioni."
    else:
        return f"• Il Margine EBITDA ({format_euro(az_ebitda)}%) supera la mediana settoriale ({format_euro(set_ebitda)}%). Valori più elevati indicano una maggiore capacità dell'impresa di generare reddito dalla gestione caratteristica prima di ammortamenti e svalutazioni, evidenziando una superiore efficienza operativa."

def get_analisi_ebit(az_ebit, set_ebit):
    if az_ebit < set_ebit:
        return f"• Il Margine EBIT ({format_euro(az_ebit)}%) si colloca al di sotto del target mediano ({format_euro(set_ebit)}%). Tale dato segnala una minore capacità di generare reddito operativo in relazione ai ricavi conseguiti a valle dell'assorbimento dei costi fissi operativi e degli ammortamenti."
    else:
        return f"• Il Margine EBIT ({format_euro(az_ebit)}%) si posiziona al di sopra della mediana di settore ({format_euro(set_ebit)}%). Questo livello indica una maggiore capacità di conseguire un risultato operativo soddisfacente dopo aver considerato gli ammortamenti e le svalutazioni."

def get_analisi_margine_profitto_tag(az_prof, set_prof):
    if az_prof < set_prof:
        return f"• Il Margine di Profitto ({format_euro(az_prof)}%) risulta inferiore alla mediana settoriale ({format_euro(set_prof)}%). Tale andamento denota una minore capacità di trasformare i ricavi in utile netto, evidenziando criticità nell'assorbimento della gestione straordinaria, degli oneri finanziari o del carico fiscale."
    else:
        return f"• Il Margine di Profitto ({format_euro(az_prof)}%) supera il parametro mediano del settore ({format_euro(set_prof)}%). Valori più elevati indicano una maggiore capacità dell'impresa di convertire i ricavi in risultato netto finale, confermando un'efficace gestione degli oneri extra-caratteristici."

# =================================================================
# 🟠 INDICATORI PATRIMONIALI (100% SEPARATI)
# =================================================================

def get_analisi_struttura1(az_str1):
    if az_str1 >= 1:
        return f"• L'Indice primario di struttura ({format_euro(az_str1)}), superiore o uguale all'unità, indica che il capitale proprio, il quale non ha vincoli di scadenza, ha finanziato interamente le immobilizzazioni, caratterizzate da tempi di disinvestimento medio-lunghi."
    else:
        return f"• L'Indice primario di struttura ({format_euro(az_str1)}), risultando inferiore ad uno, segnala che una parte delle immobilizzazioni è stata finanziata mediante capitale di terzi, con potenziale obbligo di rimborso nel breve termine."

def get_analisi_struttura2(az_str2):
    if az_str2 >= 1:
        return f"• L'Indice secondario di struttura ({format_euro(az_str2)}), superiore o uguale all'unità, conferma che il capitale permanente, costituito dal capitale proprio e dai debiti a medio-lunga scadenza, ha finanziato interamente gli asset immobilizzati."
    else:
        return f"• L'Indice secondario di struttura ({format_euro(az_str2)}), essendo inferiore ad uno, indica che una parte dell'attivo immobilizzato è finanziata attraverso capitale di terzi a breve scadenza, determinando uno squilibrio temporale tra fonti e impieghi."

def get_analisi_gearing_tag(az_gear, set_gear):
    if az_gear <= set_gear:
        return f"• Il Gearing ({format_euro(az_gear)}%) si attesta al di sotto del parametro mediano del comparto ({format_euro(set_gear)}%). Tali valori contenuti indicano una limitata dipendenza dall'indebitamento oneroso e una solida autonomia rispetto ai creditori."
    else:
        return f"• Il Gearing ({format_euro(az_gear)}%) supera la mediana di settore ({format_euro(set_gear)}%). Valori più elevati segnalano un maggiore ricorso al capitale di terzi per il finanziamento aziendale, determinando un incremento del rischio finanziario e una minore autonomia."

# =================================================================
# 🔵 INDICATORI FINANZIARI (100% SEPARATI)
# =================================================================

def get_analisi_current_ratio_tag(az_cr, set_cr):
    if az_cr >= 1:
        return f"• Il Current Ratio ({format_euro(az_cr)}), superiore o uguale all'unità, indica che le attività a breve termine sono sufficienti a coprire integralmente i debiti esigibili nel breve periodo, evidenziando una situazione di equilibrio d'esercizio."
    else:
        return f"• Il Current Ratio ({format_euro(az_cr)}), inferiore ad uno, segnala l'incapacità delle attività correnti di far fronte alle passività correnti, configurando una potenziale tensione di liquidità all'interno della struttura d'esercizio."

def get_analisi_quick_ratio_tag(az_qr, set_qr):
    if az_qr >= 1:
        return f"• Il Quick Ratio ({format_euro(az_qr)}), superiore o uguale all'unità, indica che le risorse prontamente liquidabili sono sufficienti a garantire la copertura dei debiti a breve termine senza ricorrere alla vendita delle rimanenze di magazzino."
    else:
        return f"• Il Quick Ratio ({format_euro(az_qr)}), essendo inferiore ad uno, evidenzia una dipendenza, almeno parziale, dalla monetizzazione delle scorte o da ulteriori fonti di finanziamento esterne per soddisfare gli impegni immediati."

def get_analisi_rotazione_tag(az_rot, set_rot):
    if az_rot < set_rot:
        return f"• L'Indice di rotazione del capitale investito ({format_euro(az_rot)}) risulta inferiore alla mediana del comparto ({format_euro(set_rot)}). Valori più contenuti segnalano una minore capacità del capitale investito di tradursi in ricavi, denotando un impiego meno efficiente degli asset operativi."
    else:
        return f"• L'Indice di rotazione del capitale investito ({format_euro(az_rot)}) supera la mediana settoriale ({format_euro(set_rot)}). Valori più elevati indicano una maggiore capacità dell'impresa di generare ricavi attraverso le risorse investite, evidenziando un efficiente utilizzo del capitale."


# =====================================================================
# ☢️ LA LAVATRICE NUCLEARE (Ricostruisce l'XML di Word da zero)
# =====================================================================

def lavatrice_nucleare(template_path):
    doc_temp = docx.Document(template_path)
    
    def ripara_paragrafo(p):
        testo = p.text
        # Se trova un tag di tabella o grafico
        if '{{' in testo and '}}' in testo and any(tag in testo for tag in ['{{ tabella_', '{{ grafico_', '{{tabella_', '{{grafico_']):
            # Taglia il testo usando i tag come forbici
            frammenti = re.split(r'(\{\{.*?\}\})', testo)
            
            for frammento in frammenti:
                pulito = frammento.strip()
                if pulito: # Se non è vuoto
                    # Crea un paragrafo puro e totalmente indipendente
                    nuovo_p = p.insert_paragraph_before('')
                    nuovo_p.clear()
                    nuovo_p.add_run(pulito)
                    try:
                        nuovo_p.style = p.style
                    except:
                        pass
            # Svuota il paragrafo originale corrotto, lasciando un a capo pulito obbligatorio
            p._p.getparent().remove(p._p)

    # 1. Scansiona tutto il documento
    for p in doc_temp.paragraphs:
        ripara_paragrafo(p)
        
    # 2. Scansiona le vecchie tabelle (se ci hai lasciato tag dentro per sbaglio)
    for table in doc_temp.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    ripara_paragrafo(p)
                # Obbliga ogni cella a finire con un a capo (Regola d'oro di Word)
                try:
                    cell.add_paragraph("")
                except:
                    pass

    temp_path = os.path.join(tempfile.gettempdir(), f"template_nucleare_{uuid.uuid4().hex[:8]}.docx")
    doc_temp.save(temp_path)
    return temp_path

# =====================================================================
# MOTORE DI CALCOLO E IMPAGINAZIONE
# =====================================================================
def genera_report_word(zip_buffer, template_path, azienda_target, df_orbis, settore_nace, num_max_soc_orbis, modalita_teaser=False):
    
    # 🛡️ ANTIDOTO: Disintegra i caratteri invisibili di Excel che corrompono Word
    caratteri_proibiti = dict.fromkeys(range(0, 32))
    caratteri_proibiti.pop(9, None); caratteri_proibiti.pop(10, None); caratteri_proibiti.pop(13, None)
    azienda_target = str(azienda_target).translate(caratteri_proibiti)
    settore_nace = str(settore_nace).translate(caratteri_proibiti)
    df_orbis = df_orbis.replace(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', regex=True)

    df_orbis = df_orbis.copy()

    # --- PULIZIA DATI INIZIALE (Elimina tutti gli n.d. testuali) ---
    df_orbis = df_orbis.replace(['n.d.', 'n.a.', 'n.s.', 'N.D.', 'N.A.', 'N.S.', 'n.d', ' '], np.nan)
    # ---------------------------------------------------------------

    colonne_da_convertire = [
        'Totale valore della produzione migl EUR 2024', 'Totale Attivo migl EUR 2024',
        'Numero dipendenti 2024', 'Margine di Profitto (*) % 2024', 'Margine EBITDA (*) % 2024',
        'Margine EBIT (*) % 2024', 'Indice di Struttura 1° livello (*) 2024',
        'Indice di Struttura 2° livello (*) 2024', 'Gearing (*) % 2024',
        'Indice di Rotazione del Capitale Investito (*) 2024', 'Current Ratio (*) 2024', 'Quick Ratio (*) 2024'
    ]
    for c in colonne_da_convertire:
        if c in df_orbis.columns:
            df_orbis[c] = pd.to_numeric(df_orbis[c], errors='coerce')

    col_ragione = [c for c in df_orbis.columns if 'ragione' in str(c).lower()][0]
    col_nuts = [c for c in df_orbis.columns if 'nuts2' in str(c).lower() or 'nuts 2' in str(c).lower()]
    col_regione = col_nuts[0] if col_nuts else None

    df_orbis['Forma Giuridica Pulita'] = df_orbis['Forma giuridica nazionale'].astype(str).str.replace(r'\s*\(Italia\)', '', regex=True).str.strip()
    
    # 🔠 DIZIONARIO ABBREVIAZIONI FORME GIURIDICHE
    mappatura_fg = {
        'società a responsabilità limitata semplificata': 'S.r.l.s.',
        "societa' a responsabilita' limitata semplificata": 'S.r.l.s.',
        "societa a responsabilita limitata semplificata": 'S.r.l.s.',
        'società a responsabilità limitata': 'S.r.l.',
        "societa' a responsabilita' limitata": 'S.r.l.',
        "societa a responsabilita limitata": 'S.r.l.',
        'società per azioni': 'S.p.A.',
        "societa' per azioni": 'S.p.A.',
        "societa per azioni": 'S.p.A.',
        'società in nome collettivo': 'S.n.c.',
        "societa' in nome collettivo": 'S.n.c.',
        "societa in nome collettivo": 'S.n.c.',
        'società in accomandita semplice': 'S.a.s.',
        "societa' in accomandita semplice": 'S.a.s.',
        "societa in accomandita semplice": 'S.a.s.',
        'società in accomandita per azioni': 'S.a.p.a.',
        "societa' in accomandita per azioni": 'S.a.p.a.',
        "societa in accomandita per azioni": 'S.a.p.a.',
        'società cooperativa': 'Soc. Coop.',
        "societa' cooperativa": 'Soc. Coop.',
        "societa cooperativa": 'Soc. Coop.',
        'società semplice': 'S.s.',
        "societa' semplice": 'S.s.',
        "societa semplice": 'S.s.',
        'società consortile': 'Soc. Cons.',
        "societa' consortile": 'Soc. Cons.',
        "societa consortile": 'Soc. Cons.'
    }

    def abbrevia_forma_giuridica(fg_estesa):
        if pd.isna(fg_estesa): return "N.D."
        fg_lower = str(fg_estesa).lower().strip()
        for chiave, abbreviazione in mappatura_fg.items():
            if chiave in fg_lower: return abbreviazione
        if "s.r.l." in fg_lower or "srl" in fg_lower: return "S.r.l."
        if "s.p.a." in fg_lower or "spa" in fg_lower: return "S.p.A."
        return str(fg_estesa).title()

    df_orbis['Forma Giuridica Pulita'] = df_orbis['Forma Giuridica Pulita'].apply(abbrevia_forma_giuridica)

    def get_macro(nuts2):
        code = str(nuts2)[:3]
        if code == 'ITC': return 'Nord Ovest'
        elif code == 'ITH': return 'Nord Est'
        elif code == 'ITI': return 'Centro'
        elif code in ['ITF', 'ITG']: return 'Sud e Isole'
        else: return 'Altro'
    
    if col_regione: df_orbis['Macroregione'] = df_orbis[col_regione].apply(get_macro)
    else: df_orbis['Macroregione'] = 'Altro'

    tot_imprese_settore = len(df_orbis)
    tot_ricavi_settore = df_orbis['Totale valore della produzione migl EUR 2024'].sum()
    tot_attivo_settore = df_orbis['Totale Attivo migl EUR 2024'].sum()
    tot_dipendenti_settore = df_orbis['Numero dipendenti 2024'].sum()

    # 🟢 LOGICA DETTAGLIATA FORME GIURIDICHE (Top 1, Top 2 e Altre)
    fg_counts = df_orbis['Forma Giuridica Pulita'].value_counts()

    # Forma Giuridica #1 (Maggioranza assoluta)
    fg_1_name = str(fg_counts.index[0]).strip() if len(fg_counts) > 0 else "N.D."
    fg_1_num = int(fg_counts.iloc[0]) if len(fg_counts) > 0 else 0
    fg_1_perc = (fg_1_num / tot_imprese_settore) * 100 if tot_imprese_settore > 0 else 0
    
    # Variabili per retro-compatibilità con il resto del tuo codice
    fg_maggioranza = fg_1_name 
    num_fg_maggioranza = fg_1_num
    perc_fg_maggioranza = fg_1_perc

    # Forma Giuridica #2 (La seconda forza del mercato)
    fg_2_name = str(fg_counts.index[1]).strip() if len(fg_counts) > 1 else None
    fg_2_num = int(fg_counts.iloc[1]) if len(fg_counts) > 1 else 0
    fg_2_perc = (fg_2_num / tot_imprese_settore) * 100 if tot_imprese_settore > 0 else 0

    # Forme Giuridiche #3+ (Tutte le "briciole" sommate insieme, se esistono)
    fg_altre_num = sum(fg_counts.iloc[2:]) if len(fg_counts) > 2 else 0
    fg_altre_perc = (fg_altre_num / tot_imprese_settore) * 100 if tot_imprese_settore > 0 else 0

    df_target = df_orbis[df_orbis[col_ragione].astype(str).str.lower().str.contains(azienda_target.lower().strip(), regex=False, na=False)]
    
    def get_val_and_rank(col_name, is_lower_better=False):
        if col_name not in df_orbis.columns or df_target.empty: return "n.d.", "n.d.", "n.d."
        target_idx = df_target.index[0]
        valore = df_target.at[target_idx, col_name]
        rank_naz = df_orbis[col_name].rank(ascending=is_lower_better, method='min')
        pos_naz = int(rank_naz.loc[target_idx]) if pd.notna(rank_naz.loc[target_idx]) else "n.d."
        if col_regione:
            rank_reg = df_orbis.groupby(col_regione)[col_name].rank(ascending=is_lower_better, method='min')
            pos_reg = int(rank_reg.loc[target_idx]) if pd.notna(rank_reg.loc[target_idx]) else "n.d."
        else: pos_reg = "n.d."
        return format_euro(valore), pos_naz, pos_reg

    mg_prof, rnk_naz_prof, rnk_reg_prof = get_val_and_rank('Margine di Profitto (*) % 2024', False)
    mg_ebitda, rnk_naz_ebitda, rnk_reg_ebitda = get_val_and_rank('Margine EBITDA (*) % 2024', False)
    mg_ebit, rnk_naz_ebit, rnk_reg_ebit = get_val_and_rank('Margine EBIT (*) % 2024', False)
    ind_str1, rnk_naz_strut1, rnk_reg_strut1 = get_val_and_rank('Indice di Struttura 1° livello (*) 2024', False)
    ind_str2, rnk_naz_strut2, rnk_reg_strut2 = get_val_and_rank('Indice di Struttura 2° livello (*) 2024', False)
    gearing, rnk_naz_gear, rnk_reg_gear = get_val_and_rank('Gearing (*) % 2024', True) 
    ind_rot, rnk_naz_rot, rnk_reg_rot = get_val_and_rank('Indice di Rotazione del Capitale Investito (*) 2024', False)
    ind_cr, rnk_naz_cr, rnk_reg_cr = get_val_and_rank('Current Ratio (*) 2024', False)
    ind_qr, rnk_naz_qr, rnk_reg_qr = get_val_and_rank('Quick Ratio (*) 2024', False)

    if not df_target.empty:
        riga = df_target.iloc[0]
        p_iva = riga.get('Codice fiscale/Partita IVA', 'N.D.')
        forma_giuridica = riga.get('Forma Giuridica Pulita', 'N.D.')
        macroregione_target = riga.get('Macroregione', 'N.D.')
        num_fg_target = fg_counts.get(forma_giuridica, 0)
        perc_fg_target = (num_fg_target / tot_imprese_settore) * 100 if tot_imprese_settore > 0 else 0
        ricavi_mgl = riga.get('Totale valore della produzione migl EUR 2024')
        attivo_mgl = riga.get('Totale Attivo migl EUR 2024')
        dipendenti = riga.get('Numero dipendenti 2024')
        ricavi_mln = ricavi_mgl / 1000 if pd.notna(ricavi_mgl) else 0
        attivo_mln = attivo_mgl / 1000 if pd.notna(attivo_mgl) else 0
        perc_ricavi_panel = (ricavi_mgl / tot_ricavi_settore * 100) if tot_ricavi_settore > 0 and pd.notna(ricavi_mgl) else 0
        perc_attivo_panel = (attivo_mgl / tot_attivo_settore * 100) if tot_attivo_settore > 0 and pd.notna(attivo_mgl) else 0
        tot_dip_area = df_orbis[df_orbis['Macroregione'] == macroregione_target]['Numero dipendenti 2024'].sum()
        perc_dip_area = (dipendenti / tot_dip_area * 100) if tot_dip_area > 0 and pd.notna(dipendenti) else 0
        # --- NUOVI CALCOLI TERRITORIALI ---
        regione_grezza = str(riga.get(col_regione, 'N.D.'))
        regione_target_pulita = regione_grezza.split(' - ')[1] if ' - ' in regione_grezza else regione_grezza
        tot_imprese_regione = len(df_orbis[df_orbis[col_regione] == riga.get(col_regione)]) if col_regione else 0
        perc_imprese_regione = (tot_imprese_regione / tot_imprese_settore * 100) if tot_imprese_settore > 0 else 0
        tot_ricavi_macro_mgl = df_orbis[df_orbis['Macroregione'] == macroregione_target]['Totale valore della produzione migl EUR 2024'].sum()
        tot_ricavi_macro_mln = tot_ricavi_macro_mgl / 1000
        perc_ricavi_macroregione = (tot_ricavi_macro_mgl / tot_ricavi_settore * 100) if tot_ricavi_settore > 0 else 0
        perc_ricavi_target_su_macro = (ricavi_mgl / tot_ricavi_macro_mgl * 100) if tot_ricavi_macro_mgl > 0 and pd.notna(ricavi_mgl) else 0

        # =======================================================
        # --- NUOVI CALCOLI CATEGORIA (PER FORMA GIURIDICA) ---
        # =======================================================
        # Filtriamo il database prendendo SOLO le aziende con la stessa forma giuridica (es. solo le S.p.A.)
        df_categoria = df_orbis[df_orbis['Forma Giuridica Pulita'] == forma_giuridica]
        
        tot_ricavi_categoria = df_categoria['Totale valore della produzione migl EUR 2024'].sum()
        tot_attivo_categoria = df_categoria['Totale Attivo migl EUR 2024'].sum()
        
        perc_ricavi_categoria = (ricavi_mgl / tot_ricavi_categoria * 100) if tot_ricavi_categoria > 0 and pd.notna(ricavi_mgl) else 0
        perc_attivo_categoria = (attivo_mgl / tot_attivo_categoria * 100) if tot_attivo_categoria > 0 and pd.notna(attivo_mgl) else 0

        try: quartile_ricavi_target = pd.qcut(df_orbis['Totale valore della produzione migl EUR 2024'].dropna(), 4, labels=[1, 2, 3, 4]).loc[riga.name]
        except: quartile_ricavi_target = "N.D."
        try: quartile_attivo_target = pd.qcut(df_orbis['Totale Attivo migl EUR 2024'].dropna(), 4, labels=[1, 2, 3, 4]).loc[riga.name]
        except: quartile_attivo_target = "N.D."
    else:
        p_iva, forma_giuridica, macroregione_target, num_fg_target, perc_fg_target, ricavi_mln, attivo_mln, dipendenti = "N.D.", "N.D.", "N.D.", 0, 0, 0, 0, 0
        perc_ricavi_panel, perc_attivo_panel, perc_dip_area, quartile_ricavi_target, quartile_attivo_target = 0, 0, 0, "N.D.", "N.D."
        regione_target_pulita = "N.D."
        perc_imprese_regione, tot_ricavi_macro_mln, perc_ricavi_macroregione, perc_ricavi_target_su_macro = 0, 0, 0, 0
        perc_ricavi_categoria, perc_attivo_categoria = 0, 0  # <--- AGGIUNTA QUESTA RIGA!

    # 🧽 PULITURA RAGIONE SOCIALE
    ragione_sociale_pulita = str(azienda_target)
    sostituzioni_nomi = {
        "SOCIETA' A RESPONSABILITA' LIMITATA SEMPLIFICATA": "S.r.l.s.",
        "SOCIETÀ A RESPONSABILITÀ LIMITATA SEMPLIFICATA": "S.r.l.s.",
        "SOCIETA A RESPONSABILITA LIMITATA SEMPLIFICATA": "S.r.l.s.",
        "SOCIETA' A RESPONSABILITA' LIMITATA": "S.r.l.",
        "SOCIETÀ A RESPONSABILITÀ LIMITATA": "S.r.l.",
        "SOCIETA A RESPONSABILITA LIMITATA": "S.r.l.",
        "SOCIETA' PER AZIONI": "S.p.A.",
        "SOCIETÀ PER AZIONI": "S.p.A.",
        "SOCIETA PER AZIONI": "S.p.A.",
        "SOCIETA' IN NOME COLLETTIVO": "S.n.c.",
        "SOCIETÀ IN NOME COLLETTIVO": "S.n.c.",
        "SOCIETA IN NOME COLLETTIVO": "S.n.c.",
        "SOCIETA' IN ACCOMANDITA SEMPLICE": "S.a.s.",
        "SOCIETÀ IN ACCOMANDITA SEMPLICE": "S.a.s.",
        "SOCIETA IN ACCOMANDITA SEMPLICE": "S.a.s.",
        "SOCIETA' IN ACCOMANDITA PER AZIONI": "S.a.p.a.",
        "SOCIETÀ IN ACCOMANDITA PER AZIONI": "S.a.p.a.",
        "SOCIETA IN ACCOMANDITA PER AZIONI": "S.a.p.a.",
        "SOCIETA' COOPERATIVA": "Soc. Coop.",
        "SOCIETÀ COOPERATIVA": "Soc. Coop.",
        "SOCIETA COOPERATIVA": "Soc. Coop.",
        "SOCIETA' SEMPLICE": "S.s.",
        "SOCIETÀ SEMPLICE": "S.s.",
        "SOCIETA SEMPLICE": "S.s.",
        "SOCIETA' CONSORTILE": "Soc. Cons.",
        "SOCIETÀ CONSORTILE": "Soc. Cons.",
        "SOCIETA CONSORTILE": "Soc. Cons."
    }
    for lungo, corto in sostituzioni_nomi.items():
        ragione_sociale_pulita = re.sub(re.escape(lungo), corto, ragione_sociale_pulita, flags=re.IGNORECASE)

    # 🏷️ PULITURA NACE MULTIPLI
    settore_nace_str = str(settore_nace).strip()
    codici_estratti = []
    descrizioni_estratte = []
    matches = list(re.finditer(r'(\d{3,4})\s*-\s*(.*?)(?=\s*(?:[,;|]|\n)?\s*\d{3,4}\s*-|$)', settore_nace_str))

    if matches:
        for match in matches:
            cod_grezzo = match.group(1).strip()
            desc_testo = match.group(2).strip().rstrip(",;| ")
            cod_pulito = f"{cod_grezzo[:2]}.{cod_grezzo[2:]}"
            codici_estratti.append(cod_pulito)
            descrizioni_estratte.append(desc_testo)
            
        def formatta_italiano(lista):
            if len(lista) == 1: return lista[0]
            return ", ".join(lista[:-1]) + " e " + lista[-1]
            
        cod_nace_pulito = formatta_italiano(codici_estratti)
        desc_nace_pulita = formatta_italiano(descrizioni_estratte)
    else:
        cod_nace_pulito = "N.D."
        desc_nace_pulita = settore_nace_str

    # 🏛️ ESPANSIONE FORMA GIURIDICA (Per il testo discorsivo)
    dizionario_espansioni = {
        'S.r.l.s.': 'Società a Responsabilità Limitata Semplificata',
        'S.r.l.': 'Società a Responsabilità Limitata',
        'S.p.A.': 'Società per Azioni',
        'S.n.c.': 'Società in Nome Collettivo',
        'S.a.s.': 'Società in Accomandita Semplice',
        'S.a.p.a.': 'Società in Accomandita per Azioni',
        'Soc. Coop.': 'Società Cooperativa',
        'S.s.': 'Società Semplice',
        'Soc. Cons.': 'Società Consortile'
    }
    fg_espansa = dizionario_espansioni.get(forma_giuridica, forma_giuridica)


    # 🟢 FIX P.IVA: Ripristina gli zeri iniziali "mangiati" da Excel/Pandas
    if pd.notna(p_iva) and str(p_iva).strip().lower() not in ['n.d.', 'nan', '']:
        try:
            # 1. Lo passiamo a float e poi a int per eliminare eventuali ".0" finali
            p_iva_pulita = str(int(float(p_iva)))
            # 2. Aggiunge gli zeri a sinistra finché la lunghezza non torna a 11 cifre
            p_iva = p_iva_pulita.zfill(11)
        except ValueError:
            # Se dovesse contenere lettere (es. codici esteri o errori), lo lascia a testo
            p_iva = str(p_iva)
    else:
        p_iva = "n.d."

    # 🟢 MOTORE NARRATIVO: COMPOSIZIONE SOCIETARIA DETTAGLIATA (Versione Ultra-Resistente)
    target_fg = str(forma_giuridica).strip()

    if target_fg.lower() == fg_1_name.lower():
        # Caso A: L'azienda in analisi fa parte della Maggioranza #1
        testo_fg = f"rappresenta la veste giuridica dominante assoluta del comparto, costituendo da sola il {format_euro(fg_1_perc)}% delle realtà censite ({f'{fg_1_num:,}'.replace(',', '.')} unità)."
        
        # Se esiste una seconda forma giuridica, continuiamo la frase
        if fg_2_name:
            testo_fg = testo_fg.rstrip(".") + f", mentre la restante parte del mercato è composta in larga misura da {fg_2_name} ({format_euro(fg_2_perc)}%, {f'{fg_2_num:,}'.replace(',', '.')} unità)"
            if fg_altre_num > 0:
                testo_fg += f" e, in minor parte, da altre configurazioni societarie miste ({format_euro(fg_altre_perc)}%, {f'{fg_altre_num:,}'.replace(',', '.')} unità)."
            else:
                testo_fg += "."
    else:
        # Caso B: L'azienda in analisi NON è la Maggioranza #1
        testo_fg = f"si inserisce in un comparto caratterizzato in larga parte da {fg_1_name}, struttura che controlla il {format_euro(fg_1_perc)}% delle componenti societarie ({f'{fg_1_num:,}'.replace(',', '.')} unità). "
        
        if fg_2_name and target_fg.lower() == fg_2_name.lower():
            # Il target è esattamente il #2
            testo_fg += f"A seguire si posiziona proprio la veste legale dell'azienda in analisi ({target_fg}), che rappresenta il {format_euro(fg_2_perc)}% del panel ({f'{fg_2_num:,}'.replace(',', '.')} unità)"
            if fg_altre_num > 0:
                testo_fg += f", affiancata in minor misura da altre configurazioni societarie miste ({format_euro(fg_altre_perc)}%, {f'{fg_altre_num:,}'.replace(',', '.')} unità)."
            else:
                testo_fg += "."
        elif fg_2_name:
            # Il target è dal #3 in giù
            testo_fg += f"A seguire troviamo una forte presenza di {fg_2_name} ({format_euro(fg_2_perc)}%, {f'{fg_2_num:,}'.replace(',', '.')} unità), mentre la veste societaria dell'azienda in analisi si colloca nel restante {format_euro(fg_altre_perc)}% del mercato, insieme ad altre configurazioni minoritarie."
        else:
            # Fail-safe nel caso ci fossero anomalie strane
            testo_fg += f"L'azienda in analisi si inserisce in questo contesto con una quota del {format_euro(perc_fg_target)}%."

    # ----------------------------------------------------
    # COSTRUZIONE DEL DIZIONARIO (Dati per il Word)
    context = {
        'ragione_sociale': ragione_sociale_pulita, 
        'codice_nace': cod_nace_pulito, 
        'descr_settore': desc_nace_pulita, 
        'partita_iva': p_iva, 
        'forma_giuridica': forma_giuridica,
        'forma_giuridica_espansa': fg_espansa,
        'testo_forma_giuridica_strutturato': testo_fg,
        'perc_fg': format_euro(perc_fg_target), 'num_fg': f"{num_fg_target:,}".replace(',', '.'),
        'fg_maggioranza': fg_maggioranza, 'num_fg_maggioranza': f"{num_fg_maggioranza:,}".replace(',', '.'),
        'perc_fg_maggioranza': format_euro(perc_fg_maggioranza),
        'classe_dimensionale': 'Grande Impresa' if ricavi_mln > 50 else ('Media Impresa' if ricavi_mln > 10 else 'Piccola Impresa'),
        'ricavi_mln': format_euro(ricavi_mln),
        'perc_ricavi_panel': format_euro(perc_ricavi_panel),
        'perc_ricavi_categoria': format_euro(perc_ricavi_categoria), # <-- Ora calcola il VERO dato!
        'quartile_ricavi': quartile_ricavi_target,
        'attivo_mln': format_euro(attivo_mln),
        'perc_attivo_panel': format_euro(perc_attivo_panel),
        'perc_attivo_categoria': format_euro(perc_attivo_categoria), # <-- Aggiunto anche per l'Attivo!
        'quartile_attivo': quartile_attivo_target,
        'perc_dip_area': format_euro(perc_dip_area), 'macroregione': macroregione_target,
        'regione_target': regione_target_pulita,
        'perc_imprese_regione': format_euro(perc_imprese_regione),
        'perc_ricavi_macroregione': format_euro(perc_ricavi_macroregione),
        'tot_ricavi_macro_mln': format_euro(tot_ricavi_macro_mln),
        'perc_ricavi_target_su_macro': format_euro(perc_ricavi_target_su_macro),
        'mg_ebitda': mg_ebitda, 'rnk_naz_ebitda': rnk_naz_ebitda, 'rnk_reg_ebitda': rnk_reg_ebitda,
        'mg_ebit': mg_ebit, 'rnk_naz_ebit': rnk_naz_ebit, 'rnk_reg_ebit': rnk_reg_ebit,
        'mg_prof': mg_prof, 'rnk_naz_prof': rnk_naz_prof, 'rnk_reg_prof': rnk_reg_prof,
        'ind_str1': ind_str1, 'rnk_naz_strut1': rnk_naz_strut1, 'rnk_reg_strut1': rnk_reg_strut1,
        'ind_str2': ind_str2, 'rnk_naz_strut2': rnk_naz_strut2, 'rnk_reg_strut2': rnk_reg_strut2,
        'gearing': gearing, 'rnk_naz_gear': rnk_naz_gear, 'rnk_reg_gear': rnk_reg_gear,
        'ind_rot_cap': ind_rot, 'rnk_naz_rot': rnk_naz_rot, 'rnk_reg_rot': rnk_reg_rot,
        'ind_cr': ind_cr, 'rnk_naz_cr': rnk_naz_cr, 'rnk_reg_cr': rnk_reg_cr,
        'ind_qr': ind_qr, 'rnk_naz_qr': rnk_naz_qr, 'rnk_reg_qr': rnk_reg_qr,
        'tot_imprese': f"{tot_imprese_settore:,}".replace(',', '.'), 'tot_ricavi': format_euro(tot_ricavi_settore),
        'tot_attivo': format_euro(tot_attivo_settore), 'tot_dipendenti': format_euro(tot_dipendenti_settore, 0),
        
        'perc_ricavi_target_su_macro': format_euro(perc_ricavi_target_su_macro),
        'impatto_territoriale': 'N.D.',

        # Variabili di default (sovrascritte dal motore sottostante)
        'rating_eco': 'N.D.', 'rating_patr': 'N.D.', 'rating_fin': 'N.D.', 'rating_tot': 'N.D.', 'rating_comb': 'N.D.',
        'rating_piu_presente': 'N.D.', 'assoc_desc_rating_lett': 'N.D.',
        'rat1_piu_pres_num': 'N.D.', 'rat1_piu_pres_categ': 'N.D.', 
        'rat2_piu_pres_num': 'N.D.', 'rat2_piu_pres_categ': 'N.D.',
        'rat3_piu_pres_num': 'N.D.', 'rat3_piu_pres_categ': 'N.D.', 
        'rating_piu_pres': 'N.D.', 'rating_piu_pres_num_tot': 'N.D.',
        'num_max_soc': 'N.D.', 'num_soc_valide': 'N.D.', 'perc_su_istat': '100', 'max_soc_istat': 'N.D.',
        'tab_territorio': [], 'tab_bench_territorio': []
    }

    # =================================================================
    # 📊 MOTORE DI CALCOLO NATIVO DEI RATING (Senza Excel!)
    # =================================================================
    # *** MODIFICA RISPETTO ALL'ORIGINALE ***
    # Il punteggio ora rispecchia il rango del terzile:
    #   1 = primo terzile  = valori più alti (MIGLIORI) per metriche dirette
    #   2 = secondo terzile = valori medi
    #   3 = terzo terzile  = valori più bassi (PEGGIORI)
    # Nell'originale era invertito: 3 = migliore, 1 = peggiore.
    # Il default NaN diventa 3 (peggiore) invece di 1.
    def punteggio_diretto(val, t1, t2):
        if pd.isna(val): return 3  # MODIFICATO: era return 1
        return 1 if val >= t2 else (2 if val >= t1 else 3)  # MODIFICATO: era 3/.../1

    def punteggio_inverso(val, t1, t2):
        if pd.isna(val): return 3  # MODIFICATO: era return 1
        return 1 if val <= t1 else (2 if val <= t2 else 3)  # MODIFICATO: era 3/.../1

    # MODIFICATO: due funzioni distinte perché assegna_lettera è usata in due contesti:
    # 1. assegna_lettera_area → per pts_eco/fin/pat (somma di 1/2/3 con 1=best, basso=buono)
    # 2. assegna_lettera      → per pts_totali/Benchmark Totale (valori_lettere A=3, alto=buono)
    def assegna_lettera_area(punti):  # NUOVO: soglia invertita, solo per somme di terzili area
        if pd.isna(punti): return 'C'
        return 'A' if punti <= 4 else ('B' if punti <= 7 else 'C')

    def assegna_lettera(punti):  # INVARIATA rispetto all'originale: usata per Benchmark Totale
        if pd.isna(punti): return 'C'
        return 'A' if punti >= 8 else ('B' if punti >= 5 else 'C')

    # Identificatori metriche
    c_prof = 'Margine di Profitto (*) % 2024'
    c_ebitda = 'Margine EBITDA (*) % 2024'
    c_ebit = 'Margine EBIT (*) % 2024'
    c_rot = 'Indice di Rotazione del Capitale Investito (*) 2024'
    c_quick = 'Quick Ratio (*) 2024'
    c_curr = 'Current Ratio (*) 2024'
    c_str1 = 'Indice di Struttura 1° livello (*) 2024'
    c_str2 = 'Indice di Struttura 2° livello (*) 2024'
    c_gear = 'Gearing (*) % 2024'

    metriche_dirette = [c_prof, c_ebitda, c_ebit, c_rot, c_quick, c_curr, c_str1, c_str2]
    metriche_inverse = [c_gear]

    df_rating = df_orbis.copy()

    # 1. Calcolo matematico dei Terzili e Assegnazione Punti
    for m in metriche_dirette + metriche_inverse:
        if m in df_rating.columns:
            t1 = df_rating[m].quantile(1/3)
            t2 = df_rating[m].quantile(2/3)
            if m in metriche_inverse:
                df_rating[f'pts_{m}'] = df_rating[m].apply(lambda x: punteggio_inverso(x, t1, t2))
            else:
                df_rating[f'pts_{m}'] = df_rating[m].apply(lambda x: punteggio_diretto(x, t1, t2))
        else:
            df_rating[f'pts_{m}'] = 3  # Fallback: colonna assente → terzo terzile (MODIFICATO: era 1)

    # 2. Somma Punti per Area (Economico, Finanziario, Patrimoniale)
    df_rating['pts_eco'] = df_rating[f'pts_{c_prof}'] + df_rating[f'pts_{c_ebitda}'] + df_rating[f'pts_{c_ebit}']
    df_rating['pts_fin'] = df_rating[f'pts_{c_rot}'] + df_rating[f'pts_{c_quick}'] + df_rating[f'pts_{c_curr}']
    df_rating['pts_pat'] = df_rating[f'pts_{c_str1}'] + df_rating[f'pts_{c_str2}'] + df_rating[f'pts_{c_gear}']

    # 3. Assegnazione Lettere (A, B, C) — usa assegna_lettera_area (soglia invertita)
    df_rating['Rating Economico'] = df_rating['pts_eco'].apply(assegna_lettera_area)
    df_rating['Rating Finanziario'] = df_rating['pts_fin'].apply(assegna_lettera_area)
    df_rating['Rating Patrimoniale'] = df_rating['pts_pat'].apply(assegna_lettera_area)

    # 4. Rating Combinato Testuale (es: AAA, BBB, ABC)
    df_rating['Rating Combinato'] = df_rating['Rating Economico'] + df_rating['Rating Patrimoniale'] + df_rating['Rating Finanziario']

    # 5. Benchmark Totale (Media delle Lettere)
    valori_lettere = {'A': 3, 'B': 2, 'C': 1}
    df_rating['pts_totali'] = df_rating['Rating Economico'].map(valori_lettere) + \
                              df_rating['Rating Finanziario'].map(valori_lettere) + \
                              df_rating['Rating Patrimoniale'].map(valori_lettere)
    df_rating['Benchmark Totale'] = df_rating['pts_totali'].apply(assegna_lettera)

    
    # =================================================================
    # 🎯 ESTRAZIONE E INIEZIONE NEL TEMPLATE WORD
    # =================================================================
    
    # Preleviamo la riga esatta dell'azienda target
    df_target_rating = df_rating[df_rating[col_ragione].astype(str).str.lower().str.contains(azienda_target.lower().strip(), regex=False, na=False)]

    if not df_target_rating.empty:
        riga_t = df_target_rating.iloc[0]
        context['rating_eco'] = riga_t['Rating Economico']
        context['rating_fin'] = riga_t['Rating Finanziario']
        context['rating_patr'] = riga_t['Rating Patrimoniale']
        context['rating_tot'] = riga_t['Rating Combinato']
        context['rating_comb'] = riga_t['Benchmark Totale']

        context['dip_target'] = f"{int(dipendenti):,}".replace(',', '.') if pd.notna(dipendenti) else "n.d."
       
        # Calcoli matematici reali per il Box Economico
        r_eco_target = riga_t['Rating Economico']
        num_eco_fascia_real = len(df_rating[df_rating['Rating Economico'] == r_eco_target])
        perc_eco_fascia_real = (num_eco_fascia_real / tot_imprese_settore) * 100 if tot_imprese_settore > 0 else 0
       
        context['nr_rating_eco'] = f"{num_eco_fascia_real:,}".replace(',', '.')
        context['perc_rating_eco'] = format_euro(perc_eco_fascia_real)
        context['perc_rating_eco_parte'] = format_euro((1 / num_eco_fascia_real) * 100) if num_eco_fascia_real > 0 else "0,00"
       
        # Calcoli matematici reali per il Box Patrimoniale
        r_patr_target = riga_t['Rating Patrimoniale']
        num_patr_fascia_real = len(df_rating[df_rating['Rating Patrimoniale'] == r_patr_target])
        perc_patr_fascia_real = (num_patr_fascia_real / tot_imprese_settore) * 100 if tot_imprese_settore > 0 else 0
       
        context['nr_rating_patr'] = f"{num_patr_fascia_real:,}".replace(',', '.')
        context['perc_rating_patr'] = format_euro(perc_patr_fascia_real)
        context['perc_rating_patr_parte'] = format_euro((1 / num_patr_fascia_real) * 100) if num_patr_fascia_real > 0 else "0,00"

        # Calcoli matematici reali per il Box Finanziario
        r_fin_target = riga_t['Rating Finanziario']
        num_fin_fascia_real = len(df_rating[df_rating['Rating Finanziario'] == r_fin_target])
        perc_fin_fascia_real = (num_fin_fascia_real / tot_imprese_settore) * 100 if tot_imprese_settore > 0 else 0
       
        context['nr_rating_fin'] = f"{num_fin_fascia_real:,}".replace(',', '.')
        context['perc_rating_fin'] = format_euro(perc_fin_fascia_real)
        context['perc_rating_fin_parte'] = format_euro((1 / num_fin_fascia_real) * 100) if num_fin_fascia_real > 0 else "0,00"


    # --- 🤖 MOTORE NARRATIVO (Testi Dinamici e Tecnici in base ai Dati Reali) --- 

    def get_impatto_territoriale(perc, nome, ricavi_formattati):
        if perc >= 5.0:
            return f"Con ricavi pari a {ricavi_formattati} mln di EUR, {nome} incide in maniera determinante sulla creazione di ricchezza locale, confermandosi un player di assoluto riferimento sul piano territoriale grazie a un impatto pari al {format_euro(perc)}% rispetto al totale dei ricavi dell'area."
        elif perc >= 1.0:
            return f"Con ricavi pari a {ricavi_formattati} mln di EUR, {nome} fornisce un contributo significativo alla creazione di ricchezza locale, consolidando una posizione di rilievo sul piano territoriale con un'incidenza pari al {format_euro(perc)}% rispetto ai ricavi complessivi dell'area."
        elif perc >= 0.1:
            return f"Con ricavi pari a {ricavi_formattati} mln di EUR, {nome} partecipa attivamente al tessuto economico locale, rappresentando una stabile realtà territoriale con un'incidenza pari al {format_euro(perc)}% rispetto ai ricavi complessivi dell'area."
        else:
            return f"Con ricavi pari a {ricavi_formattati} mln di EUR, {nome} opera all'interno di un mercato territoriale ampio e competitivo, contribuendo al tessuto economico locale con un'incidenza pari al {format_euro(perc)}% rispetto ai ricavi complessivi dell'area."

    def get_testo_totale(rating):
        if rating == 'A':
            return "Questo risultato riflette un profilo economico-patrimoniale solido, con indicatori di redditività, indipendenza finanziaria ed equilibrio di breve termine superiori ai parametri mediani del settore."
        elif rating == 'B':
            return "Questo risultato riflette un profilo di sostanziale stabilità complessiva, evidenziando un assetto economico-patrimoniale solido pur con margini di miglioramento specifici nell'ottimizzazione degli asset e dei margini operativi."
        elif rating == 'C':
            return "Questo risultato evidenzia elementi di vulnerabilità nella struttura economica, patrimoniale e finanziaria, suggerendo l'opportunità di interventi mirati per stabilizzare i flussi di cassa, riequilibrare le fonti di finanziamento o recuperare marginalità operativa."
        return "Non sono disponibili dati sufficienti per elaborare un giudizio complessivo accurato."

    def get_testo_eco(rating):
        if rating == 'A':
            return "I margini operativi (EBITDA ed EBIT) e il margine di profitto si collocano stabilmente al di sopra dei parametri mediani di settore. L'impresa mostra una buona efficienza nel generare reddito dalla gestione caratteristica e nel trasformare i ricavi in risultato netto."
        elif rating == 'B':
            return "La redditività operativa e netta risulta adeguata alle dinamiche di settore. L'azienda presenta una buona capacità di generare reddito operativo, seppur con spazi di ottimizzazione nell'assorbimento dei costi di gestione e degli oneri accessori per incrementare l'efficienza della struttura economica."
        elif rating == 'C':
            return "L'analisi dei margini segnala una minore capacità di trasformare i ricavi in risultato operativo e netto. Il posizionamento al di sotto dei parametri settoriali denota una minore efficienza della struttura economica e un'elevata incidenza dei costi di gestione."
        return "Dati economici non disponibili o insufficienti."

    def get_testo_patr(rating):
        if rating == 'A':
            return "La struttura patrimoniale rappresenta un punto di solidità dell'azienda. Il capitale proprio copre interamente gli investimenti a lungo termine (immobilizzazioni), garantendo una buona indipendenza dai vincoli di rimborso a breve termine."
        elif rating == 'B':
            return "La struttura patrimoniale appare equilibrata. L'azienda finanzia una parte adeguata delle proprie immobilizzazioni attraverso capitale permanente, sebbene il ricorso al capitale di terzi per sostenere gli impieghi a medio-lungo termine debba essere costantemente monitorato."
        elif rating == 'C':
            return "La struttura patrimoniale evidenzia uno squilibrio nella correlazione temporale tra fonti e impieghi. Una porzione significativa delle immobilizzazioni risulta finanziata mediante capitale di terzi con obbligo di rimborso nel breve termine, esponendo l'azienda a potenziali rischi di rifinanziamento."
        return "Dati patrimoniali non disponibili o insufficienti."

    def get_testo_fin(rating):
        if rating == 'A':
            return "L'equilibrio della struttura finanziaria di breve termine è solido. Le attività correnti coprono integralmente i debiti esigibili nell'esercizio e le risorse prontamente liquidabili assicurano una buona solvibilità immediata, senza necessità di smobilizzare le scorte."
        elif rating == 'B':
            return "L'equilibrio della struttura finanziaria di breve termine è sufficiente a coprire gli impieghi correnti. Tuttavia, la liquidità immediata potrebbe presentare una parziale dipendenza dalla monetizzazione delle scorte o dall'incasso dei crediti per soddisfare tutti gli impegni a breve."
        elif rating == 'C':
            return "La struttura finanziaria registra potenziali tensioni di liquidità. L'incapacità delle attività correnti o prontamente liquidabili di far fronte agevolmente alle passività correnti segnala il rischio di dover ricorrere a ulteriori fonti di finanziamento esterne."
        return "Dati finanziari non disponibili o insufficienti."

    def get_testo_sintesi(rating):
        if rating == 'A':
            return "In sintesi, l'integrazione di questi equilibri mostra un'impresa con una struttura economica, patrimoniale e finanziaria solida. L'efficienza nell'impiego del capitale investito e la solidità delle fonti di finanziamento supportano i futuri percorsi di crescita."
        elif rating == 'B':
            return "In sintesi, il coordinamento tra struttura patrimoniale e liquidità finanziaria assicura all'impresa la continuità operativa. Risulta tuttavia opportuno monitorare la capacità del capitale investito di tradursi in ricavi, al fine di non erodere i margini di sicurezza nel lungo periodo."
        elif rating == 'C':
            return "In sintesi, la combinazione dei tre equilibri rivela inefficienze nella struttura economica, patrimoniale e finanziaria che penalizzano l'impiego del capitale e l'autonomia monetaria. È opportuno intervenire per ripristinare un migliore allineamento temporale tra le fonti di finanziamento e gli impieghi."
        return ""

    def get_intro_benchmark_eco(rating):
        if rating == 'A':
            return f'Rispetto al Benchmark Economico, la valutazione "{rating}" evidenzia una redditività operativa e netta superiore ai parametri mediani del comparto di riferimento:'
        elif rating == 'B':
            return f'Rispetto al Benchmark Economico, la valutazione "{rating}" evidenzia una capacità di generare reddito adeguata e allineata ai valori mediani del comparto di riferimento, pur con specifici margini di intervento:'
        elif rating == 'C':
            return f'Rispetto al Benchmark Economico, la valutazione "{rating}" sottolinea una marginalità operativa e netta contratta, al di sotto dei parametri mediani espressi dal comparto di riferimento:'
        return "Rispetto al Benchmark Economico, i dati a disposizione non consentono di esprimere una valutazione completa, come riassunto di seguito:"

    # =================================================================
    # 🟢 INDICATORI ECONOMICI (Valore vs Mediana) - Formattati a Bullet Points
    # =================================================================
    def get_analisi_margini_operativi(az_ebitda, set_ebitda, az_ebit, set_ebit, descr_settore):
        if az_ebitda < set_ebitda:
            txt_ebitda = f"• Il **Margine EBITDA** ({format_euro(az_ebitda)}%) risulta inferiore alla mediana settoriale ({format_euro(set_ebitda)}%). Valori più contenuti segnalano una minore capacità di trasformare i ricavi in margine operativo lordo, denotando una minore efficienza della gestione caratteristica prima degli ammortamenti e delle svalutazioni.\n"
        else:
            txt_ebitda = f"• Il **Margine EBITDA** ({format_euro(az_ebitda)}%) supera la mediana settoriale ({format_euro(set_ebitda)}%). Valori più elevati indicano una maggiore capacità dell'impresa di generare reddito dalla gestione caratteristica prima di ammortamenti e svalutazioni, evidenziando una superiore efficienza operativa.\n"

        if az_ebit < set_ebit:
            txt_ebit = f"• Il **Margine EBIT** ({format_euro(az_ebit)}%) si colloca al di sotto del target mediano ({format_euro(set_ebit)}%). Tale dato segnala una minore capacità di generare reddito operativo in relazione ai ricavi conseguiti a valle dell'assorbimento dei costi e degli ammortamenti."
        else:
            txt_ebit = f"• Il **Margine EBIT** ({format_euro(az_ebit)}%) si posiziona al di sopra della mediana di settore ({format_euro(set_ebit)}%). Questo livello indica una maggiore capacità di conseguire un risultato operativo soddisfacente dopo aver considerato gli ammortamenti e le svalutazioni."

        return f"Le risultanze relative al mercato ({descr_settore}) evidenziano che:\n{txt_ebitda}{txt_ebit}"

    def get_analisi_margine_profitto(az_prof, set_prof, descr_settore):
        if az_prof < set_prof:
            return f"\n• Il **Margine di Profitto** ({format_euro(az_prof)}%) risulta inferiore alla mediana settoriale ({format_euro(set_prof)}%). Tale andamento denota una minore capacità di trasformare i ricavi in utile netto, evidenziando criticità nell'assorbimento della gestione straordinaria, degli oneri finanziari o del carico fiscale."
        else:
            return f"\n• Il **Margine di Profitto** ({format_euro(az_prof)}%) supera il parametro mediano del settore ({format_euro(set_prof)}%). Valori più elevati indicano una maggiore capacità dell'impresa di convertire i ricavi in risultato netto finale, confermando una gestione ottimizzata degli oneri extra-caratteristici."

    # =================================================================
    # 🟠 INDICATORI PATRIMONIALI (Valore vs 1 e Gearing vs Mediana)
    # =================================================================
    def get_intro_benchmark_patr(rating):
        if rating == 'A':
            return "L'analisi attesta una buona robustezza nella copertura degli investimenti, posizionando gli indici strutturali dell'azienda al di sopra delle soglie di sicurezza:"
        elif rating == 'B':
            return "L'analisi evidenzia un rapporto tra capitale, debiti e immobilizzazioni adeguato e in linea con i parametri mediani del mercato di riferimento:"
        elif rating == 'C':
            return "L'analisi segnala squilibri nella correlazione temporale tra le fonti di copertura e le immobilizzazioni aziendali rispetto ai parametri di sicurezza:"
        return "L'analisi non consente di esprimere una valutazione completa sui rischi a lungo termine a causa di dati insufficienti:"

    def get_analisi_indici_struttura(az_str1, set_str1, az_str2, set_str2):
        if az_str1 >= 1:
            txt_s1 = f"• L'**Indice primario di struttura** ({format_euro(az_str1)}), superiore o uguale all'unità, indica che il capitale proprio, il quale non ha vincoli di scadenza, ha finanziato interamente le immobilizzazioni, caratterizzate da tempi di disinvestimento medio-lunghi.\n"
        else:
            txt_s1 = f"• L'**Indice primario di struttura** ({format_euro(az_str1)}), risultando inferiore ad uno, segnala che una parte delle immobilizzazioni è stata finanziata mediante capitale di terzi, con potenziale obbligo di rimborso nel breve termine.\n"

        if az_str2 >= 1:
            txt_s2 = f"• L'**Indice secondario di struttura** ({format_euro(az_str2)}), superiore o uguale all'unità, conferma che il capitale permanente, costituito dal capitale proprio e dai debiti a medio-lunga scadenza, ha finanziato interamente gli asset immobilizzati."
        else:
            txt_s2 = f"• L'**Indice secondario di struttura** ({format_euro(az_str2)}), essendo inferiore ad uno, indica che una parte dell'attivo immobilizzato è finanziata attraverso capitale di terzi a breve scadenza, determinando uno squilibrio temporale tra fonti e impieghi."

        return f"Analizzando la provvista allargata:\n{txt_s1}{txt_s2}"

    def get_analisi_gearing(az_gear, set_gear):
        if az_gear <= set_gear:
            return f"\n• Il **Gearing** ({format_euro(az_gear)}%) si attesta al di sotto del parametro mediano del comparto ({format_euro(set_gear)}%). Tali valori più contenuti indicano una limitata dipendenza dell'impresa dall'indebitamento oneroso e una solida autonomia rispetto ai creditori."
        else:
            return f"\n• Il **Gearing** ({format_euro(az_gear)}%) supera la mediana di settore ({format_euro(set_gear)}%). Valori più elevati segnalano un maggiore ricorso al capitale di terzi per il finanziamento aziendale, determinando un incremento del rischio finanziario e una minore autonomia."

    # =================================================================
    # 🔵 INDICATORI FINANZIARI (Valore vs 1 e Rotazione vs Mediana)
    # =================================================================
    def get_intro_benchmark_fin(rating):
        if rating == 'A':
            return "L'analisi degli indicatori correnti evidenzia una condizione di stabilità nel breve periodo e una buona capacità di generare ricavi tramite l'impiego delle risorse, superiore ai parametri mediani del settore:"
        elif rating == 'B':
            return "L'analisi degli indicatori correnti evidenzia un equilibrio finanziario di breve termine adeguato a coprire le passività correnti, in linea con gli standard del settore:"
        elif rating == 'C':
            return "L'analisi degli indicatori correnti segnala potenziali tensioni di liquidità e un impiego meno efficiente delle risorse rispetto alle soglie di sicurezza:"
        return "L'analisi non consente di esprimere una valutazione completa sulla gestione della liquidità:"

    def get_analisi_current_ratio(az_cr, set_cr):
        if az_cr >= 1:
            return f"• Il **Current Ratio** ({format_euro(az_cr)}), essendo superiore o uguale all'unità, indica che le attività a breve termine sono sufficienti a coprire integralmente i debiti esigibili nel breve periodo, evidenziando una situazione di equilibrio d'esercizio.\n"
        else:
            return f"• Il **Current Ratio** ({format_euro(az_cr)}), risultando inferiore ad uno, segnala l'incapacità delle attività correnti di far fronte alle passività correnti, configurando una potenziale tensione di liquidità all'interno della struttura d'esercizio.\n"

    def get_analisi_quick_ratio(az_qr, set_qr):
        if az_qr >= 1:
            return f"• Il **Quick Ratio** ({format_euro(az_qr)}), superiore o uguale all'unità, indica che le risorse prontamente liquidabili sono sufficienti a garantire la copertura dei debiti a breve termine senza ricorrere alla vendita forzata delle rimanenze di magazzino.\n"
        else:
            return f"• Il **Quick Ratio** ({format_euro(az_qr)}), essendo inferiore ad uno, evidenzia una dipendenza, almeno parziale, dalla monetizzazione delle scorte o da ulteriori fonti di finanziamento esterne per soddisfare gli impegni immediati.\n"

    def get_analisi_rotazione(az_rot, set_rot, descr_settore):
        if az_rot < set_rot:
            return f"• L'**Indice di rotazione del capitale investito** ({format_euro(az_rot)}) risulta inferiore alla mediana del comparto ({format_euro(set_rot)}). Valori più contenuti segnalano una minore capacità del capitale investito di tradursi in ricavi, denotando un impiego meno efficiente degli asset operativi."
        else:
            return f"• L'**Indice di rotazione del capitale investito** ({format_euro(az_rot)}) supera la mediana settoriale ({format_euro(set_rot)}). Valori più elevati indicano una maggiore capacità dell'impresa di generare ricavi attraverso le risorse investite, evidenziando un efficiente utilizzo del capitale."

    def get_analisi_posizionamento_fin(az_cr, az_qr, set_cr, set_qr):
        # La frase conclusiva sotto ai 3 bullet point finanziari
        if az_cr >= set_cr and az_qr >= set_qr:
            return "Il posizionamento finale nelle griglie distributive consolida un profilo di elevata affidabilità finanziaria, registrando indicatori correnti sistematicamente allineati o superiori ai parametri dei competitor caratteristici."
        elif az_cr < set_cr and az_qr < set_qr:
            return "Il posizionamento finale colloca la società nelle fasce contratte della distribuzione settoriale. Gli indicatori riflettono la necessità di velocizzare la rotazione degli asset circolanti per alleviare il potenziale rischio di insolvenza nel breve periodo."
        else:
            return "Il posizionamento finale riflette risultanze asimmetriche nel panel di settore. Sebbene la solvibilità generale risulti presidiata in linea con i valori mediani, permangono mirati elementi di sfasamento monetario sul ciclo di liquidazione immediata delle rimanenze."

    def get_analisi_combinata(eco, patr, fin):
        dict_eco = {'A': "un'ottima marginalità operativa", 'B': "una redditività caratteristica in linea col mercato", 'C': "una debole capacità di trasformare i ricavi in margini"}
        dict_patr = {'A': "un'indipendenza finanziaria che copre le immobilizzazioni", 'B': "un capitale permanente adeguato agli asset", 'C': "un forte squilibrio nell'indebitamento a sostegno degli impieghi"}
        dict_fin = {'A': "una rotazione ottimale per coprire le passività a breve", 'B': "un Current Ratio adeguato a onorare i debiti esigibili", 'C': "un Quick Ratio dipendente dallo smobilizzo delle rimanenze"}

        if eco in ['A','B','C'] and patr in ['A','B','C'] and fin in ['A','B','C']:
            return f"riflette un'azienda che poggia su {dict_patr[patr]}, associata a {dict_eco[eco]} e {dict_fin[fin]}."
        return "riflette un'azienda per la quale non è possibile tracciare un profilo combinato completo a causa di dati mancanti."

    def get_descr_fascia_appartenenza(rating):
        if rating == 'A': return "la solidità reddituale, patrimoniale e finanziaria"
        elif rating == 'B': return "l'adeguatezza agli indici di stabilità settoriale"
        elif rating == 'C': return "la fascia che necessita di consolidamento per le tensioni su margini e liquidità"
        return "una porzione non classificabile"

    def get_intro_divario_strutturale(az_ebitda, set_ebitda):
        if az_ebitda >= set_ebitda:
            return "un posizionamento favorevole in termini di ottimizzazione della gestione operativa e di struttura economica da parte di"
        else:
            return "un significativo divario prestazionale tra la mediana del comparto e l'efficienza della struttura economica di"

    # =========================================================================================
    # LE ALTRE FUNZIONI NARRATIVE PER LA PARTE DI TREND STORICO (Restano invariate o snellite)
    # =========================================================================================
    def get_analisi_trend_ebitda(az_ebitda, set_ebitda):
        if az_ebitda >= set_ebitda:
            return f"ha registrato performance favorevoli, con un valore al 2024 pari al {format_euro(az_ebitda)}%. Questo risultato dimostra una maggiore capacità della struttura economica di generare reddito dalla gestione caratteristica, superando la mediana del settore ({format_euro(set_ebitda)}%)."
        else:
            return f"ha registrato evidenti segnali di contrazione, con un valore al 2024 pari al {format_euro(az_ebitda)}%. Questo risultato si colloca al di sotto della mediana settoriale ({format_euro(set_ebitda)}%), segnalando una minore efficienza nell'assorbimento dei costi correnti."

    def get_asimmetria_ebitda(az_ebitda, set_ebitda):
        if az_ebitda >= set_ebitda: return "una spiccata forza del Margine EBITDA dell'azienda rispetto ai parametri mediani di settore."
        else: return "una chiara criticità nella tenuta della struttura economica aziendale rispetto al benchmark mediano del comparto."

    def get_confronto_costi_settore(az_ebitda, set_ebitda):
        if az_ebitda >= set_ebitda: return "ottimizzata e strutturalmente protetta rispetto ai parametri economici medi del mercato di riferimento"
        else: return "maggiormente gravosa e penalizzante rispetto ai parametri economici mediani del mercato di riferimento"

    def get_impatto_costi_su_margine(az_ebitda, set_ebitda):
        if az_ebitda >= set_ebitda: return "evidenzia una spiccata efficienza nell'assorbimento dei costi operativi correnti a salvaguardia del margine lordo caratteristico"
        else: return "segnala una minore efficienza nell'assorbimento dei costi industriali interni, comprimendo il potenziale residuo del Margine EBITDA"

    def get_analisi_trend_ebit(az_ebit, set_ebit):
        if az_ebit >= set_ebit: return f"Questo dato dimostra un'elevata redditività dopo aver considerato gli ammortamenti e le svalutazioni, con un valore pari al {format_euro(az_ebit)}% che supera stabilmente la mediana."
        else: return f"Questo dato riflette una minore capacità di conseguire un risultato operativo soddisfacente in rapporto ai ricavi di vendita conseguiti, posizionandosi sotto la mediana settoriale al {format_euro(az_ebit)}%."

    def get_confronto_ebit_settore(az_ebit, set_ebit):
        if az_ebit >= set_ebit: return "L'efficienza nell'impiego operativo delle risorse aziendali è confermata dal posizionamento nettamente superiore alla mediana dell'EBIT."
        else: return "La capacità di generare reddito operativo a valle dei costi fissi si attesta al di sotto dei livelli di riferimento mediani espressi dal settore in esame."

    def get_sintesi_ebit_mediana(az_ebit, set_ebit):
        if az_ebit >= set_ebit: return f"al di sopra della mediana di settore (pari a {format_euro(set_ebit)}%) per quanto concerne l'utile operativo EBIT"
        else: return f"al di sotto della mediana di settore (pari a {format_euro(set_ebit)}%) per quanto concerne l'utile operativo EBIT"

    def get_analisi_verticale_ebit(az_ebitda, az_ebit, set_ebitda, set_ebit):
        spread_azienda = az_ebitda - az_ebit
        spread_settore = set_ebitda - set_ebit
        if spread_azienda > spread_settore + 3.0: return "un cosiddetto 'effetto forbice' più marcato rispetto al settore: le quote di ammortamento e svalutazione erodono la redditività generata dai ricavi in misura superiore alla media di comparto, appesantendo la struttura economica a valle della gestione lorda caratteristica"
        else: return "un bilanciamento della struttura economica in linea con il settore: l'incidenza degli ammortamenti sul Valore della Produzione risulta paragonabile a quella dei competitor, senza un ulteriore assottigliamento della redditività nel passaggio da EBITDA a EBIT"

    def get_efficienza_struttura_asset(az_ebit, set_ebit):
        if az_ebit >= set_ebit: return "efficiente nell'impiego, nel rinnovo e nella valorizzazione delle risorse e della struttura a disposizione"
        else: return "caratterizzata da una minore efficienza nell'utilizzo strutturale delle proprie immobilizzazioni operative"

    def get_analisi_trend_profitto(az_prof, set_prof):
        if az_prof >= set_prof: return f"Questo dato, pari al {format_euro(az_prof)}%, dimostra un'elevata efficacia nella gestione dei costi accessori, degli oneri finanziari e del carico fiscale complessivo."
        else: return f"Questo dato ({format_euro(az_prof)}%) riflette una redditività netta complessivamente più limitata, risentendo del peso degli oneri extra-caratteristici."

    def get_confronto_profitto_settore(az_prof, set_prof):
        if az_prof >= set_prof: return f"La struttura dell'impresa supera la performance mediana dei concorrenti (pari a {format_euro(set_prof)}%), distinguendosi per una forte propensione alla generazione di utile netto di periodo."
        else: return f"La capacità di convertire i ricavi in risultato netto si colloca al di sotto della performance mediana espressa dal settore ({format_euro(set_prof)}%)."

    def get_prospettiva_redditivita_futura(az_prof, set_prof):
        if az_prof >= set_prof: return "dimostra di poter sostenere le dinamiche gestionali future grazie all'ottimizzazione del proprio ciclo dell'utile netto."
        else: return "deve focalizzarsi sul recupero dell'efficienza sui costi complessivi per migliorare la trasformazione del fatturato in utile netto."

    def get_posizionamento_margine_profitto_fine(az_prof, set_prof):
        if az_prof >= set_prof: return "su livelli superiori alla mediana di settore, confermando una buona capacità di trasformare i ricavi in risultato netto"
        else: return "su livelli compressi che limitano la redditività finale in confronto al comparto"

    def get_sintesi_bilancio_finale(az_prof, set_prof):
        if az_prof >= set_prof: return "una più elevata efficacia nella gestione dei costi lungo tutta la catena del valore aziendale"
        else: return "il forte peso delle gestioni accessorie, degli oneri del debito o del carico fiscale rispetto alla ricchezza economica operativa generata"

    def get_impatto_gestione_caratteristica(az_ebitda, set_ebitda):
        if az_ebitda >= set_ebitda: return "pienamente efficiente rispetto ai costi diretti correnti"
        else: return "debole, disperdendo gran parte del valore operativo in costi di produzione e spese d'esercizio"

    def get_impatto_oneri_accessori(az_ebit, az_prof):
        if (az_ebit - az_prof) > 4.0: return "assorbe le residue disponibilità operative, contraendo la marginalità netta a causa del peso della gestione extra-operativa"
        else: return "impatta in modo proporzionato, consentendo di registrare un utile finale in linea con i target attesi"

    def get_interpretazione_risultato_eco(az_ebitda, set_ebitda, az_prof, set_prof):
        if az_ebitda >= set_ebitda and az_prof >= set_prof: return "conferma in maniera analitica la piena solidità del modello e la gestione ottimale dei costi lungo tutta la catena del valore."
        elif az_ebitda < set_ebitda and az_prof < set_prof: return "indica uno squilibrio nell'area economica per la debole conversione dei ricavi lungo la filiera produttiva aziendale."
        else: return "restituisce una dinamica asimmetrica tra l'efficienza industriale caratteristica (EBITDA) e il rendimento netto complessivo."

    def get_stato_struttura_operativa(az_ebit, set_ebit):
        if az_ebit >= set_ebit: return "altamente performante nella massimizzazione dei ricavi e del reddito caratteristico"
        else: return "vulnerabile a causa di un'efficienza non ottimizzata sul rendimento degli asset caratteristici"

    def get_gestione_costi_ricavi_core(az_ebitda, set_ebitda):
        if az_ebitda >= set_ebitda: return "gestisce in modo efficiente i costi operativi esterni, generando un surplus soddisfacente rispetto ai ricavi di vendita"
        else: return "subisce un forte assorbimento dei ricavi a causa di una rigida incidenza dei costi di gestione operativa correnti"

    def get_implicazione_finale_eco(az_prof, set_prof):
        if az_prof >= set_prof: return "Mantenendo tale assetto organizzativo, l'impresa garantisce una solida autonomia nel lungo periodo grazie alla propria capacità di trasformazione dei ricavi in utile netto."
        else: return "La società necessita di interventi mirati per ottimizzare l'incidenza dei costi operativi e extra-caratteristici, al fine di recuperare redditività complessiva e autonomia economica."

    def get_sintesi_quadriennio_patr(az_str1, az_str2):
        if az_str1 >= 1.0 and az_str2 >= 1.0: return "un'elevata stabilità della struttura patrimoniale e l'assenza di rischi derivanti dal disequilibrio tra investimenti e capitale permanente"
        elif az_str1 < 1.0 and az_str2 < 1.0: return "uno squilibrio nella correlazione temporale tra fonti e impieghi, con potenziale immobilizzazione di capitale circolante a copertura di asset duraturi"
        else: return "un posizionamento asimmetrico e parzialmente sbilanciato nell'assetto di medio-lungo termine della struttura patrimoniale"

    def get_analisi_copertura_attivo_fisso(az_str1):
        if az_str1 >= 1.0: return "il capitale proprio finanzia integralmente gli investimenti a lungo periodo, non presentando vincoli di scadenza e riducendo l'esposizione verso terzi"
        else: return "si evidenzia una parziale copertura degli asset fissi, dipendendo in larga misura dal debito verso terzi, il che obbliga a un puntuale rimborso a breve"

    def get_confronto_struttura_settore(az_str1, set_str1):
        if az_str1 >= set_str1: return "L'assetto patrimoniale dell'azienda si distacca in modo favorevole rispetto alla mediana settoriale, limitando l'uso della leva finanziaria."
        else: return "L'assetto patrimoniale sconta un divario rispetto al posizionamento mediano dei competitor in termini di capitalizzazione autonoma."

    def get_analisi_soglia_struttura1(az_str1):
        if az_str1 >= 1.0: return f"L'Indice primario di struttura, attestandosi a {format_euro(az_str1)}, indica che il capitale proprio ha finanziato interamente le immobilizzazioni, garantendo la sicurezza degli investimenti."
        else: return f"L'Indice primario di struttura, attestandosi a {format_euro(az_str1)}, segnala che una parte delle immobilizzazioni è stata finanziata mediante capitale di terzi, erodendo i parametri di sicurezza."

    def get_confronto_mediana_struttura1(az_str1, set_str1):
        if az_str1 >= set_str1: return f"Il posizionamento dell'Indice risulta nettamente superiore alla mediana del comparto di riferimento (pari a {format_euro(set_str1)})."
        else: return f"L'azienda sconta un ritardo nella solidità del proprio Indice rispetto ai parametri mediani di settore (attestati a {format_euro(set_str1)})."

    def get_implicazione_copertura_attivo(az_str1):
        if az_str1 >= 1.0: return "gode di un assetto bilanciato e immune ai tempi di disinvestimento medio-lunghi delle immobilizzazioni."
        else: return "è soggetta a forte pressione temporale sui rimborsi dei capitali prestati da terzi a causa dello squilibrio d'impiego."

    def get_analisi_trend_struttura2(az_str2):
        if az_str2 >= 1.0: return f"L'Indice secondario di struttura ({format_euro(az_str2)}) segnala che il capitale permanente ha finanziato interamente le immobilizzazioni, assicurando correlazione temporale."
        else: return f"L'Indice secondario di struttura ({format_euro(az_str2)}) indica che una parte delle immobilizzazioni è stata finanziata attraverso capitale a breve termine, determinando uno squilibrio temporale."

    def get_confronto_settore_struttura2(az_str2, set_str2):
        if az_str2 >= set_str2: return "Tale dinamica risulta superiore alle performance di copertura strutturale dei competitor."
        else: return "Il trend sconta uno svantaggio in confronto alla correlazione temporale mediana del mercato."

    def get_conclusione_struttura2(az_str2):
        if az_str2 >= 1.0: return "riduce sensibilmente il rischio legato a rinegoziazioni del debito per la copertura delle immobilizzazioni fisse."
        else: return "comporta il rischio di tensioni finanziarie al momento del rinnovo degli affidamenti bancari a breve scadenza."

    def get_evoluzione_vantaggio_competitivo(az_str1, set_str1):
        if az_str1 >= set_str1: return "riesce a consolidare un vantaggio patrimoniale di rilievo all'interno del proprio segmento competitivo."
        else: return "necessita di colmare il divario patrimoniale rispetto ai principali operatori del comparto per preservare il rating."

    def get_resilienza_credit_crunch(az_str2):
        if az_str2 >= 1.0: return "è fortemente tutelato contro eventuali restrizioni creditizie grazie alla completa copertura fornita dal capitale permanente"
        else: return "risulta maggiormente esposto agli effetti delle restrizioni creditizie a causa dello squilibrio tra fonti correnti e immobilizzazioni"

    def get_flessibilita_strategica_asset(az_str1):
        if az_str1 >= 1.0: return "una buona flessibilità per nuovi piani di sviluppo strategico e investimenti durevoli."
        else: return "una rigidità nel finanziare nuovi progetti di rinnovamento degli asset e degli impianti fisici."

    def get_intro_analisi_gearing(az_gear, set_gear):
        if az_gear <= set_gear: return "L'analisi conferma un contenimento del Gearing coerente con la solidità della struttura patrimoniale."
        else: return "L'analisi evidenzia una marcata esposizione verso il capitale di terzi per il finanziamento aziendale."

    def get_andamento_storico_gearing(az_gear, set_gear):
        if az_gear <= set_gear: return f"Il dato rivela un Gearing oneroso pari al {format_euro(az_gear)}%. Valori contenuti indicano una limitata dipendenza dall'indebitamento bancario e una maggiore solidità complessiva."
        else: return f"Il dato palesa un Gearing elevato pari al {format_euro(az_gear)}%. Valori più alti segnalano un forte ricorso ai terzi, determinando un incremento del rischio finanziario."

    def get_confronto_gearing_settore(az_gear, set_gear):
        if az_gear <= set_gear: return f"l'esposizione debitoria risulta nettamente inferiore rispetto alla mediana del comparto ({format_euro(set_gear)}%), riducendo l'indice di rischio aziendale."
        else: return f"l'esposizione debitoria risulta del {format_euro(az_gear)}%, sensibilmente superiore rispetto al valore mediano di settore ({format_euro(set_gear)}%), incrementando il rischio d'insolvenza."

    def get_reazione_contesto_gearing(az_gear, set_gear):
        if az_gear <= set_gear: return "garantisce all'impresa una buona autonomia dai vincoli creditizi, riducendo l'indebitamento oneroso"
        else: return "richiede attenzione costante alla sostenibilità degli oneri, la cui leva agisce incrementando il rischio sui margini"

    def get_sintesi_valore_gearing(az_gear, set_gear):
        if az_gear <= set_gear: return "conferma la solidità aziendale e il limitato ricorso al capitale di terzi."
        else: return "rappresenta l'elemento di maggiore leva finanziaria che incide sull'autonomia economica complessiva."

    def get_analisi_rischio_default(az_gear, set_gear):
        if az_gear <= set_gear: return "una minore esposizione ai rischi di insolvenza legati ai debiti onerosi o alle pressioni sui rimborsi."
        else: return "una potenziale tensione nella capacità di rimborso in scenari di stress e restrizione del credito."

    def get_conclusione_autonomia_finanziaria(az_gear, set_gear):
        if az_gear <= set_gear: return "conferma la validità della propria politica di autofinanziamento, limitando la dipendenza esterna."
        else: return "necessita di un rafforzamento del patrimonio attraverso ricapitalizzazioni o utili portati a nuovo per mitigare il debito."

    def get_sintesi_finale_patr(az_str1, az_gear, set_gear):
        if az_str1 >= 1.0 and az_gear <= set_gear: return "conferma una solida copertura degli indici di struttura del passivo, fondata su capitale permanente."
        else: return "evidenzia la necessità di ricalibrare gli squilibri temporali e ridurre il ricorso al capitale di terzi per stabilizzare gli asset."

    def get_impatto_rischio_sistemico(az_gear, set_gear):
        if az_gear <= set_gear: return "riduce i rischi derivanti dall'irrigidimento del mercato creditizio o dall'aumento dei tassi"
        else: return "amplifica la vulnerabilità dell'azienda nei confronti delle variazioni dei tassi di interesse applicati dal sistema bancario"

    def get_posizionamento_competitivo_patr(az_str1, set_str1):
        if az_str1 >= set_str1: return "si colloca favorevolmente nelle classifiche settoriali in termini di stabilità strutturale a medio-lungo termine,"
        else: return "occupa i gradini più bassi della distribuzione per solidità e coordinamento temporale delle coperture,"

    def get_sintesi_modello_operativo_fin(az_cr, az_qr, az_rot, set_rot):
        solvibile = az_cr >= 1.0 and az_qr >= 1.0
        efficiente = az_rot >= set_rot
        if solvibile and efficiente: return "una capacità adeguata di convertire le proprie risorse in liquidità per onorare il debito corrente, unita a una rotazione commerciale superiore alla mediana di settore."
        elif solvibile and not efficiente: return "una solvibilità di breve termine adeguata, a fronte di una rotazione del capitale investito inferiore alla mediana di settore."
        elif not solvibile and efficiente: return "una rotazione del capitale investito superiore alla mediana di settore, a fronte di un certo livello di dipendenza dalla monetizzazione delle scorte per onorare gli impegni correnti."
        else: return "un modello d'esercizio che sconta parziali rallentamenti operativi e un certo livello di dipendenza dalla monetizzazione delle scorte correnti."

    def get_analisi_trend_current_ratio(az_cr):
        if az_cr >= 1.0: return f"L'azienda dispone sistematicamente di attività a breve termine sufficienti a coprire integralmente i debiti esigibili, registrando un coefficiente di {format_euro(az_cr)}."
        else: return f"Si segnala una contrazione delle attività liquidabili in rapporto alle passività esigibili (coefficiente a {format_euro(az_cr)}), segno di tensione corrente."

    def get_reazione_contesto_liquidita(az_cr):
        if az_cr >= 1.0: return "ha garantito costanti eccedenze di cassa (Current Ratio > 1) per prevenire tensioni finanziarie d'esercizio"
        else: return "ha manifestato l'incapacità delle attività correnti di far fronte alle passività a breve scadenza, limitando l'operatività"

    def get_confronto_current_ratio_mediana(az_cr, set_cr):
        if az_cr >= set_cr: return f"supera il parametro mediano del Current Ratio ({format_euro(set_cr)}), evidenziando una solidità a breve scadenza superiore rispetto ai competitor."
        else: return f"subisce un divario in negativo rispetto ai valori mediani del Current Ratio ({format_euro(set_cr)}), esponendo la liquidità a potenziali contrazioni monetarie."

    def get_rapporto_attivita_passivita_brevi(az_cr):
        if az_cr >= 1.0: return "le attività correnti coprono con margine adeguato gli oneri di breve termine, confermando l'equilibrio monetario"
        else: return "il livello dei debiti a breve eccede le disponibilità di pari durata, svalutando l'equilibrio della liquidità d'impresa"

    def get_capacita_generazione_liquidita(az_cr):
        if az_cr >= 1.0: return "conferma una situazione di pieno equilibrio finanziario di breve termine (Current Ratio >= 1)."
        else: return "configura l'incapacità temporanea delle attività correnti di estinguere le scadenze (Current Ratio < 1)."

    def get_analisi_quick_ratio_soglia(az_qr):
        if az_qr >= 1.0: return f"L'indice depurato dal magazzino si attesta a {format_euro(az_qr)}. Valori uguali o superiori all'unità indicano che la cassa copre gli impegni senza dover smobilizzare il magazzino."
        else: return f"L'indice depurato dal magazzino si contrae a {format_euro(az_qr)}. Valori inferiori ad uno evidenziano una dipendenza dalla vendita delle scorte per non fallire gli impegni di breve periodo."

    def get_implicazione_liquidita_immediata(az_qr):
        if az_qr >= 1.0: return "la flessibilità immediata garantisce la copertura integrale senza vendite forzate a sconto (Quick Ratio)."
        else: return "l'azienda dipenderà fisiologicamente dalla monetizzazione delle scorte per soddisfare i pagamenti a brevissimo termine (Quick Ratio)."

    def get_andamento_liquidita_primaria(az_qr):
        if az_qr >= 1.0: return "una gestione ottimizzata della liquidità primaria, potendo contare su solide riserve bancarie o crediti certi."
        else: return "una chiara carenza monetaria immediata che espone l'equilibrio d'esercizio alla rotazione dei magazzini fisici."

    def get_confronto_quick_ratio_settore(az_qr, set_qr):
        if az_qr >= set_qr: return f"superiore ai livelli mediani del Quick Ratio ({format_euro(set_qr)}), proteggendo le casse dall'obbligo di smobilizzo coatto delle merci."
        elif az_qr >= 1.0: return f"al di sotto della mediana di mercato nel Quick Ratio ({format_euro(set_qr)}), un divario da monitorare rispetto ai competitor pur restando la liquidità immediata su livelli di per sé adeguati."
        else: return f"al di sotto della mediana di mercato nel Quick Ratio ({format_euro(set_qr)}), forzando l'impresa verso un realizzo più intensivo delle scorte per onorare gli impegni a breve."

    def get_copertura_debiti_breve_quick(az_qr):
        if az_qr >= 1.0: return "risultano ampiamente sufficienti per coprire gli esborsi immediati (Quick Ratio >= 1)."
        else: return "risultano parziali, confermando lo squilibrio e la debolezza del Quick Ratio inferiore all'unità."

    def get_interpretazione_flussi_cassa_quick(az_qr):
        if az_qr >= 1.0: return "testimonia una buona velocità di conversione del circolante netto, senza gravare sulle rimanenze commerciali."
        else: return "dimostra inefficienze del ciclo d'incasso, con la chiara dipendenza dalle scorte che affossa il Quick Ratio."

    def get_confronto_rotazione_mediana(az_rot, set_rot):
        if az_rot >= set_rot: return f"L'efficienza d'utilizzo si colloca su livelli favorevoli, risultando superiore al parametro mediano ({format_euro(set_rot)}) e confermando un'ottimale resa."
        else: return f"L'Indice di rotazione si colloca al di sotto della mediana settoriale ({format_euro(set_rot)}), evidenziando un impiego poco performante del capitale."

    def get_interpretazione_modello_rotazione(az_rot, set_rot):
        if az_rot >= set_rot: return "si distingue per valori elevati dell'Indice di Rotazione, a dimostrazione di una spiccata capacità dell'impresa di generare ricavi attraverso le risorse investite e di una maggiore efficienza operativa."
        else: return "si caratterizza per valori contenuti dell'Indice di Rotazione, segnalando una ridotta capacità del capitale di tradursi in ricavi e una minore efficienza operativa nell'allocazione degli asset."

    def get_analisi_anomalia_rotazione(az_rot, set_rot):
        if az_rot >= set_rot: return "l'efficienza commerciale e d'incasso si riflette in modo positivo sulla produttività dell'Indice di Rotazione."
        else: return "la carenza di fatturato rispetto al capitale immesso penalizza l'andamento della Rotazione d'esercizio."

    def get_rapporto_capitale_fatturato(az_rot, set_rot):
        if az_rot >= set_rot: return "risulta produttivo in termini di traduzione in ricavi lordi"
        else: return "risulta scarsamente produttivo, a causa di un impiego rigido delle risorse immobilizzate che deprime il rapporto con i ricavi lordi"

    def get_motivazione_strutturale_rotazione(az_rot, set_rot):
        if az_rot >= set_rot: return "un modello di business capace di sfruttare efficacemente gli investimenti storici operati."
        else: return "un eccesso di asset non pienamente operativi che ingessa il rapporto tra capitale investito e vendite realizzate."

    def get_sintesi_finale_fin(az_cr, az_qr, az_rot, set_rot):
        solvibile = az_cr >= 1.0 and az_qr >= 1.0
        efficiente = az_rot >= set_rot
        if solvibile and efficiente: return "un giudizio di equilibrio sulle metriche di solvibilità di breve termine, unito a una rotazione dell'attivo superiore alla mediana di settore."
        elif solvibile and not efficiente: return "un giudizio di equilibrio sulla solvibilità di breve termine, a fronte di una rotazione dell'attivo da migliorare rispetto alla mediana di settore."
        elif not solvibile and efficiente: return "una rotazione dell'attivo superiore alla mediana di settore, a fronte di margini di miglioramento sulla solvibilità di breve termine."
        else: return "debolezze su più fronti: l'indice segnala tensioni di cassa correnti unite a un'inefficienza nella rotazione operativa del capitale."

    def get_motivazione_rating_fin(az_cr, az_qr):
        if az_cr >= 1.0 and az_qr >= 1.0: return "l'azienda dispone di risorse sufficienti a coprire i debiti (Current Ratio >= 1) senza svendere rimanenze (Quick Ratio >= 1)."
        elif az_cr >= 1.0: return "l'azienda copre i debiti a breve con le attività correnti nel loro complesso (Current Ratio >= 1), ma tale copertura dipende in parte dallo smobilizzo delle rimanenze (Quick Ratio < 1)."
        else: return "l'azienda risente di un'incapacità parziale di estinguere le passività a breve senza attingere al magazzino."

    def get_gestione_tesoreria_fin(az_cr, az_qr):
        if az_cr >= 1.0 and az_qr >= 1.0: return "assicura un equilibrio di breve termine solido e privo di esposizione allo smobilizzo merci."
        elif az_cr >= 1.0: return "assicura la copertura complessiva dei debiti a breve, pur con una quota di dipendenza dallo smobilizzo delle rimanenze."
        else: return "evidenzia la necessità stringente di monitorare la cassa immediata per far fronte agli impegni non dilazionabili."

    def get_priorita_strategica_fin(az_rot, set_rot):
        if az_rot >= set_rot: return "È possibile focalizzarsi su piani di sviluppo commerciale per sfruttare la redditività di rotazione già superiore alla mediana di settore."
        else: return "È opportuno valutare la dismissione di asset non core o un incremento delle vendite per tradurre in flussi di cassa il capitale immesso."

    def get_sintesi_profilo_integrato(rating_comb):
        if rating_comb == 'A': return "un modello solido, patrimonialmente indipendente e con buone coperture di liquidità a breve."
        elif rating_comb == 'B': return "un assetto complessivamente equilibrato che assicura la continuità in linea col mercato di riferimento."
        else: return "uno scenario di vulnerabilità su più fronti (marginalità o cassa) disallineato rispetto alla concorrenza."

    def get_sintesi_posizionamento_lungo_periodo(rating_comb):
        if rating_comb == 'A': return "in grado di assorbire con maggiore margine le perturbazioni di mercato, grazie a solidi indici di solvibilità e redditività operativa."
        elif rating_comb == 'B': return "resiliente e atta a garantire la sopravvivenza d'esercizio mantenendo un attento bilanciamento dei costi."
        else: return "esposta a un maggiore rischio di tensione di liquidità e di erosione dei margini, in presenza di una leva finanziaria elevata."

    def get_conclusione_patrimoniale(az_str1, az_gear, set_gear):
        if az_str1 >= 1.0 and az_gear <= set_gear: return "La copertura tramite mezzi propri e il Gearing contenuto riducono l'esposizione dell'azienda a eventuali restrizioni del credito bancario."
        else: return "È prioritario avviare azioni volte a riequilibrare il capitale, riducendo l'esposizione al debito per mitigare il rischio tassi."

    def get_conclusione_economica(az_ebitda, set_ebitda, az_prof, set_prof):
        if az_ebitda >= set_ebitda and az_prof >= set_prof: return "Confermano la capacità di generare reddito lordo dalla gestione caratteristica e di ottimizzare gli oneri fino all'utile netto finale."
        else: return "Indicano l'opportunità di rivedere l'incidenza dei costi operativi e di ridurre il peso del carico fiscale o degli oneri finanziari."

    def get_conclusione_finanziaria_dettaglio(az_cr, az_qr):
        if az_cr >= 1.0 and az_qr >= 1.0: return "Conferma coperture di cassa senza tensioni evidenti e non subordinate allo svuotamento dei magazzini."
        elif az_cr >= 1.0: return "Conferma la copertura dei debiti a breve nel complesso, pur con una componente di dipendenza dallo smobilizzo dei magazzini."
        else: return "Rivela fragilità strutturali nel circolante netto, a conferma dell'opportunità di incassare o dismettere scorte in tempi più contratti."

    def get_raccomandazione_finale(rating_comb):
        if rating_comb == 'A': return "dispone di un buon margine per accedere a nuovo credito per investimenti o espansione settoriale."
        elif rating_comb == 'B': return "dovrebbe concentrare l'azione manageriale sull'ottimizzazione del circolante per compiere uno scale-up stabile."
        else: return "dovrebbe strutturare un piano di rafforzamento patrimoniale e operativo, valutando iniezioni di capitale e un contenimento mirato dei costi non strategici."


    context['descr_rating_tot'] = get_testo_totale(context['rating_comb'])
    context['descr_rating_eco'] = get_testo_eco(context['rating_eco'])
    context['descr_rating_patr'] = get_testo_patr(context['rating_patr'])
    context['descr_rating_fin'] = get_testo_fin(context['rating_fin'])
    context['descr_sintesi'] = get_testo_sintesi(context['rating_comb'])
    
    # =================================================================
    # 🟢 ESTRAZIONE VALORI NUMERICI REALI (FLOAT) E MEDIANE DI SETTORE (2024)
    # =================================================================
    if not df_target.empty:
        riga_target = df_target.iloc[0]
        val_az_ebitda_24 = pd.to_numeric(riga_target.get('Margine EBITDA (*) % 2024'), errors='coerce')
        val_az_ebit_24 = pd.to_numeric(riga_target.get('Margine EBIT (*) % 2024'), errors='coerce')
        val_az_profitto_24 = pd.to_numeric(riga_target.get('Margine di Profitto (*) % 2024'), errors='coerce')
        val_az_strut1_24 = pd.to_numeric(riga_target.get('Indice di Struttura 1° livello (*) 2024'), errors='coerce')
        val_az_strut2_24 = pd.to_numeric(riga_target.get('Indice di Struttura 2° livello (*) 2024'), errors='coerce')
        val_az_gearing_24 = pd.to_numeric(riga_target.get('Gearing (*) % 2024'), errors='coerce')
        val_az_cr_24 = pd.to_numeric(riga_target.get('Current Ratio (*) 2024'), errors='coerce')
        val_az_qr_24 = pd.to_numeric(riga_target.get('Quick Ratio (*) 2024'), errors='coerce')
        val_az_rot_24 = pd.to_numeric(riga_target.get('Indice di Rotazione del Capitale Investito (*) 2024'), errors='coerce')
    else:
        val_az_ebitda_24 = val_az_ebit_24 = val_az_profitto_24 = 0
        val_az_strut1_24 = val_az_strut2_24 = val_az_gearing_24 = 0
        val_az_cr_24 = val_az_qr_24 = val_az_rot_24 = 0

    # Mediane del settore (2024)
    val_set_ebitda_24 = df_orbis['Margine EBITDA (*) % 2024'].median() if 'Margine EBITDA (*) % 2024' in df_orbis.columns else 0
    val_set_ebit_24 = df_orbis['Margine EBIT (*) % 2024'].median() if 'Margine EBIT (*) % 2024' in df_orbis.columns else 0
    val_set_profitto_24 = df_orbis['Margine di Profitto (*) % 2024'].median() if 'Margine di Profitto (*) % 2024' in df_orbis.columns else 0
    val_set_strut1_24 = df_orbis['Indice di Struttura 1° livello (*) 2024'].median() if 'Indice di Struttura 1° livello (*) 2024' in df_orbis.columns else 0
    val_set_strut2_24 = df_orbis['Indice di Struttura 2° livello (*) 2024'].median() if 'Indice di Struttura 2° livello (*) 2024' in df_orbis.columns else 0
    val_set_gearing_24 = df_orbis['Gearing (*) % 2024'].median() if 'Gearing (*) % 2024' in df_orbis.columns else 0
    val_set_cr_24 = df_orbis['Current Ratio (*) 2024'].median() if 'Current Ratio (*) 2024' in df_orbis.columns else 0
    val_set_qr_24 = df_orbis['Quick Ratio (*) 2024'].median() if 'Quick Ratio (*) 2024' in df_orbis.columns else 0
    val_set_rot_24 = df_orbis['Indice di Rotazione del Capitale Investito (*) 2024'].median() if 'Indice di Rotazione del Capitale Investito (*) 2024' in df_orbis.columns else 0

    # =================================================================
    # 🎯 POPOLAMENTO REALE DEL DIZIONARIO CON CHIAMATE POSIZIONALI CORRETTE
    # =================================================================
    context['intro_benchmark_eco'] = get_intro_benchmark_eco(context['rating_eco'])
    context['analisi_margini_operativi'] = get_analisi_margini_operativi(val_az_ebitda_24, val_set_ebitda_24, val_az_ebit_24, val_set_ebit_24, desc_nace_pulita)
    context['analisi_margine_profitto'] = get_analisi_margine_profitto(val_az_profitto_24, val_set_profitto_24, desc_nace_pulita)

    context['intro_benchmark_patr'] = get_intro_benchmark_patr(context['rating_patr'])
    context['analisi_indici_struttura'] = get_analisi_indici_struttura(val_az_strut1_24, val_set_strut1_24, val_az_strut2_24, val_set_strut2_24)
    context['analisi_gearing'] = get_analisi_gearing(val_az_gearing_24, val_set_gearing_24)

    context['intro_benchmark_fin'] = get_intro_benchmark_fin(context['rating_fin'])
    context['analisi_rotazione'] = get_analisi_rotazione(val_az_rot_24, val_set_rot_24, desc_nace_pulita)
    context['analisi_current_ratio'] = get_analisi_current_ratio(val_az_cr_24, val_set_cr_24)
    context['analisi_quick_ratio'] = get_analisi_quick_ratio(val_az_qr_24, val_set_qr_24)
    context['analisi_posizionamento_fin'] = get_analisi_posizionamento_fin(val_az_cr_24, val_az_qr_24, val_set_cr_24, val_set_qr_24)

    # -----------------------------------------------------
    # INIEZIONE TAG COMPLETAMENTE SEPARATI PER WORD
    # -----------------------------------------------------
    context['intro_margini'] = get_intro_benchmark_eco(context['rating_eco']) + " " + get_intro_margini(desc_nace_pulita)
    context['analisi_ebitda'] = get_analisi_ebitda(val_az_ebitda_24, val_set_ebitda_24)
    context['analisi_ebit'] = get_analisi_ebit(val_az_ebit_24, val_set_ebit_24)
    context['analisi_margine_profitto'] = get_analisi_margine_profitto_tag(val_az_profitto_24, val_set_profitto_24)

    context['analisi_struttura1'] = get_analisi_struttura1(val_az_strut1_24)
    context['analisi_struttura2'] = get_analisi_struttura2(val_az_strut2_24)
    context['analisi_gearing'] = get_analisi_gearing_tag(val_az_gearing_24, val_set_gearing_24)

    context['analisi_current_ratio'] = get_analisi_current_ratio_tag(val_az_cr_24, val_set_cr_24)
    context['analisi_quick_ratio'] = get_analisi_quick_ratio_tag(val_az_qr_24, val_set_qr_24)
    context['analisi_rotazione'] = get_analisi_rotazione_tag(val_az_rot_24, val_set_rot_24)


    # Calcolo dei totali per la fascia di Rating dell'Azienda
    target_glob = context['rating_comb']
    
    if target_glob in ['A', 'B', 'C'] and 'Benchmark Totale' in df_rating.columns:
        num_soc_fascia_tot = len(df_rating[df_rating['Benchmark Totale'] == target_glob])
        tot_valide_bench = len(df_rating[df_rating['Benchmark Totale'].isin(['A', 'B', 'C'])])
        perc_soc_fascia_tot = (num_soc_fascia_tot / tot_valide_bench * 100) if tot_valide_bench > 0 else 0
        
        num_eco_fascia = len(df_rating[df_rating['Rating Economico'] == target_glob])
        num_patr_fascia = len(df_rating[df_rating['Rating Patrimoniale'] == target_glob])
        num_fin_fascia = len(df_rating[df_rating['Rating Finanziario'] == target_glob])
    else:
        num_soc_fascia_tot, perc_soc_fascia_tot, num_eco_fascia, num_patr_fascia, num_fin_fascia = 0, 0, 0, 0, 0

    context['analisi_combinata'] = get_analisi_combinata(context['rating_eco'], context['rating_patr'], context['rating_fin'])
    context['descr_fascia_appartenenza'] = get_descr_fascia_appartenenza(context['rating_comb'])
    context['num_soc_fascia_tot'] = f"{num_soc_fascia_tot:,}".replace(',', '.')
    context['perc_soc_fascia_tot'] = format_euro(perc_soc_fascia_tot)
    context['num_eco_fascia'] = f"{num_eco_fascia:,}".replace(',', '.')
    context['num_patr_fascia'] = f"{num_patr_fascia:,}".replace(',', '.')
    context['num_fin_fascia'] = f"{num_fin_fascia:,}".replace(',', '.')

    context['intro_divario_strutturale'] = get_intro_divario_strutturale(val_az_ebitda_24, val_set_ebitda_24)
    context['analisi_trend_ebitda'] = get_analisi_trend_ebitda(val_az_ebitda_24, val_set_ebitda_24)
    context['asimmetria_ebitda'] = get_asimmetria_ebitda(val_az_ebitda_24, val_set_ebitda_24)
    context['confronto_costi_settore'] = get_confronto_costi_settore(val_az_ebitda_24, val_set_ebitda_24)
    context['impatto_costi_su_margine'] = get_impatto_costi_su_margine(val_az_ebitda_24, val_set_ebitda_24)

    context['analisi_trend_ebit'] = get_analisi_trend_ebit(val_az_ebit_24, val_set_ebit_24)
    context['confronto_ebit_settore'] = get_confronto_ebit_settore(val_az_ebit_24, val_set_ebit_24)
    context['sintesi_ebit_mediana'] = get_sintesi_ebit_mediana(val_az_ebit_24, val_set_ebit_24)
    context['analisi_verticale_ebit'] = get_analisi_verticale_ebit(val_az_ebitda_24, val_az_ebit_24, val_set_ebitda_24, val_set_ebit_24)
    context['efficienza_struttura_asset'] = get_efficienza_struttura_asset(val_az_ebit_24, val_set_ebit_24)

    context['analisi_trend_profitto'] = get_analisi_trend_profitto(val_az_profitto_24, val_set_profitto_24)
    context['confronto_profitto_settore'] = get_confronto_profitto_settore(val_az_profitto_24, val_set_profitto_24)
    context['prospettiva_redditivita_futura'] = get_prospettiva_redditivita_futura(val_az_profitto_24, val_set_profitto_24)
    context['posizionamento_margine_profitto_fine'] = get_posizionamento_margine_profitto_fine(val_az_profitto_24, val_set_profitto_24)
    context['sintesi_bilancio_finale'] = get_sintesi_bilancio_finale(val_az_profitto_24, val_set_profitto_24)

    context['impatto_gestione_caratteristica'] = get_impatto_gestione_caratteristica(val_az_ebitda_24, val_set_ebitda_24)
    context['impatto_oneri_accessori'] = get_impatto_oneri_accessori(val_az_ebit_24, val_az_profitto_24)
    context['interpretazione_risultato_eco'] = get_interpretazione_risultato_eco(val_az_ebitda_24, val_set_ebitda_24, val_az_profitto_24, val_set_profitto_24)
    context['stato_struttura_operativa'] = get_stato_struttura_operativa(val_az_ebit_24, val_set_ebit_24)
    context['gestione_costi_ricavi_core'] = get_gestione_costi_ricavi_core(val_az_ebitda_24, val_set_ebitda_24)
    context['implicazione_finale_eco'] = get_implicazione_finale_eco(val_az_profitto_24, val_set_profitto_24)

    context['sintesi_quadriennio_patr'] = get_sintesi_quadriennio_patr(val_az_strut1_24, val_az_strut2_24)
    context['analisi_copertura_attivo_fisso'] = get_analisi_copertura_attivo_fisso(val_az_strut1_24)
    context['confronto_struttura_settore'] = get_confronto_struttura_settore(val_az_strut1_24, val_set_strut1_24)
    context['analisi_soglia_struttura1'] = get_analisi_soglia_struttura1(val_az_strut1_24)
    context['confronto_mediana_struttura1'] = get_confronto_mediana_struttura1(val_az_strut1_24, val_set_strut1_24)
    context['implicazione_copertura_attivo'] = get_implicazione_copertura_attivo(val_az_strut1_24)

    context['analisi_trend_struttura2'] = get_analisi_trend_struttura2(val_az_strut2_24)
    context['confronto_settore_struttura2'] = get_confronto_settore_struttura2(val_az_strut2_24, val_set_strut2_24)
    context['conclusione_struttura2'] = get_conclusione_struttura2(val_az_strut2_24)
    context['evoluzione_vantaggio_competitivo'] = get_evoluzione_vantaggio_competitivo(val_az_strut1_24, val_set_strut1_24)
    context['resilienza_credit_crunch'] = get_resilienza_credit_crunch(val_az_strut2_24)
    context['flessibilita_strategica_asset'] = get_flessibilita_strategica_asset(val_az_strut1_24)

    context['intro_analisi_gearing'] = get_intro_analisi_gearing(val_az_gearing_24, val_set_gearing_24)
    context['andamento_storico_gearing'] = get_andamento_storico_gearing(val_az_gearing_24, val_set_gearing_24)
    context['confronto_gearing_settore'] = get_confronto_gearing_settore(val_az_gearing_24, val_set_gearing_24)
    context['reazione_contesto_gearing'] = get_reazione_contesto_gearing(val_az_gearing_24, val_set_gearing_24)
    context['sintesi_valore_gearing'] = get_sintesi_valore_gearing(val_az_gearing_24, val_set_gearing_24)
    context['analisi_rischio_default'] = get_analisi_rischio_default(val_az_gearing_24, val_set_gearing_24)
    context['conclusione_autonomia_finanziaria'] = get_conclusione_autonomia_finanziaria(val_az_gearing_24, val_set_gearing_24)
    context['sintesi_finale_patr'] = get_sintesi_finale_patr(val_az_strut1_24, val_az_gearing_24, val_set_gearing_24)
    context['impatto_rischio_sistemico'] = get_impatto_rischio_sistemico(val_az_gearing_24, val_set_gearing_24)
    context['posizionamento_competitivo_patr'] = get_posizionamento_competitivo_patr(val_az_strut1_24, val_set_strut1_24)

    context['sintesi_modello_operativo_fin'] = get_sintesi_modello_operativo_fin(val_az_cr_24, val_az_qr_24, val_az_rot_24, val_set_rot_24)
    context['analisi_trend_current_ratio'] = get_analisi_trend_current_ratio(val_az_cr_24)
    context['reazione_contesto_liquidita'] = get_reazione_contesto_liquidita(val_az_cr_24)
    context['confronto_current_ratio_mediana'] = get_confronto_current_ratio_mediana(val_az_cr_24, val_set_cr_24)
    context['rapporto_attivita_passivita_brevi'] = get_rapporto_attivita_passivita_brevi(val_az_cr_24)
    context['capacita_generazione_liquidita'] = get_capacita_generazione_liquidita(val_az_cr_24)

    context['analisi_quick_ratio_soglia'] = get_analisi_quick_ratio_soglia(val_az_qr_24)
    context['implicazione_liquidita_immediata'] = get_implicazione_liquidita_immediata(val_az_qr_24)
    context['andamento_liquidita_primaria'] = get_andamento_liquidita_primaria(val_az_qr_24)
    context['confronto_quick_ratio_settore'] = get_confronto_quick_ratio_settore(val_az_qr_24, val_set_qr_24)
    context['copertura_debiti_breve_quick'] = get_copertura_debiti_breve_quick(val_az_qr_24)
    context['interpretazione_flussi_cassa_quick'] = get_interpretazione_flussi_cassa_quick(val_az_qr_24)

    context['confronto_rotazione_mediana'] = get_confronto_rotazione_mediana(val_az_rot_24, val_set_rot_24)
    context['interpretazione_modello_rotazione'] = get_interpretazione_modello_rotazione(val_az_rot_24, val_set_rot_24)
    context['analisi_anomalia_rotazione'] = get_analisi_anomalia_rotazione(val_az_rot_24, val_set_rot_24)
    context['rapporto_capitale_fatturato'] = get_rapporto_capitale_fatturato(val_az_rot_24, val_set_rot_24)
    context['motivazione_strutturale_rotazione'] = get_motivazione_strutturale_rotazione(val_az_rot_24, val_set_rot_24)
    context['sintesi_finale_fin'] = get_sintesi_finale_fin(val_az_cr_24, val_az_qr_24, val_az_rot_24, val_set_rot_24)
    context['motivazione_rating_fin'] = get_motivazione_rating_fin(val_az_cr_24, val_az_qr_24)
    context['gestione_tesoreria_fin'] = get_gestione_tesoreria_fin(val_az_cr_24, val_az_qr_24)
    context['priorita_strategica_fin'] = get_priorita_strategica_fin(val_az_rot_24, val_set_rot_24)

    context['sintesi_profilo_integrato'] = get_sintesi_profilo_integrato(context['rating_comb'])
    context['sintesi_posizionamento_lungo_periodo'] = get_sintesi_posizionamento_lungo_periodo(context['rating_comb'])
    context['conclusione_patrimoniale'] = get_conclusione_patrimoniale(val_az_strut1_24, val_az_gearing_24, val_set_gearing_24)
    context['conclusione_economica'] = get_conclusione_economica(val_az_ebitda_24, val_set_ebitda_24, val_az_profitto_24, val_set_profitto_24)
    context['conclusione_finanziaria_dettaglio'] = get_conclusione_finanziaria_dettaglio(val_az_cr_24, val_az_qr_24)
    context['raccomandazione_finale'] = get_raccomandazione_finale(context['rating_comb'])
    context['impatto_territoriale'] = get_impatto_territoriale(perc_ricavi_target_su_macro, ragione_sociale_pulita, format_euro(ricavi_mln))

    # Costruzione delle Medaglie di Settore!!!
    conteggi = df_rating['Benchmark Totale'].value_counts()
    
    # =================================================================
    # 📚 NOTA METODOLOGICA E DIZIONARIO ISTAT
    # =================================================================
    
    # 1. Società valide (post-filtri) e Iniziali (dal foglio Risultati passato da finhack)
    totale_valide = len(df_rating)
    
    context['num_soc_valide'] = f"{len(df_rating):,}".replace(',', '.')
    context['num_max_soc'] = str(num_max_soc_orbis) # <--- Prende il 25.232 da Finhack!

    # 2. Dizionario ISTAT interno (tutti i settori)
    istat_db = {
        "00.10": 4751988,
        "B": 1881,
        "06": 16,
        "06.1": 6,
        "06.10": 6,
        "06.2": 10,
        "06.20": 10,
        "07": 2,
        "07.1": 1,
        "07.10": 1,
        "07.2": 1,
        "07.29": 1,
        "08": 1789,
        "08.1": 1627,
        "08.11": 768,
        "08.12": 859,
        "08.9": 162,
        "08.91": 4,
        "08.92": 3,
        "08.93": 22,
        "08.99": 133,
        "09": 74,
        "09.1": 44,
        "09.10": 44,
        "09.9": 30,
        "09.90": 30,
        "C": 355908,
        "10": 48051,
        "10.1": 3270,
        "10.11": 1353,
        "10.12": 108,
        "10.13": 1809,
        "10.2": 428,
        "10.20": 428,
        "10.3": 1714,
        "10.31": 44,
        "10.32": 127,
        "10.39": 1543,
        "10.4": 2780,
        "10.41": 2773,
        "10.42": 7,
        "10.5": 3084,
        "10.51": 2727,
        "10.52": 357,
        "10.6": 902,
        "10.61": 894,
        "10.62": 8,
        "10.7": 30951,
        "10.71": 25954,
        "10.72": 1346,
        "10.73": 3651,
        "10.8": 4411,
        "10.81": 11,
        "10.82": 564,
        "10.83": 897,
        "10.84": 294,
        "10.85": 987,
        "10.86": 253,
        "10.89": 1405,
        "10.9": 511,
        "10.91": 374,
        "10.92": 137,
        "11": 3203,
        "11.0": 3203,
        "11.01": 659,
        "11.02": 1582,
        "11.03": 5,
        "11.04": 66,
        "11.05": 684,
        "11.06": 3,
        "11.07": 204,
        "12": 10,
        "12.0": 10,
        "12.00": 10,
        "13": 10660,
        "13.1": 1153,
        "13.10": 1153,
        "13.2": 1328,
        "13.20": 1328,
        "13.3": 1787,
        "13.30": 1787,
        "13.9": 6392,
        "13.91": 682,
        "13.92": 3035,
        "13.93": 123,
        "13.94": 139,
        "13.95": 249,
        "13.96": 1073,
        "13.99": 1091,
        "14": 27641,
        "14.1": 25001,
        "14.11": 592,
        "14.12": 426,
        "14.13": 16853,
        "14.14": 880,
        "14.19": 6250,
        "14.2": 463,
        "14.20": 463,
        "14.3": 2177,
        "14.31": 434,
        "14.39": 1743,
        "15": 12408,
        "15.1": 6171,
        "15.11": 1575,
        "15.12": 4596,
        "15.2": 6237,
        "15.20": 6237,
        "16": 19873,
        "16.1": 2088,
        "16.10": 2088,
        "16.2": 17785,
        "16.21": 339,
        "16.22": 126,
        "16.23": 11838,
        "16.24": 1276,
        "16.29": 4206,
        "17": 3248,
        "17.1": 224,
        "17.11": 3,
        "17.12": 221,
        "17.2": 3024,
        "17.21": 1246,
        "17.22": 188,
        "17.23": 998,
        "17.24": 14,
        "17.29": 578,
        "18": 12375,
        "18.1": 12251,
        "18.11": 25,
        "18.12": 8968,
        "18.13": 2265,
        "18.14": 993,
        "18.2": 124,
        "18.20": 124,
        "19": 259,
        "19.1": 2,
        "19.10": 2,
        "19.2": 257,
        "19.20": 257,
        "20": 4242,
        "20.1": 925,
        "20.11": 81,
        "20.12": 81,
        "20.13": 133,
        "20.14": 90,
        "20.15": 184,
        "20.16": 344,
        "20.17": 12,
        "20.2": 33,
        "20.20": 33,
        "20.3": 721,
        "20.30": 721,
        "20.4": 1537,
        "20.41": 434,
        "20.42": 1103,
        "20.5": 1007,
        "20.51": 106,
        "20.52": 81,
        "20.53": 99,
        "20.59": 721,
        "20.6": 19,
        "20.60": 19,
        "21": 489,
        "21.1": 93,
        "21.10": 93,
        "21.2": 396,
        "21.20": 396,
        "22": 9109,
        "22.1": 1348,
        "22.19": 1261,
        "22.2": 7761,
        "22.21": 795,
        "22.22": 1452,
        "22.23": 745,
        "22.29": 4769,
        "23": 16150,
        "23.1": 2922,
        "23.11": 8,
        "23.12": 2142,
        "23.13": 53,
        "23.14": 32,
        "23.19": 687,
        "23.2": 87,
        "23.20": 87,
        "23.3": 494,
        "23.31": 260,
        "23.32": 234,
        "23.4": 2144,
        "23.41": 1854,
        "23.42": 61,
        "23.43": 6,
        "23.44": 36,
        "23.49": 187,
        "23.5": 107,
        "23.51": 46,
        "23.52": 61,
        "23.6": 2541,
        "23.61": 960,
        "23.62": 73,
        "23.63": 871,
        "23.64": 63,
        "23.65": 30,
        "23.69": 544,
        "23.7": 7055,
        "23.70": 7055,
        "23.9": 800,
        "23.91": 207,
        "23.99": 593,
        "24": 2948,
        "24.1": 370,
        "24.10": 370,
        "24.2": 350,
        "24.20": 350,
        "24.3": 817,
        "24.31": 14,
        "24.32": 37,
        "24.33": 622,
        "24.34": 144,
        "24.4": 572,
        "24.41": 163,
        "24.42": 249,
        "24.43": 28,
        "24.44": 57,
        "24.45": 75,
        "24.5": 839,
        "24.51": 117,
        "24.52": 31,
        "24.53": 437,
        "24.54": 254,
        "25": 71302,
        "25.1": 28486,
        "25.11": 12415,
        "25.12": 16071,
        "25.2": 535,
        "25.21": 97,
        "25.29": 438,
        "25.3": 91,
        "25.30": 91,
        "25.4": 114,
        "25.40": 114,
        "25.5": 1478,
        "25.50": 1478,
        "25.6": 23176,
        "25.61": 4506,
        "25.62": 18670,
        "25.7": 3877,
        "25.71": 303,
        "25.72": 523,
        "25.73": 3051,
        "25.9": 13545,
        "25.91": 74,
        "25.92": 149,
        "25.93": 585,
        "25.94": 321,
        "25.99": 12416,
        "26": 4818,
        "26.1": 1799,
        "26.11": 880,
        "26.12": 919,
        "26.2": 534,
        "26.20": 534,
        "26.3": 577,
        "26.30": 577,
        "26.4": 232,
        "26.40": 232,
        "26.5": 905,
        "26.51": 832,
        "26.52": 73,
        "26.6": 584,
        "26.60": 584,
        "26.7": 186,
        "26.70": 186,
        "26.8": 1,
        "26.80": 1,
        "27": 7108,
        "27.1": 1944,
        "27.11": 1010,
        "27.12": 934,
        "27.2": 79,
        "27.20": 79,
        "27.3": 780,
        "27.31": 14,
        "27.32": 402,
        "27.33": 364,
        "27.4": 1184,
        "27.40": 1184,
        "27.5": 473,
        "27.51": 329,
        "27.52": 144,
        "27.9": 2648,
        "27.90": 2648,
        "28": 18161,
        "28.1": 2232,
        "28.11": 172,
        "28.12": 406,
        "28.13": 400,
        "28.14": 792,
        "28.15": 462,
        "28.2": 7130,
        "28.21": 401,
        "28.22": 1107,
        "28.23": 128,
        "28.24": 33,
        "28.25": 992,
        "28.29": 4469,
        "28.3": 1322,
        "28.30": 1322,
        "28.4": 1664,
        "28.41": 627,
        "28.49": 1037,
        "28.9": 5813,
        "28.91": 567,
        "28.92": 420,
        "28.94": 805,
        "28.95": 285,
        "28.96": 479,
        "28.99": 1626,
        "29": 2451,
        "29.1": 143,
        "29.10": 143,
        "29.2": 826,
        "29.20": 826,
        "29.3": 1482,
        "29.31": 247,
        "29.32": 1235,
        "30": 3070,
        "30.1": 1953,
        "30.11": 616,
        "30.12": 1337,
        "30.2": 95,
        "30.20": 95,
        "30.3": 198,
        "30.30": 198,
        "30.4": 1,
        "30.40": 1,
        "30.9": 823,
        "30.91": 374,
        "30.92": 435,
        "30.99": 14,
        "31": 16439,
        "31.0": 16439,
        "31.01": 2102,
        "31.02": 619,
        "31.03": 517,
        "31.09": 13201,
        "32": 26718,
        "32.1": 7144,
        "32.11": 19,
        "32.12": 5534,
        "32.13": 1591,
        "32.2": 821,
        "32.20": 821,
        "32.3": 462,
        "32.30": 462,
        "32.4": 333,
        "32.40": 333,
        "32.5": 14780,
        "32.50": 14780,
        "32.9": 3178,
        "32.91": 179,
        "32.99": 2999,
        "33": 35175,
        "33.1": 27448,
        "33.11": 1761,
        "33.12": 17760,
        "33.13": 1710,
        "33.14": 1281,
        "33.15": 3399,
        "33.16": 138,
        "33.17": 188,
        "33.19": 1211,
        "33.2": 7727,
        "33.20": 7727,
        "D": 13269,
        "35": 13269,
        "35.1": 12485,
        "35.11": 10845,
        "35.12": 9,
        "35.13": 175,
        "35.14": 1456,
        "35.2": 550,
        "35.21": 63,
        "35.22": 179,
        "35.23": 308,
        "35.3": 234,
        "35.30": 234,
        "E": 9120,
        "36": 750,
        "36.0": 750,
        "36.00": 750,
        "37": 1474,
        "37.0": 1474,
        "37.00": 1474,
        "38": 6211,
        "38.1": 1935,
        "38.11": 1624,
        "38.12": 311,
        "38.2": 1076,
        "38.21": 931,
        "38.22": 145,
        "38.3": 3200,
        "38.31": 484,
        "38.32": 2716,
        "39": 685,
        "39.0": 685,
        "39.00": 685,
        "F": 544886,
        "41": 145938,
        "41.1": 5013,
        "41.10": 5013,
        "41.2": 140925,
        "41.20": 140925,
        "42": 8020,
        "42.1": 4105,
        "42.11": 3832,
        "42.12": 205,
        "42.13": 68,
        "42.2": 774,
        "42.21": 336,
        "42.22": 438,
        "42.9": 3141,
        "42.91": 344,
        "42.99": 2797,
        "43": 390928,
        "43.1": 10681,
        "43.11": 1371,
        "43.12": 8176,
        "43.13": 1134,
        "43.2": 144130,
        "43.21": 67851,
        "43.22": 62600,
        "43.29": 13679,
        "43.3": 219632,
        "43.31": 7258,
        "43.32": 22770,
        "43.33": 20575,
        "43.34": 35337,
        "43.39": 133692,
        "43.9": 16485,
        "43.91": 5173,
        "43.99": 11312,
        "45": 121034,
        "45.1": 27891,
        "45.11": 26204,
        "45.19": 1687,
        "45.2": 74920,
        "45.20": 74920,
        "45.3": 11010,
        "45.31": 5556,
        "45.32": 5454,
        "45.4": 7213,
        "45.40": 7213,
        "46": 345876,
        "46.1": 191132,
        "46.11": 3362,
        "46.12": 5291,
        "46.13": 10104,
        "46.14": 8886,
        "46.15": 11424,
        "46.16": 10446,
        "46.17": 32307,
        "46.18": 47873,
        "46.19": 61439,
        "46.2": 6969,
        "46.22": 1441,
        "46.23": 1180,
        "46.24": 1109,
        "46.31": 7503,
        "46.32": 2357,
        "46.34": 5533,
        "46.35": 68,
        "46.4": 43425,
        "46.41": 2710,
        "46.42": 9693,
        "46.45": 3284,
        "46.46": 5664,
        "46.48": 2108,
        "46.49": 10714,
        "46.5": 6730,
        "46.52": 2117,
        "46.62": 2544,
        "46.65": 890,
        "46.66": 1593,
        "46.69": 10342,
        "46.7": 37210,
        "46.72": 3251,
        "46.75": 3186,
        "46.76": 3372,
        "46.77": 4174,
        "46.9": 8368,
        "46.90": 8368,
        "47": 517833,
        "47.1": 55518,
        "47.11": 43948,
        "47.19": 11570,
        "47.2": 90280,
        "47.21": 13365,
        "47.22": 19637,
        "47.23": 4957,
        "47.24": 4454,
        "47.25": 4263,
        "47.26": 30978,
        "47.29": 12626,
        "47.3": 12428,
        "47.30": 12428,
        "47.4": 10547,
        "47.41": 4977,
        "47.42": 4968,
        "47.43": 602,
        "47.5": 60003,
        "47.51": 8265,
        "47.52": 25009,
        "47.53": 1301,
        "47.54": 2596,
        "47.59": 22832,
        "47.6": 29715,
        "47.61": 3117,
        "47.62": 15817,
        "47.63": 484,
        "47.64": 7481,
        "47.65": 2816,
        "47.7": 168226,
        "47.71": 54535,
        "47.72": 10888,
        "47.73": 20210,
        "47.74": 3285,
        "47.75": 10625,
        "47.76": 15562,
        "47.77": 12299,
        "47.78": 36556,
        "47.79": 4266,
        "47.8": 63459,
        "47.81": 19893,
        "47.82": 24459,
        "47.89": 19107,
        "47.9": 27657,
        "47.91": 20418,
        "47.99": 7239,
        "H": 120213,
        "49": 90562,
        "49.1": 17,
        "49.10": 17,
        "49.2": 22,
        "49.20": 22,
        "49.3": 32255,
        "49.31": 911,
        "49.32": 27996,
        "49.39": 3348,
        "49.4": 58252,
        "49.41": 56836,
        "49.42": 1416,
        "49.5": 16,
        "49.50": 16,
        "50": 2586,
        "50.1": 1104,
        "50.10": 1104,
        "50.2": 149,
        "50.20": 149,
        "50.3": 1226,
        "50.30": 1226,
        "50.4": 107,
        "50.40": 107,
        "51": 131,
        "51.1": 112,
        "51.10": 112,
        "51.2": 19,
        "51.21": 19,
        "52": 21781,
        "52.1": 1489,
        "52.10": 1489,
        "52.2": 20292,
        "52.21": 6991,
        "52.22": 2005,
        "52.23": 430,
        "52.24": 3065,
        "52.29": 7801,
        "53": 5153,
        "53.1": 1,
        "53.10": 1,
        "53.2": 5152,
        "53.20": 5152,
        "I": 338404,
        "55": 68529,
        "55.1": 22729,
        "55.10": 22729,
        "55.2": 43935,
        "55.20": 43935,
        "55.3": 1618,
        "55.30": 1618,
        "55.9": 247,
        "55.90": 247,
        "56": 269875,
        "56.1": 165258,
        "56.10": 165258,
        "56.2": 3672,
        "56.21": 2222,
        "56.29": 1450,
        "56.3": 100945,
        "56.30": 100945,
        "J": 125648,
        "58": 5726,
        "58.1": 4886,
        "58.11": 1929,
        "58.12": 7,
        "58.13": 401,
        "58.14": 1423,
        "58.19": 1126,
        "58.2": 840,
        "58.21": 22,
        "58.29": 818,
        "59": 10716,
        "59.1": 8798,
        "59.11": 6134,
        "59.12": 1800,
        "59.13": 220,
        "59.14": 644,
        "59.2": 1918,
        "59.20": 1918,
        "60": 1267,
        "60.1": 702,
        "60.10": 702,
        "60.2": 565,
        "60.20": 565,
        "61": 3512,
        "61.1": 264,
        "61.10": 264,
        "61.2": 119,
        "61.20": 119,
        "61.3": 24,
        "61.30": 24,
        "61.9": 3105,
        "61.90": 3105,
        "62": 63327,
        "62.0": 63327,
        "62.01": 20750,
        "62.02": 29753,
        "62.03": 1063,
        "62.09": 11761,
        "63": 41100,
        "63.1": 37562,
        "63.11": 35135,
        "63.12": 2427,
        "63.9": 3538,
        "63.91": 509,
        "63.99": 3029,
        "K": 120248,
        "64": 23826,
        "64.1": 389,
        "64.11": 1,
        "64.19": 388,
        "64.2": 20522,
        "64.20": 20522,
        "64.3": 79,
        "64.30": 79,
        "64.9": 2836,
        "64.91": 144,
        "64.92": 477,
        "64.99": 2215,
        "65": 182,
        "65.1": 175,
        "65.11": 59,
        "65.12": 116,
        "65.2": 7,
        "65.20": 7,
        "66": 96240,
        "66.1": 39609,
        "66.11": 5,
        "66.12": 103,
        "66.19": 39501,
        "66.2": 56354,
        "66.21": 5097,
        "66.22": 50779,
        "66.29": 478,
        "66.3": 277,
        "66.30": 277,
        "L": 250621,
        "68": 250621,
        "68.1": 29568,
        "68.10": 29568,
        "68.2": 159486,
        "68.20": 159486,
        "68.3": 61567,
        "68.31": 42633,
        "68.32": 18934,
        "M": 934227,
        "69": 319307,
        "69.1": 182798,
        "69.10": 182798,
        "69.2": 136509,
        "69.20": 136509,
        "70": 117805,
        "70.1": 2544,
        "70.10": 2544,
        "70.2": 115261,
        "70.21": 9207,
        "70.22": 106054,
        "71.1": 237345,
        "71.11": 82790,
        "71.12": 154555,
        "71.2": 8385,
        "71.20": 8385,
        "72": 16440,
        "72.1": 14471,
        "72.11": 8369,
        "72.19": 6102,
        "72.2": 1969,
        "72.20": 1969,
        "73": 36336,
        "73.1": 30972,
        "73.11": 28986,
        "73.12": 1986,
        "73.2": 5364,
        "73.20": 5364,
        "74": 179867,
        "74.1": 47980,
        "74.10": 47980,
        "74.2": 19783,
        "74.20": 19783,
        "74.3": 8228,
        "74.30": 8228,
        "74.9": 103876,
        "74.90": 103876,
        "75": 18742,
        "75.0": 18742,
        "75.00": 18742,
        "N": 195326,
        "77": 17973,
        "77.1": 4677,
        "77.11": 4337,
        "77.12": 340,
        "77.2": 4202,
        "77.21": 3362,
        "77.22": 70,
        "77.29": 770,
        "77.3": 7969,
        "77.31": 123,
        "77.32": 1601,
        "77.33": 557,
        "77.34": 755,
        "77.35": 122,
        "77.39": 4811,
        "77.4": 1125,
        "77.40": 1125,
        "78": 1137,
        "78.1": 996,
        "78.10": 996,
        "78.2": 133,
        "78.20": 133,
        "78.3": 8,
        "78.30": 8,
        "79": 22215,
        "79.1": 9132,
        "79.11": 7451,
        "79.12": 1681,
        "79.9": 13083,
        "79.90": 13083,
        "80": 2415,
        "80.1": 1238,
        "80.10": 1238,
        "80.2": 209,
        "80.20": 209,
        "80.3": 968,
        "80.30": 968,
        "81": 64408,
        "81.1": 2853,
        "81.10": 2853,
        "81.2": 40117,
        "81.21": 32882,
        "81.22": 1845,
        "81.29": 5390,
        "81.3": 21438,
        "81.30": 21438,
        "82": 87178,
        "82.1": 12964,
        "82.11": 9054,
        "82.19": 3910,
        "82.2": 1360,
        "82.20": 1360,
        "82.3": 4892,
        "82.30": 4892,
        "82.9": 67962,
        "82.91": 1418,
        "82.92": 2258,
        "82.99": 64286,
        "P": 54890,
        "85": 54890,
        "85.1": 1468,
        "85.10": 1468,
        "85.2": 269,
        "85.20": 269,
        "85.3": 1405,
        "85.31": 344,
        "85.32": 1061,
        "85.4": 855,
        "85.41": 210,
        "85.42": 645,
        "85.5": 48632,
        "85.51": 15321,
        "85.52": 5543,
        "85.53": 4591,
        "85.59": 23177,
        "85.6": 2261,
        "85.60": 2261,
        "Q": 384369,
        "86": 364376,
        "86.1": 877,
        "86.10": 877,
        "86.2": 210234,
        "86.21": 65517,
        "86.22": 89820,
        "86.23": 54897,
        "86.9": 153265,
        "86.90": 153265,
        "87": 6777,
        "87.1": 930,
        "87.10": 930,
        "87.2": 573,
        "87.20": 573,
        "87.3": 3593,
        "87.30": 3593,
        "87.9": 1681,
        "87.90": 1681,
        "88": 13216,
        "88.1": 2941,
        "88.10": 2941,
        "88.9": 10275,
        "88.91": 4063,
        "88.99": 6212,
        "R": 91859,
        "90": 45822,
        "90.0": 45822,
        "90.01": 15585,
        "90.02": 9439,
        "90.03": 20516,
        "90.04": 282,
        "91": 1228,
        "91.0": 1228,
        "91.01": 594,
        "91.02": 304,
        "91.03": 230,
        "91.04": 100,
        "92": 7305,
        "92.0": 7305,
        "92.00": 7305,
        "93": 37504,
        "93.1": 20843,
        "93.11": 3894,
        "93.12": 2471,
        "93.13": 4470,
        "93.19": 10008,
        "93.2": 16661,
        "93.21": 1775,
        "93.29": 14886,
        "S": 226376,
        "95": 24924,
        "95.1": 5795,
        "95.11": 3378,
        "95.12": 2417,
        "95.2": 19129,
        "95.21": 959,
        "95.22": 3150,
        "95.23": 2324,
        "95.24": 5705,
        "95.25": 1571,
        "95.29": 5420,
        "96": 201452,
        "96.0": 201452,
        "96.01": 14453,
        "96.02": 141437,
        "96.03": 6857,
        "96.04": 2946,
        "96.09": 35759,
    }

    totale_istat = istat_db.get(cod_nace_pulito, 0)
    
    if totale_istat > 0:
        context['max_soc_istat'] = f"{totale_istat:,}".replace(',', '.')
    else:
        context['max_soc_istat'] = "N/D"

    # 3. Calcolo percentuale automatica
    try:
        if totale_istat > 0:
            percentuale = (totale_valide / totale_istat) * 100
            context['perc_su_istat'] = f"{percentuale:.2f}".replace('.', ',')
        else:
            context['perc_su_istat'] = "N/D"
    except (ValueError, TypeError, ZeroDivisionError):
        context['perc_su_istat'] = "N/D"

    # --- 

    if len(conteggi) > 0:
        rat1_nome = str(conteggi.index[0]).strip()
        rat1_num = conteggi.iloc[0]
        context['rating_piu_presente'] = rat1_nome
        context['rating_piu_pres'] = rat1_nome
        context['rat1_piu_pres_categ'] = rat1_nome
        context['rat1_piu_pres_num'] = f"{rat1_num:,}".replace(',', '.')
        context['rating_piu_pres_num_tot'] = f"{rat1_num:,}".replace(',', '.')

        if 'A' in rat1_nome: desc = 'Positivo / Solido'
        elif 'B' in rat1_nome: desc = 'Medio / Equilibrato'
        else: desc = 'Rischioso / Critico'
        context['assoc_desc_rating_lett'] = desc

    if len(conteggi) > 1:
        context['rat2_piu_pres_categ'] = str(conteggi.index[1]).strip()
        context['rat2_piu_pres_num'] = f"{conteggi.iloc[1]:,}".replace(',', '.')

    if len(conteggi) > 2:
        context['rat3_piu_pres_categ'] = str(conteggi.index[2]).strip()
        context['rat3_piu_pres_num'] = f"{conteggi.iloc[2]:,}".replace(',', '.')
    
    # =================================================================
    # VINCITORI DELLE SINGOLE AREE (Per il paragrafo descrittivo)
    # =================================================================
    
    # Economico
    eco_counts = df_rating['Rating Economico'].value_counts()
    if len(eco_counts) > 0:
        context['rat_eco_piu_pres'] = str(eco_counts.index[0]).strip()
        context['num_eco_piu_pres'] = f"{eco_counts.iloc[0]:,}".replace(',', '.')
    else:
        context['rat_eco_piu_pres'], context['num_eco_piu_pres'] = "N.D.", "0"

    # Patrimoniale
    patr_counts = df_rating['Rating Patrimoniale'].value_counts()
    if len(patr_counts) > 0:
        context['rat_patr_piu_pres'] = str(patr_counts.index[0]).strip()
        context['num_patr_piu_pres'] = f"{patr_counts.iloc[0]:,}".replace(',', '.')
    else:
        context['rat_patr_piu_pres'], context['num_patr_piu_pres'] = "N.D.", "0"

    # Finanziario
    fin_counts = df_rating['Rating Finanziario'].value_counts()
    if len(fin_counts) > 0:
        context['rat_fin_piu_pres'] = str(fin_counts.index[0]).strip()
        context['num_fin_piu_pres'] = f"{fin_counts.iloc[0]:,}".replace(',', '.')
    else:
        context['rat_fin_piu_pres'], context['num_fin_piu_pres'] = "N.D.", "0"
    
    # =================================================================
    # FORMA DELLA DISTRIBUZIONE (Asimmetria e Curtosi sui Margini)
    # Utilizziamo l'EBITDA 2024 come proxy rappresentativo della redditività
    # =================================================================
    col_ebitda = 'Margine EBITDA (*) % 2024'
    if col_ebitda in df_orbis.columns:
        skew_val = df_orbis[col_ebitda].skew()
        kurt_val = df_orbis[col_ebitda].kurt() # In Pandas > 0 è leptocurtica

        # Analisi Asimmetria
        if skew_val > 0.3:
            context['tipo_asimmetria'] = "un'asimmetria positiva"
            context['rel_media_mediana'] = "dall'evidente scostamento tra la media (maggiore) e la mediana"
        elif skew_val < -0.3:
            context['tipo_asimmetria'] = "un'asimmetria negativa"
            context['rel_media_mediana'] = "dall'evidente scostamento tra la media (minore) e la mediana"
        else:
            context['tipo_asimmetria'] = "una sostanziale simmetria"
            context['rel_media_mediana'] = "dal generale allineamento tra i valori di media e mediana"

        # Analisi Curtosi
        if kurt_val > 0.5:
            context['tipo_curtosi'] = "leptocurtiche"
        elif kurt_val < -0.5:
            context['tipo_curtosi'] = "platicurtiche"
        else:
            context['tipo_curtosi'] = "mesocurtiche"
    else:
        context['tipo_asimmetria'] = "dati non sufficienti"
        context['rel_media_mediana'] = "N.D."
        context['tipo_curtosi'] = "non definibili"

    # =================================================================
    # FORMA DELLA DISTRIBUZIONE (Asimmetria e Curtosi Patrimoniale)
    # Utilizziamo l'Indice di Struttura 1° livello come proxy rappresentativo
    # =================================================================
    col_strut = 'Indice di Struttura 1° livello (*) 2024'
    if col_strut in df_orbis.columns:
        skew_patr = df_orbis[col_strut].skew()
        kurt_patr = df_orbis[col_strut].kurt()

        # Analisi Asimmetria Patrimoniale
        if skew_patr > 0.3:
            context['tipo_asimmetria_patr'] = "asimmetriche, in questo caso positive,"
        elif skew_patr < -0.3:
            context['tipo_asimmetria_patr'] = "asimmetriche, in questo caso negative,"
        else:
            context['tipo_asimmetria_patr'] = "sostanzialmente simmetriche"

        # Analisi Curtosi Patrimoniale
        if kurt_patr > 0.5:
            context['tipo_curtosi_patr'] = "leptocurtiche"
        elif kurt_patr < -0.5:
            context['tipo_curtosi_patr'] = "platicurtiche"
        else:
            context['tipo_curtosi_patr'] = "mesocurtiche"
    else:
        context['tipo_asimmetria_patr'] = "non valutabili"
        context['tipo_curtosi_patr'] = "non definibili"
    
    # =================================================================
    # FORMA DELLA DISTRIBUZIONE (Asimmetria e Curtosi Finanziaria)
    # Utilizziamo il Current Ratio come proxy rappresentativo
    # =================================================================
    col_fin = 'Current Ratio (*) 2024'
    if col_fin in df_orbis.columns:
        skew_fin = df_orbis[col_fin].skew()
        kurt_fin = df_orbis[col_fin].kurt()

        # Analisi Asimmetria Finanziaria
        if skew_fin > 0.3:
            context['tipo_asimmetria_fin'] = "un'asimmetria positiva"
        elif skew_fin < -0.3:
            context['tipo_asimmetria_fin'] = "un'asimmetria negativa"
        else:
            context['tipo_asimmetria_fin'] = "una sostanziale simmetria"

        # Analisi Curtosi Finanziaria
        if kurt_fin > 0.5:
            context['tipo_curtosi_fin'] = "leptocurtiche"
        elif kurt_fin < -0.5:
            context['tipo_curtosi_fin'] = "platicurtiche"
        else:
            context['tipo_curtosi_fin'] = "mesocurtiche"
    else:
        context['tipo_asimmetria_fin'] = "non valutabili"
        context['tipo_curtosi_fin'] = "non definibili"
    
    # ---

    # 🟢 ATTIVIAMO LA LAVATRICE NUCLEARE
    template_pulito = lavatrice_nucleare(template_path)
    doc = DocxTemplate(template_pulito)

    # =================================================================
    # COSTRUZIONE TABELLA 1 DINAMICA DA PYTHON
    # =================================================================
    sd_tab1 = doc.new_subdoc()
    t1 = sd_tab1.add_table(rows=1, cols=9)
    t1.style = 'Table Grid' # Stile classico a griglia di Word

    # 1. Creiamo le Intestazioni in Grassetto
    intestazioni = ['Regioni', 'Imprese V.A.', 'Imprese %', 'Ricavi V.A.', 'Ricavi %', 'Attivo V.A.', 'Attivo %', 'Dip. V.A.', 'Dip. %']
    for i, intestazione in enumerate(intestazioni):
        t1.rows[0].cells[i].text = intestazione
        t1.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    # 2. Calcoli Territoriali
    df_terr = df_orbis.copy()
    def pulisci_regione(x):
        if pd.isna(x): return 'Altro'
        return str(x).split(' - ')[1] if ' - ' in str(x) else str(x)
    df_terr['Reg_Clean'] = df_terr[col_regione].apply(pulisci_regione) if col_regione else 'Altro'

    tot_imp = len(df_terr)
    tot_ric = df_terr['Totale valore della produzione migl EUR 2024'].sum()
    tot_att = df_terr['Totale Attivo migl EUR 2024'].sum()
    tot_dip = df_terr['Numero dipendenti 2024'].sum()

    pivot_terr = df_terr.groupby(['Macroregione', 'Reg_Clean']).agg({
        col_ragione: 'count', 'Totale valore della produzione migl EUR 2024': 'sum', 
        'Totale Attivo migl EUR 2024': 'sum', 'Numero dipendenti 2024': 'sum'
    }).reset_index()

    # 3. Compilazione Righe
    for macro in ['Nord Ovest', 'Nord Est', 'Centro', 'Sud e Isole', 'Altro']:
        df_m = pivot_terr[pivot_terr['Macroregione'] == macro]
        if df_m.empty: continue
        
       # Inserimento Singole Regioni
        for _, r in df_m.iterrows():
            row_cells = t1.add_row().cells
            nome_regione = str(r['Reg_Clean'])
            
            row_cells[0].text = nome_regione
            row_cells[1].text = f"{int(r[col_ragione]):,}".replace(',', '.')
            row_cells[2].text = format_euro((r[col_ragione]/tot_imp)*100) if tot_imp else "0,00"
            row_cells[3].text = format_euro(r['Totale valore della produzione migl EUR 2024'])
            row_cells[4].text = format_euro((r['Totale valore della produzione migl EUR 2024']/tot_ric)*100) if tot_ric else "0,00"
            row_cells[5].text = format_euro(r['Totale Attivo migl EUR 2024'])
            row_cells[6].text = format_euro((r['Totale Attivo migl EUR 2024']/tot_att)*100) if tot_att else "0,00"
            row_cells[7].text = f"{int(r['Numero dipendenti 2024']):,}".replace(',', '.')
            row_cells[8].text = format_euro((r['Numero dipendenti 2024']/tot_dip)*100) if tot_dip else "0,00"

            # 🎯 EVIDENZIA LA REGIONE DELL'AZIENDA (Sfondo azzurro e grassetto)
            if nome_regione.strip().lower() == regione_target_pulita.strip().lower():
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                for cell in row_cells:
                    shd = parse_xml(r'<w:shd %s w:val="clear" w:color="auto" w:fill="DDEBF7"/>' % nsdecls('w'))
                    cell._tc.get_or_add_tcPr().append(shd)
                    if len(cell.paragraphs[0].runs) > 0:
                        cell.paragraphs[0].runs[0].bold = True
            
        # Inserimento Totale Macroregione (Grassetto)
        row_cells = t1.add_row().cells
        row_cells[0].text = f"TOTALE {macro.upper()}"
        row_cells[0].paragraphs[0].runs[0].bold = True
        row_cells[1].text = f"{int(df_m[col_ragione].sum()):,}".replace(',', '.')
        row_cells[2].text = format_euro((df_m[col_ragione].sum()/tot_imp)*100) if tot_imp else "0,00"
        row_cells[3].text = format_euro(df_m['Totale valore della produzione migl EUR 2024'].sum())
        row_cells[4].text = format_euro((df_m['Totale valore della produzione migl EUR 2024'].sum()/tot_ric)*100) if tot_ric else "0,00"
        row_cells[5].text = format_euro(df_m['Totale Attivo migl EUR 2024'].sum())
        row_cells[6].text = format_euro((df_m['Totale Attivo migl EUR 2024'].sum()/tot_att)*100) if tot_att else "0,00"
        row_cells[7].text = f"{int(df_m['Numero dipendenti 2024'].sum()):,}".replace(',', '.')
        row_cells[8].text = format_euro((df_m['Numero dipendenti 2024'].sum()/tot_dip)*100) if tot_dip else "0,00"

    # Inserimento Totale Italia Finale (Tutto in Grassetto)
    row_cells = t1.add_row().cells
    row_cells[0].text = "ITALIA (TOTALE)"
    row_cells[1].text = f"{int(tot_imp):,}".replace(',', '.')
    row_cells[2].text = "100,00"
    row_cells[3].text = format_euro(tot_ric)
    row_cells[4].text = "100,00"
    row_cells[5].text = format_euro(tot_att)
    row_cells[6].text = "100,00"
    row_cells[7].text = f"{int(tot_dip):,}".replace(',', '.')
    row_cells[8].text = "100,00"
    for cell in row_cells: cell.paragraphs[0].runs[0].bold = True

    for row in t1.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10.5)

    # Iniettiamo la tabella nel documento
    context['tabella_1_dinamica'] = sd_tab1


    # =================================================================
    # 🏗️ COSTRUZIONE TABELLA 2: CLASSIFICAZIONE POSIZIONAMENTO
    # =================================================================
    sd_tab2 = doc.new_subdoc()
    t2 = sd_tab2.add_table(rows=3, cols=4)
    t2.style = 'Table Grid'

    # Riga 0: Ragione Sociale (Uniamo le 4 celle in una sola)
    cella_unita = t2.cell(0, 0)
    cella_unita.merge(t2.cell(0, 3))
    cella_unita.text = context.get('ragione_sociale', 'Azienda Target')
    cella_unita.paragraphs[0].runs[0].bold = True

    # Riga 1: Intestazioni
    headers_t2 = ['Benchmark Economico', 'Benchmark Patrimoniale', 'Benchmark Finanziario', 'Rating Combinato']
    for i, h in enumerate(headers_t2):
        t2.cell(1, i).text = h
        t2.cell(1, i).paragraphs[0].runs[0].bold = True

    # Riga 2: Valori (es. A, B, C, ABC)
    valori_t2 = [
        context.get('rating_eco', 'N.D.'), 
        context.get('rating_patr', 'N.D.'), 
        context.get('rating_fin', 'N.D.'), 
        context.get('rating_tot', 'N.D.')
    ]
    for i, v in enumerate(valori_t2):
        t2.cell(2, i).text = str(v)

    # Iniettiamo la tabella nel documento
    context['tabella_2_dinamica'] = sd_tab2

    # =================================================================
    # 🏗️ COSTRUZIONE TABELLE 3, 4, 5: RANKING AZIENDALI
    # =================================================================

    # --- TABELLA 3: ECONOMICO ---
    sd_tab3 = doc.new_subdoc()
    t3 = sd_tab3.add_table(rows=4, cols=4)
    t3.style = 'Table Grid'
    t3.cell(0, 1).merge(t3.cell(0, 3))
    t3.cell(0, 1).text = 'Equilibrio Economico'
    t3.cell(0, 1).paragraphs[0].runs[0].bold = True

    t3.cell(1, 1).text = "Margine EBITDA"
    t3.cell(1, 2).text = "Margine EBIT"
    t3.cell(1, 3).text = "Margine di Profitto"
    for i in range(1, 4):
        t3.cell(1, i).paragraphs[0].runs[0].bold = True

    t3.cell(2, 0).text = 'Ranking Nazionale'
    t3.cell(2, 0).paragraphs[0].runs[0].bold = True
    t3.cell(2, 1).text = f"{context.get('rnk_naz_ebitda', 'N.D.')}°"
    t3.cell(2, 2).text = f"{context.get('rnk_naz_ebit', 'N.D.')}°"
    t3.cell(2, 3).text = f"{context.get('rnk_naz_prof', 'N.D.')}°"

    t3.cell(3, 0).text = 'Ranking Regionale'
    t3.cell(3, 0).paragraphs[0].runs[0].bold = True
    t3.cell(3, 1).text = f"{context.get('rnk_reg_ebitda', 'N.D.')}°"
    t3.cell(3, 2).text = f"{context.get('rnk_reg_ebit', 'N.D.')}°"
    t3.cell(3, 3).text = f"{context.get('rnk_reg_prof', 'N.D.')}°"
    context['tabella_3_dinamica'] = sd_tab3

    # --- TABELLA 4: PATRIMONIALE ---
    sd_tab4 = doc.new_subdoc()
    t4 = sd_tab4.add_table(rows=4, cols=4)
    t4.style = 'Table Grid'
    t4.cell(0, 1).merge(t4.cell(0, 3))
    t4.cell(0, 1).text = 'Equilibrio Patrimoniale'
    t4.cell(0, 1).paragraphs[0].runs[0].bold = True

    t4.cell(1, 1).text = "Indice Struttura 1° Liv."
    t4.cell(1, 2).text = "Indice Struttura 2° Liv."
    t4.cell(1, 3).text = "Gearing"
    for i in range(1, 4):
        t4.cell(1, i).paragraphs[0].runs[0].bold = True

    t4.cell(2, 0).text = 'Ranking Nazionale'
    t4.cell(2, 0).paragraphs[0].runs[0].bold = True
    t4.cell(2, 1).text = f"{context.get('rnk_naz_strut1', 'N.D.')}°"
    t4.cell(2, 2).text = f"{context.get('rnk_naz_strut2', 'N.D.')}°"
    t4.cell(2, 3).text = f"{context.get('rnk_naz_gear', 'N.D.')}°"

    t4.cell(3, 0).text = 'Ranking Regionale'
    t4.cell(3, 0).paragraphs[0].runs[0].bold = True
    t4.cell(3, 1).text = f"{context.get('rnk_reg_strut1', 'N.D.')}°"
    t4.cell(3, 2).text = f"{context.get('rnk_reg_strut2', 'N.D.')}°"
    t4.cell(3, 3).text = f"{context.get('rnk_reg_gear', 'N.D.')}°"
    context['tabella_4_dinamica'] = sd_tab4

    # --- TABELLA 5: FINANZIARIO ---
    sd_tab5 = doc.new_subdoc()
    t5 = sd_tab5.add_table(rows=4, cols=4)
    t5.style = 'Table Grid'
    t5.cell(0, 1).merge(t5.cell(0, 3))
    t5.cell(0, 1).text = 'Equilibrio Finanziario'
    t5.cell(0, 1).paragraphs[0].runs[0].bold = True

    t5.cell(1, 1).text = "Current Ratio"
    t5.cell(1, 2).text = "Quick Ratio"
    t5.cell(1, 3).text = "Indice Rot.Cap.Inv."
    for i in range(1, 4):
        t5.cell(1, i).paragraphs[0].runs[0].bold = True

    t5.cell(2, 0).text = 'Ranking Nazionale'
    t5.cell(2, 0).paragraphs[0].runs[0].bold = True
    t5.cell(2, 1).text = f"{context.get('rnk_naz_cr', 'N.D.')}°"
    t5.cell(2, 2).text = f"{context.get('rnk_naz_qr', 'N.D.')}°"
    t5.cell(2, 3).text = f"{context.get('rnk_naz_rot', 'N.D.')}°"

    t5.cell(3, 0).text = 'Ranking Regionale'
    t5.cell(3, 0).paragraphs[0].runs[0].bold = True
    t5.cell(3, 1).text = f"{context.get('rnk_reg_cr', 'N.D.')}°"
    t5.cell(3, 2).text = f"{context.get('rnk_reg_qr', 'N.D.')}°"
    t5.cell(3, 3).text = f"{context.get('rnk_reg_rot', 'N.D.')}°"
    context['tabella_5_dinamica'] = sd_tab5


    # =================================================================
    # 🏗️ COSTRUZIONE TABELLA 6: RIPARTIZIONE TERRITORIALE BENCHMARK TOTALE
    # =================================================================
    sd_tab6 = doc.new_subdoc()
    t6 = sd_tab6.add_table(rows=2, cols=9)
    t6.style = 'Table Grid'

    # Uniamo le celle per le intestazioni della riga superiore
    t6.cell(0, 0).merge(t6.cell(1, 0))
    t6.cell(0, 0).text = 'Regioni'
    
    t6.cell(0, 1).merge(t6.cell(0, 2))
    t6.cell(0, 1).text = 'Totale Imprese'
    
    t6.cell(0, 3).merge(t6.cell(0, 4))
    t6.cell(0, 3).text = 'Classe A'
    
    t6.cell(0, 5).merge(t6.cell(0, 6))
    t6.cell(0, 5).text = 'Classe B'
    
    t6.cell(0, 7).merge(t6.cell(0, 8))
    t6.cell(0, 7).text = 'Classe C'

    # Sotto-intestazioni (V.A. e %)
    for col_idx in [1, 3, 5, 7]:
        t6.cell(1, col_idx).text = 'V.A.'
        t6.cell(1, col_idx+1).text = '%'

    # Formattiamo le intestazioni in grassetto
    for r_idx in [0, 1]:
        for c_idx in range(9):
            if len(t6.cell(r_idx, c_idx).paragraphs[0].runs) > 0:
                t6.cell(r_idx, c_idx).paragraphs[0].runs[0].bold = True

    # Calcoli per la Tabella 6
    df_t6 = df_rating.copy()
    def pulisci_reg(x):
        if pd.isna(x): return 'Altro'
        return str(x).split(' - ')[1] if ' - ' in str(x) else str(x)
    
    df_t6['Reg_Clean'] = df_t6[col_regione].apply(pulisci_reg) if col_regione else 'Altro'

    if 'Benchmark Totale' in df_t6.columns:
        tot_ita = len(df_t6)
        tot_ita_A = len(df_t6[df_t6['Benchmark Totale'] == 'A'])
        tot_ita_B = len(df_t6[df_t6['Benchmark Totale'] == 'B'])
        tot_ita_C = len(df_t6[df_t6['Benchmark Totale'] == 'C'])
        
        for macro in ['Nord Ovest', 'Nord Est', 'Centro', 'Sud e Isole', 'Altro']:
            df_m = df_t6[df_t6['Macroregione'] == macro]
            if df_m.empty: continue
            
            regioni_macro = df_m['Reg_Clean'].unique()
            for reg in sorted(regioni_macro):
                df_reg = df_m[df_m['Reg_Clean'] == reg]
                val_tot = len(df_reg)
                val_A = len(df_reg[df_reg['Benchmark Totale'] == 'A'])
                val_B = len(df_reg[df_reg['Benchmark Totale'] == 'B'])
                val_C = len(df_reg[df_reg['Benchmark Totale'] == 'C'])
                
                row_cells = t6.add_row().cells
                row_cells[0].text = str(reg)
                row_cells[1].text = f"{val_tot:,}".replace(',', '.')
                row_cells[2].text = format_euro((val_tot/tot_ita)*100) if tot_ita else "0,00"
                row_cells[3].text = f"{val_A:,}".replace(',', '.')
                row_cells[4].text = format_euro((val_A/tot_ita)*100) if tot_ita else "0,00"
                row_cells[5].text = f"{val_B:,}".replace(',', '.')
                row_cells[6].text = format_euro((val_B/tot_ita)*100) if tot_ita else "0,00"
                row_cells[7].text = f"{val_C:,}".replace(',', '.')
                row_cells[8].text = format_euro((val_C/tot_ita)*100) if tot_ita else "0,00"
                
            # Riga Totale Macroregione (Grassetto)
            val_tot_m = len(df_m)
            val_A_m = len(df_m[df_m['Benchmark Totale'] == 'A'])
            val_B_m = len(df_m[df_m['Benchmark Totale'] == 'B'])
            val_C_m = len(df_m[df_m['Benchmark Totale'] == 'C'])
            
            row_cells = t6.add_row().cells
            row_cells[0].text = f"TOTALE {macro.upper()}"
            row_cells[1].text = f"{val_tot_m:,}".replace(',', '.')
            row_cells[2].text = format_euro((val_tot_m/tot_ita)*100) if tot_ita else "0,00"
            row_cells[3].text = f"{val_A_m:,}".replace(',', '.')
            row_cells[4].text = format_euro((val_A_m/tot_ita)*100) if tot_ita else "0,00"
            row_cells[5].text = f"{val_B_m:,}".replace(',', '.')
            row_cells[6].text = format_euro((val_B_m/tot_ita)*100) if tot_ita else "0,00"
            row_cells[7].text = f"{val_C_m:,}".replace(',', '.')
            row_cells[8].text = format_euro((val_C_m/tot_ita)*100) if tot_ita else "0,00"
            
            for cell in row_cells: cell.paragraphs[0].runs[0].bold = True

        # Riga Totale Italia (Grassetto)
        row_cells = t6.add_row().cells
        row_cells[0].text = "ITALIA (TOTALE)"
        row_cells[1].text = f"{tot_ita:,}".replace(',', '.')
        row_cells[2].text = "100,00"
        row_cells[3].text = f"{tot_ita_A:,}".replace(',', '.')
        row_cells[4].text = format_euro((tot_ita_A/tot_ita)*100) if tot_ita else "0,00"
        row_cells[5].text = f"{tot_ita_B:,}".replace(',', '.')
        row_cells[6].text = format_euro((tot_ita_B/tot_ita)*100) if tot_ita else "0,00"
        row_cells[7].text = f"{tot_ita_C:,}".replace(',', '.')
        row_cells[8].text = format_euro((tot_ita_C/tot_ita)*100) if tot_ita else "0,00"
        
        for cell in row_cells: cell.paragraphs[0].runs[0].bold = True

    context['tabella_6_dinamica'] = sd_tab6

    # =================================================================
    # 🏗️ COSTRUZIONE TABELLA 7 E GRAFICI (COMPOSIZIONE VdP)
    # =================================================================
    # 1. RECUPERO DATI (Calcolo incidenza percentuale corretto e ANTI-CRASH)
    anni = ['2021', '2022', '2023', '2024']
    metriche = [
        ('Costo del venduto', 'Costo del venduto migl EUR'),
        ('Oneri di gestione', 'Oneri diversi di gestione migl EUR'),
        ('Proventi/Oneri fin.', 'Proventi/oneri finanziari migl EUR'),
        ('Imposte', 'Totale imposte migl EUR'),
        ('Utile/Perdita Netta', 'Utile/Perdita al netto delle imposte migl EUR')
    ]
    col_prod_prefisso = 'Totale valore della produzione migl EUR'

    dati_settore = {m[0]: [] for m in metriche}
    dati_azienda = {m[0]: [] for m in metriche}

    for nome_metrica, col_base in metriche:
        for anno in anni:
            col_num = f"{col_base} {anno}"
            col_den = f"{col_prod_prefisso} {anno}"

            # --- MEDIANA SETTORE (Blindata contro Zeri e Lettere) ---
            s_num = pd.to_numeric(df_orbis[col_num] if col_num in df_orbis.columns else pd.Series(dtype=float), errors='coerce')
            s_den = pd.to_numeric(df_orbis[col_den] if col_den in df_orbis.columns else pd.Series(dtype=float), errors='coerce')

            if not s_num.empty and not s_den.empty:
                # Trasforma gli zeri in NaN per non far esplodere la divisione
                s_den_safe = s_den.replace(0, np.nan)
                pct_series = (s_num / s_den_safe).replace([np.inf, -np.inf], np.nan).dropna()
                val_sett_pct = (pct_series.median() * 100) if not pct_series.empty else 0.0
            else:
                val_sett_pct = 0.0
            
            # Filtro di sicurezza estrema prima del grafico
            if pd.isna(val_sett_pct) or val_sett_pct == np.inf or val_sett_pct == -np.inf:
                val_sett_pct = 0.0
                
            dati_settore[nome_metrica].append(val_sett_pct)

            # --- VALORE AZIENDA TARGET (Blindato contro Zeri e Lettere) ---
            if not df_target.empty and col_num in df_target.columns and col_den in df_target.columns:
                try:
                    v_num = float(df_target.iloc[0].get(col_num, 0))
                    v_den = float(df_target.iloc[0].get(col_den, 1))
                    
                    if pd.isna(v_num): v_num = 0.0
                    if pd.isna(v_den) or v_den == 0: v_den = 1.0 # Evita div/0
                    
                    val_az_pct = (v_num / v_den) * 100
                except:
                    val_az_pct = 0.0
            else:
                val_az_pct = 0.0

            # Filtro di sicurezza estrema prima del grafico
            if pd.isna(val_az_pct) or val_az_pct == np.inf or val_az_pct == -np.inf:
                val_az_pct = 0.0

            dati_azienda[nome_metrica].append(val_az_pct)
            

        # 2. COSTRUZIONE TABELLA 7 (Word) - Layout Split-Panel #?
        sd_tab7 = doc.new_subdoc()
        t7 = sd_tab7.add_table(rows=1, cols=5)
        t7.style = 'Table Grid'
        
        # Allineamento globale della tabella al centro
        tbl_pr = t7._tbl.tblPr
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        tbl_pr.append(jc)

        # --- INTESTAZIONE PRINCIPALE ---
        headers_t7 = ['% su Valore Produzione', '2021', '2022', '2023', '2024']
        for i, h in enumerate(headers_t7):
            cell = t7.cell(0, i)
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            
            # Testo a sinistra per la prima colonna, a DESTRA per gli anni
            if i == 0:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            # Sfondo grigio scuro/elegante per l'intestazione
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            shd = parse_xml(r'<w:shd %s w:val="clear" w:color="auto" w:fill="E7E6E6"/>' % nsdecls('w'))
            cell._tc.get_or_add_tcPr().append(shd)

        # ==========================================
        # BLOCCO 1: AZIENDA
        # ==========================================
        # Riga di divisione (Unisce tutte le celle)
        row_az_title = t7.add_row().cells
        row_az_title[0].merge(row_az_title[4])
        row_az_title[0].text = f"AZIENDA: {context.get('ragione_sociale', 'Azienda Target').upper()}"
        row_az_title[0].paragraphs[0].runs[0].bold = True
        row_az_title[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Sfondo azzurrino per il titolo dell'azienda
        shd_az = parse_xml(r'<w:shd %s w:val="clear" w:color="auto" w:fill="DDEBF7"/>' % nsdecls('w'))
        row_az_title[0]._tc.get_or_add_tcPr().append(shd_az)

        # Inserimento metriche Azienda
        for nome_metrica, _ in metriche:
            row_az = t7.add_row().cells
            row_az[0].text = nome_metrica
            row_az[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            for i, val in enumerate(dati_azienda[nome_metrica]):
                testo_val = f"{val:.2f}%".replace('.', ',') if pd.notna(val) else "N.D."
                row_az[i+1].text = testo_val
                # Numeri allineati rigorosamente a destra
                row_az[i+1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # ==========================================
        # BLOCCO 2: MEDIANA SETTORE
        # ==========================================
        # Riga di divisione (Unisce tutte le celle)
        row_set_title = t7.add_row().cells
        row_set_title[0].merge(row_set_title[4])
        row_set_title[0].text = "MEDIANA DI SETTORE"
        row_set_title[0].paragraphs[0].runs[0].bold = True
        row_set_title[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Sfondo grigio tenue per il titolo del settore
        shd_set = parse_xml(r'<w:shd %s w:val="clear" w:color="auto" w:fill="F2F2F2"/>' % nsdecls('w'))
        row_set_title[0]._tc.get_or_add_tcPr().append(shd_set)

        # Inserimento metriche Settore
        for nome_metrica, _ in metriche:
            row_sett = t7.add_row().cells
            row_sett[0].text = nome_metrica
            row_sett[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            for i, val in enumerate(dati_settore[nome_metrica]):
                testo_val = f"{val:.2f}%".replace('.', ',') if pd.notna(val) else "N.D."
                row_sett[i+1].text = testo_val
                # Numeri allineati rigorosamente a destra
                row_sett[i+1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        context['tabella_7_dinamica'] = sd_tab7

    # 3. GENERAZIONE GRAFICI (Matplotlib)
    def genera_grafico_barre_impilate(dati, titolo):
        fig, ax = plt.subplots(figsize=(8, 4.5))
        
        # Colori professionali (toni di blu, grigio, verde, ecc.)
        colori = ['#2b3a67', '#496a81', '#66999b', '#b3af8f', '#ffc482']
        
        bottoms = np.zeros(len(anni))
        for i, (nome, valori) in enumerate(dati.items()):
            # Sostituiamo eventuali NaN con 0 per non far crashare il grafico
            valori_puliti = [v if pd.notna(v) else 0 for v in valori]
            ax.bar(anni, valori_puliti, bottom=bottoms, label=nome, color=colori[i], edgecolor='white')
            bottoms += np.array(valori_puliti)
            
        ax.set_title(titolo, fontsize=12, fontweight='bold', color='#333333')
        ax.set_ylabel('% sul Valore della Produzione', fontsize=10)
        ax.legend(loc='upper left', bbox_to_anchor=(1, 1), frameon=False)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.grid(axis='y', linestyle='--', alpha=0.7)
        #plt.tight_layout()
        
        # Salviamo l'immagine in memoria (senza creare file sul pc)
        mem_img = io.BytesIO()

        aggiungi_watermark_fig(fig, modalita_teaser) # <--- AGGIUNGI ", modalita_teaser"

        plt.savefig(mem_img, format='png', dpi=(150 if modalita_teaser else 300), bbox_inches='tight')
        plt.close(fig)
        mem_img.seek(0)
        return mem_img

    # Creazione Immagini e Iniezione nel Word
    img_settore_buf = genera_grafico_barre_impilate(dati_settore, "Composizione % - Mediane di Settore")
    context['grafico_settore'] = InlineImage(doc, img_settore_buf, width=Mm(155))

    img_azienda_buf = genera_grafico_barre_impilate(dati_azienda, f"Composizione % - {context.get('ragione_sociale', 'Azienda')}")
    context['grafico_azienda'] = InlineImage(doc, img_azienda_buf, width=Mm(155))
    # =================================================================

    # =================================================================
    # 🏗️ COSTRUZIONE TABELLA 8 E GRAFICO 3 (TREND MARGINE EBITDA)
    # =================================================================
    
    anni = ['2021', '2022', '2023', '2024']
    col_base_ebitda = 'Margine EBITDA (*) %'
    
    valori_settore_ebitda = []
    valori_azienda_ebitda = []
    
    # 1. Recupero Dati
    for anno in anni:
        colonna = f"{col_base_ebitda} {anno}"
        
        # Mediana Settore
        if colonna in df_orbis.columns:
            valori_settore_ebitda.append(df_orbis[colonna].median())
        else:
            valori_settore_ebitda.append(0.0)
            
        # Valore Azienda
        if not df_target.empty and colonna in df_target.columns:
            valori_azienda_ebitda.append(df_target.iloc[0][colonna])
        else:
            valori_azienda_ebitda.append(0.0)

    # 2. Costruzione Tabella 8
    sd_tab8 = doc.new_subdoc()
    t8 = sd_tab8.add_table(rows=3, cols=5)
    t8.style = 'Table Grid'
    
    # Intestazioni anni
    headers_t8 = ['', '2021', '2022', '2023', '2024']
    for i, h in enumerate(headers_t8):
        t8.cell(0, i).text = h
        t8.cell(0, i).paragraphs[0].runs[0].bold = True
        
    nome_azienda = context.get('ragione_sociale', 'Azienda Target')
    codice_nace_report = context.get('codice_nace', 'N.D.')
    
    # Riga Azienda
    t8.cell(1, 0).text = nome_azienda
    t8.cell(1, 0).paragraphs[0].runs[0].bold = True
    for i, val in enumerate(valori_azienda_ebitda):
        t8.cell(1, i+1).text = f"{val:.2f}".replace('.', ',') if pd.notna(val) else "N.D."
        
    # Riga Settore (Mediana)
    t8.cell(2, 0).text = f"Settore {codice_nace_report} (Mediana)"
    t8.cell(2, 0).paragraphs[0].runs[0].bold = True
    for i, val in enumerate(valori_settore_ebitda):
        t8.cell(2, i+1).text = f"{val:.2f}".replace('.', ',') if pd.notna(val) else "N.D."
        
    context['tabella_8_dinamica'] = sd_tab8

    # 3. Funzione per generare Grafici di Trend (Linee)
    def genera_grafico_trend(anni, val_azienda, val_settore, nome_az, titolo):
        fig, ax = plt.subplots(figsize=(8, 4.5))
        
        # Pulizia dati per il grafico (sostituiamo eventuali NaN con 0)
        val_az_clean = [v if pd.notna(v) else 0 for v in val_azienda]
        val_set_clean = [v if pd.notna(v) else 0 for v in val_settore]
        
        # Disegno delle due linee
        ax.plot(anni, val_az_clean, marker='o', linewidth=2.5, markersize=8, color='#2b3a67', label=nome_az)
        ax.plot(anni, val_set_clean, marker='s', linewidth=2, markersize=7, color='#b3af8f', linestyle='--', label='Mediana Settore')
        
        ax.set_title(titolo, fontsize=12, fontweight='bold', color='#333333')
        ax.set_ylabel('%', fontsize=10)
        ax.legend(loc='best', frameon=False)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.grid(axis='y', linestyle='--', alpha=0.5)
        #plt.tight_layout()
        
        # Salvataggio in memoria
        mem_img = io.BytesIO()

        aggiungi_watermark_fig(fig, modalita_teaser) # <--- AGGIUNGI ", modalita_teaser"

        plt.savefig(mem_img, format='png', dpi=(150 if modalita_teaser else 300), bbox_inches='tight')
        plt.close(fig)
        mem_img.seek(0)
        return mem_img

    # Generazione Immagine e Iniezione nel Word
    img_ebitda_buf = genera_grafico_trend(
        anni, valori_azienda_ebitda, valori_settore_ebitda, 
        nome_azienda, "Andamento Margine EBITDA (%)"
    )
    context['grafico_ebitda'] = InlineImage(doc, img_ebitda_buf, width=Mm(155))

    # =================================================================
    # 🏗️ COSTRUZIONE GRAFICO 4 (CONFRONTO DIRETTO EBITDA 2024)
    # =================================================================
    
    # Funzione riutilizzabile per grafici a barre comparative (Azienda vs Settore per 1 singolo anno)
    def genera_grafico_confronto_singolo(val_az, val_set, nome_az, titolo):
        fig, ax = plt.subplots(figsize=(8, 5.5))
        
        etichette = [nome_az, 'Mediana Settore']
        
        # Pulizia da eventuali NaN
        v_az_clean = val_az if pd.notna(val_az) else 0
        v_set_clean = val_set if pd.notna(val_set) else 0
        valori = [v_az_clean, v_set_clean]
        
        colori = ['#2b3a67', '#b3af8f']
        
        # Disegno delle barre
        bars = ax.bar(etichette, valori, color=colori, edgecolor='white', width=0.5)

        for bar in bars:
            yval = bar.get_height()
            if yval != 0:
                ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    yval / 2,
                    f"{yval:.2f}%".replace('.', ','),
                    ha='center', va='center',
                    fontsize=9, fontweight='bold', color='white'
                )
        
            
        ax.set_title(titolo, fontsize=11, fontweight='bold', color='#333333')
        ax.set_ylabel('%', fontsize=10)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.grid(axis='y', linestyle='--', alpha=0.5)
        #plt.tight_layout()
        
        # Salvataggio in memoria
        mem_img = io.BytesIO()

        aggiungi_watermark_fig(fig, modalita_teaser) # <--- AGGIUNGI ", modalita_teaser"
        
        plt.savefig(mem_img, format='png', dpi=(150 if modalita_teaser else 300), bbox_inches='tight')
        plt.close(fig)
        mem_img.seek(0)
        return mem_img

    # Prendiamo i dati del 2024 (che è l'ultimo valore della lista calcolata prima, quindi indice -1)
    val_az_ebitda_24 = valori_azienda_ebitda[-1] if len(valori_azienda_ebitda) > 0 else 0
    val_set_ebitda_24 = valori_settore_ebitda[-1] if len(valori_settore_ebitda) > 0 else 0

    # Generazione Immagine e Iniezione nel Word
    img_ebitda_24_buf = genera_grafico_confronto_singolo(
        val_az_ebitda_24, val_set_ebitda_24, nome_azienda, "Margine EBITDA (%) - 2024"
    )
    context['grafico_ebitda_2024'] = InlineImage(doc, img_ebitda_24_buf, width=Mm(125))


    # =================================================================
    # 🏗️ COSTRUZIONE TABELLA 9 E GRAFICO 5 (TREND MARGINE EBIT)
    # =================================================================
    
    col_base_ebit = 'Margine EBIT (*) %'
    
    valori_settore_ebit = []
    valori_azienda_ebit = []
    
    # 1. Recupero Dati
    for anno in anni: # 'anni' è la lista ['2021', '2022', '2023', '2024'] già definita sopra
        colonna = f"{col_base_ebit} {anno}"
        
        # Mediana Settore
        if colonna in df_orbis.columns:
            valori_settore_ebit.append(df_orbis[colonna].median())
        else:
            valori_settore_ebit.append(0.0)
            
        # Valore Azienda
        if not df_target.empty and colonna in df_target.columns:
            valori_azienda_ebit.append(df_target.iloc[0][colonna])
        else:
            valori_azienda_ebit.append(0.0)

    # 2. Costruzione Tabella 9
    sd_tab9 = doc.new_subdoc()
    t9 = sd_tab9.add_table(rows=3, cols=5)
    t9.style = 'Table Grid'
    
    # Intestazioni anni (riutilizziamo l'array headers_t8 che avevamo già)
    for i, h in enumerate(headers_t8):
        t9.cell(0, i).text = h
        t9.cell(0, i).paragraphs[0].runs[0].bold = True
        
    # Riga Azienda
    t9.cell(1, 0).text = nome_azienda
    t9.cell(1, 0).paragraphs[0].runs[0].bold = True
    for i, val in enumerate(valori_azienda_ebit):
        t9.cell(1, i+1).text = f"{val:.2f}".replace('.', ',') if pd.notna(val) else "N.D."
        
    # Riga Settore (Mediana)
    t9.cell(2, 0).text = f"Settore {codice_nace_report} (Mediana)"
    t9.cell(2, 0).paragraphs[0].runs[0].bold = True
    for i, val in enumerate(valori_settore_ebit):
        t9.cell(2, i+1).text = f"{val:.2f}".replace('.', ',') if pd.notna(val) else "N.D."
        
    context['tabella_9_dinamica'] = sd_tab9

    # 3. Generazione Immagine e Iniezione nel Word (Riutilizziamo la funzione magica!)
    img_ebit_buf = genera_grafico_trend(
        anni, valori_azienda_ebit, valori_settore_ebit, 
        nome_azienda, "Andamento Margine EBIT (%)"
    )
    context['grafico_ebit'] = InlineImage(doc, img_ebit_buf, width=Mm(155))


    # =================================================================
    # 🏗️ COSTRUZIONE GRAFICO 6 (CONFRONTO DIRETTO EBIT 2024)
    # =================================================================
    
    # Prendiamo i dati del 2024 (l'ultimo valore della lista calcolata nel blocco precedente)
    val_az_ebit_24 = valori_azienda_ebit[-1] if len(valori_azienda_ebit) > 0 else 0
    val_set_ebit_24 = valori_settore_ebit[-1] if len(valori_settore_ebit) > 0 else 0

    # Generazione Immagine e Iniezione nel Word (riutilizziamo la funzione a barre)
    img_ebit_24_buf = genera_grafico_confronto_singolo(
        val_az_ebit_24, val_set_ebit_24, nome_azienda, "Margine EBIT (%) - 2024"
    )
    context['grafico_ebit_2024'] = InlineImage(doc, img_ebit_24_buf, width=Mm(125))


    # =================================================================
    # 🏗️ COSTRUZIONE TABELLA 10 E GRAFICO 7 (TREND MARGINE DI PROFITTO)
    # =================================================================
    
    col_base_profitto = 'Margine di Profitto (*) %' # Assicurati che nel tuo Excel si chiami così
    
    valori_settore_profitto = []
    valori_azienda_profitto = []
    
    # 1. Recupero Dati
    for anno in anni:
        colonna = f"{col_base_profitto} {anno}"
        
        # Mediana Settore
        if colonna in df_orbis.columns:
            valori_settore_profitto.append(df_orbis[colonna].median())
        else:
            valori_settore_profitto.append(0.0)
            
        # Valore Azienda
        if not df_target.empty and colonna in df_target.columns:
            valori_azienda_profitto.append(df_target.iloc[0][colonna])
        else:
            valori_azienda_profitto.append(0.0)

    # 2. Costruzione Tabella 10
    sd_tab10 = doc.new_subdoc()
    t10 = sd_tab10.add_table(rows=3, cols=5)
    t10.style = 'Table Grid'
    
    # Intestazioni anni (questa volta la prima cella ha del testo)
    headers_t10 = ['Margine di Profitto', '2021', '2022', '2023', '2024']
    for i, h in enumerate(headers_t10):
        t10.cell(0, i).text = h
        t10.cell(0, i).paragraphs[0].runs[0].bold = True
        
    # Riga Azienda
    t10.cell(1, 0).text = nome_azienda
    t10.cell(1, 0).paragraphs[0].runs[0].bold = True
    for i, val in enumerate(valori_azienda_profitto):
        t10.cell(1, i+1).text = f"{val:.2f}".replace('.', ',') if pd.notna(val) else "N.D."
        
    # Riga Settore (Mediana)
    t10.cell(2, 0).text = f"Settore {codice_nace_report} (Mediana)"
    t10.cell(2, 0).paragraphs[0].runs[0].bold = True
    for i, val in enumerate(valori_settore_profitto):
        t10.cell(2, i+1).text = f"{val:.2f}".replace('.', ',') if pd.notna(val) else "N.D."
        
    context['tabella_10_dinamica'] = sd_tab10

    # 3. Generazione Immagine e Iniezione nel Word
    img_profitto_buf = genera_grafico_trend(
        anni, valori_azienda_profitto, valori_settore_profitto, 
        nome_azienda, "Andamento Margine di Profitto (%)"
    )
    context['grafico_profitto'] = InlineImage(doc, img_profitto_buf, width=Mm(155))

    # =================================================================
    # 🏗️ COSTRUZIONE GRAFICO 8 (CONFRONTO DIRETTO PROFITTO 2024)
    # =================================================================
    
    # Prendiamo i dati del 2024 (l'ultimo valore della lista calcolata nel blocco precedente)
    val_az_profitto_24 = valori_azienda_profitto[-1] if len(valori_azienda_profitto) > 0 else 0
    val_set_profitto_24 = valori_settore_profitto[-1] if len(valori_settore_profitto) > 0 else 0

    # Generazione Immagine e Iniezione nel Word (riutilizziamo la funzione a barre)
    img_profitto_24_buf = genera_grafico_confronto_singolo(
        val_az_profitto_24, val_set_profitto_24, nome_azienda, "Margine di Profitto (%) - 2024"
    )
    context['grafico_profitto_2024'] = InlineImage(doc, img_profitto_24_buf, width=Mm(125))

    # =================================================================
    # 🏗️ COSTRUZIONE TABELLA 11 E GRAFICO 9 (TREND STRUTTURA 1° LIV)
    # =================================================================
    
    col_base_strut1 = 'Indice di Struttura 1° livello (*)' # Verifica l'esatto nome nel tuo Excel senza l'anno
    
    valori_settore_strut1 = []
    valori_azienda_strut1 = []
    
    # 1. Recupero Dati
    for anno in anni:
        colonna = f"{col_base_strut1} {anno}"
        
        # Mediana Settore
        if colonna in df_orbis.columns:
            valori_settore_strut1.append(df_orbis[colonna].median())
        else:
            valori_settore_strut1.append(0.0)
            
        # Valore Azienda
        if not df_target.empty and colonna in df_target.columns:
            valori_azienda_strut1.append(df_target.iloc[0][colonna])
        else:
            valori_azienda_strut1.append(0.0)

    # 2. Costruzione Tabella 11
    sd_tab11 = doc.new_subdoc()
    t11 = sd_tab11.add_table(rows=3, cols=5)
    t11.style = 'Table Grid'
    
    # Intestazioni anni 
    headers_t11 = ['Indice Struttura 1° Liv.', '2021', '2022', '2023', '2024']
    for i, h in enumerate(headers_t11):
        t11.cell(0, i).text = h
        t11.cell(0, i).paragraphs[0].runs[0].bold = True
        
    # Riga Azienda
    t11.cell(1, 0).text = nome_azienda
    t11.cell(1, 0).paragraphs[0].runs[0].bold = True
    for i, val in enumerate(valori_azienda_strut1):
        t11.cell(1, i+1).text = f"{val:.2f}".replace('.', ',') if pd.notna(val) else "N.D."
        
    # Riga Settore (Mediana)
    t11.cell(2, 0).text = f"Settore {codice_nace_report} (Mediana)"
    t11.cell(2, 0).paragraphs[0].runs[0].bold = True
    for i, val in enumerate(valori_settore_strut1):
        t11.cell(2, i+1).text = f"{val:.2f}".replace('.', ',') if pd.notna(val) else "N.D."
        
    context['tabella_11_dinamica'] = sd_tab11

    # 3. Generazione Immagine e Iniezione nel Word
    img_strut1_buf = genera_grafico_trend(
        anni, valori_azienda_strut1, valori_settore_strut1, 
        nome_azienda, "Andamento Indice di Struttura 1° Liv."
    )
    context['grafico_strut1'] = InlineImage(doc, img_strut1_buf, width=Mm(155))


    # =================================================================
    # 🏗️ COSTRUZIONE GRAFICO 10 (CONFRONTO DIRETTO STRUTTURA 1° LIV 2024)
    # =================================================================
    
    # Prendiamo i dati del 2024 (l'ultimo valore della lista calcolata nel blocco precedente)
    val_az_strut1_24 = valori_azienda_strut1[-1] if len(valori_azienda_strut1) > 0 else 0
    val_set_strut1_24 = valori_settore_strut1[-1] if len(valori_settore_strut1) > 0 else 0

    # Generazione Immagine e Iniezione nel Word (riutilizziamo la funzione a barre)
    img_strut1_24_buf = genera_grafico_confronto_singolo(
        val_az_strut1_24, val_set_strut1_24, nome_azienda, "Indice Struttura 1° Liv. - 2024"
    )
    context['grafico_strut1_2024'] = InlineImage(doc, img_strut1_24_buf, width=Mm(125))

    # =================================================================
    # 🏗️ COSTRUZIONE TABELLA 12 E GRAFICO 11 (TREND STRUTTURA 2° LIV)
    # =================================================================
    
    col_base_strut2 = 'Indice di Struttura 2° livello (*)' # Verifica l'esatto nome nel tuo Excel senza l'anno
    
    valori_settore_strut2 = []
    valori_azienda_strut2 = []
    
    # 1. Recupero Dati
    for anno in anni:
        colonna = f"{col_base_strut2} {anno}"
        
        # Mediana Settore
        if colonna in df_orbis.columns:
            valori_settore_strut2.append(df_orbis[colonna].median())
        else:
            valori_settore_strut2.append(0.0)
            
        # Valore Azienda
        if not df_target.empty and colonna in df_target.columns:
            valori_azienda_strut2.append(df_target.iloc[0][colonna])
        else:
            valori_azienda_strut2.append(0.0)

    # 2. Costruzione Tabella 12
    sd_tab12 = doc.new_subdoc()
    t12 = sd_tab12.add_table(rows=3, cols=5)
    t12.style = 'Table Grid'
    
    # Intestazioni anni 
    headers_t12 = ['Indice Struttura 2° Liv.', '2021', '2022', '2023', '2024']
    for i, h in enumerate(headers_t12):
        t12.cell(0, i).text = h
        t12.cell(0, i).paragraphs[0].runs[0].bold = True
        
    # Riga Azienda
    t12.cell(1, 0).text = nome_azienda
    t12.cell(1, 0).paragraphs[0].runs[0].bold = True
    for i, val in enumerate(valori_azienda_strut2):
        t12.cell(1, i+1).text = f"{val:.2f}".replace('.', ',') if pd.notna(val) else "N.D."
        
    # Riga Settore (Mediana)
    t12.cell(2, 0).text = f"Settore {codice_nace_report} (Mediana)"
    t12.cell(2, 0).paragraphs[0].runs[0].bold = True
    for i, val in enumerate(valori_settore_strut2):
        t12.cell(2, i+1).text = f"{val:.2f}".replace('.', ',') if pd.notna(val) else "N.D."
        
    context['tabella_12_dinamica'] = sd_tab12

    # 3. Generazione Immagine e Iniezione nel Word
    img_strut2_buf = genera_grafico_trend(
        anni, valori_azienda_strut2, valori_settore_strut2, 
        nome_azienda, "Andamento Indice di Struttura 2° Liv."
    )
    context['grafico_strut2'] = InlineImage(doc, img_strut2_buf, width=Mm(155))

    # =================================================================
    # 🏗️ COSTRUZIONE GRAFICO 12 (CONFRONTO DIRETTO STRUTTURA 2° LIV 2024)
    # =================================================================
    
    # Prendiamo i dati del 2024 (l'ultimo valore della lista calcolata nel blocco precedente)
    val_az_strut2_24 = valori_azienda_strut2[-1] if len(valori_azienda_strut2) > 0 else 0
    val_set_strut2_24 = valori_settore_strut2[-1] if len(valori_settore_strut2) > 0 else 0

    # Generazione Immagine e Iniezione nel Word (riutilizziamo la funzione a barre)
    img_strut2_24_buf = genera_grafico_confronto_singolo(
        val_az_strut2_24, val_set_strut2_24, nome_azienda, "Indice Struttura 2° Liv. - 2024"
    )
    context['grafico_strut2_2024'] = InlineImage(doc, img_strut2_24_buf, width=Mm(125))

    # =================================================================
    # 🏗️ COSTRUZIONE TABELLA 13 E GRAFICO 13 (TREND GEARING)
    # =================================================================
    
    col_base_gearing = 'Gearing (*) %' # Verifica l'esatto nome nel tuo Excel senza l'anno
    
    valori_settore_gearing = []
    valori_azienda_gearing = []
    
    # 1. Recupero Dati
    for anno in anni:
        colonna = f"{col_base_gearing} {anno}"
        
        # Mediana Settore
        if colonna in df_orbis.columns:
            valori_settore_gearing.append(df_orbis[colonna].median())
        else:
            valori_settore_gearing.append(0.0)
            
        # Valore Azienda
        if not df_target.empty and colonna in df_target.columns:
            valori_azienda_gearing.append(df_target.iloc[0][colonna])
        else:
            valori_azienda_gearing.append(0.0)

    # 2. Costruzione Tabella 13
    sd_tab13 = doc.new_subdoc()
    t13 = sd_tab13.add_table(rows=3, cols=5)
    t13.style = 'Table Grid'
    
    # Intestazioni anni 
    headers_t13 = ['Gearing', '2021', '2022', '2023', '2024']
    for i, h in enumerate(headers_t13):
        t13.cell(0, i).text = h
        t13.cell(0, i).paragraphs[0].runs[0].bold = True
        
    # Riga Azienda
    t13.cell(1, 0).text = nome_azienda
    t13.cell(1, 0).paragraphs[0].runs[0].bold = True
    for i, val in enumerate(valori_azienda_gearing):
        t13.cell(1, i+1).text = f"{val:.2f}".replace('.', ',') if pd.notna(val) else "N.D."
        
    # Riga Settore (Mediana)
    t13.cell(2, 0).text = f"Settore {codice_nace_report} (Mediana)"
    t13.cell(2, 0).paragraphs[0].runs[0].bold = True
    for i, val in enumerate(valori_settore_gearing):
        t13.cell(2, i+1).text = f"{val:.2f}".replace('.', ',') if pd.notna(val) else "N.D."
        
    context['tabella_13_dinamica'] = sd_tab13

    # 3. Generazione Immagine e Iniezione nel Word
    img_gearing_buf = genera_grafico_trend(
        anni, valori_azienda_gearing, valori_settore_gearing, 
        nome_azienda, "Andamento Gearing (%)"
    )
    context['grafico_gearing'] = InlineImage(doc, img_gearing_buf, width=Mm(155))

    # =================================================================
    # 🏗️ COSTRUZIONE GRAFICO 14 (CONFRONTO DIRETTO GEARING 2024)
    # =================================================================
    
    # Prendiamo i dati del 2024 (l'ultimo valore della lista calcolata nel blocco precedente)
    val_az_gearing_24 = valori_azienda_gearing[-1] if len(valori_azienda_gearing) > 0 else 0
    val_set_gearing_24 = valori_settore_gearing[-1] if len(valori_settore_gearing) > 0 else 0

    # Generazione Immagine e Iniezione nel Word (riutilizziamo la funzione a barre)
    img_gearing_24_buf = genera_grafico_confronto_singolo(
        val_az_gearing_24, val_set_gearing_24, nome_azienda, "Gearing (%) - 2024"
    )
    context['grafico_gearing_2024'] = InlineImage(doc, img_gearing_24_buf, width=Mm(125))

    # =================================================================
    # 🏗️ COSTRUZIONE TABELLA 14 E GRAFICO TREND (CURRENT RATIO)
    # =================================================================
    
    col_base_cr = 'Current Ratio (*)' # Controlla che nel tuo Excel si chiami così o aggiungi la "x" se c'è
    
    valori_settore_cr = []
    valori_azienda_cr = []
    
    # 1. Recupero Dati
    for anno in anni:
        colonna = f"{col_base_cr} {anno}"
        
        # Mediana Settore
        if colonna in df_orbis.columns:
            valori_settore_cr.append(df_orbis[colonna].median())
        else:
            valori_settore_cr.append(0.0)
            
        # Valore Azienda
        if not df_target.empty and colonna in df_target.columns:
            valori_azienda_cr.append(df_target.iloc[0][colonna])
        else:
            valori_azienda_cr.append(0.0)

    # 2. Costruzione Tabella 14
    sd_tab14 = doc.new_subdoc()
    t14 = sd_tab14.add_table(rows=3, cols=5)
    t14.style = 'Table Grid'
    
    # Intestazioni anni 
    headers_t14 = ['Current Ratio', '2021', '2022', '2023', '2024']
    for i, h in enumerate(headers_t14):
        t14.cell(0, i).text = h
        t14.cell(0, i).paragraphs[0].runs[0].bold = True
        
    # Riga Azienda
    t14.cell(1, 0).text = nome_azienda
    t14.cell(1, 0).paragraphs[0].runs[0].bold = True
    for i, val in enumerate(valori_azienda_cr):
        t14.cell(1, i+1).text = f"{val:.2f}".replace('.', ',') if pd.notna(val) else "N.D."
        
    # Riga Settore (Mediana)
    t14.cell(2, 0).text = f"Settore {codice_nace_report} (Mediana)"
    t14.cell(2, 0).paragraphs[0].runs[0].bold = True
    for i, val in enumerate(valori_settore_cr):
        t14.cell(2, i+1).text = f"{val:.2f}".replace('.', ',') if pd.notna(val) else "N.D."
        
    context['tabella_14_dinamica'] = sd_tab14

    # 3. Generazione Immagine e Iniezione nel Word
    img_cr_buf = genera_grafico_trend(
        anni, valori_azienda_cr, valori_settore_cr, 
        nome_azienda, "Andamento Current Ratio"
    )
    context['grafico_cr'] = InlineImage(doc, img_cr_buf, width=Mm(155))

    # =================================================================
    # 🏗️ COSTRUZIONE GRAFICO 15 (CONFRONTO DIRETTO CURRENT RATIO 2024)
    # =================================================================
    
    # Prendiamo i dati del 2024 (l'ultimo valore della lista calcolata nel blocco precedente)
    val_az_cr_24 = valori_azienda_cr[-1] if len(valori_azienda_cr) > 0 else 0
    val_set_cr_24 = valori_settore_cr[-1] if len(valori_settore_cr) > 0 else 0

    # Generazione Immagine e Iniezione nel Word (riutilizziamo la funzione a barre)
    img_cr_24_buf = genera_grafico_confronto_singolo(
        val_az_cr_24, val_set_cr_24, nome_azienda, "Current Ratio - 2024"
    )
    context['grafico_cr_2024'] = InlineImage(doc, img_cr_24_buf, width=Mm(125))


    # =================================================================
    # 🏗️ COSTRUZIONE TABELLA 15 E GRAFICO 16 (TREND QUICK RATIO)
    # =================================================================
    
    col_base_qr = 'Quick Ratio (*)' # Controlla che nel tuo Excel la colonna si chiami così
    
    valori_settore_qr = []
    valori_azienda_qr = []
    
    # 1. Recupero Dati
    for anno in anni:
        colonna = f"{col_base_qr} {anno}"
        
        # Mediana Settore
        if colonna in df_orbis.columns:
            valori_settore_qr.append(df_orbis[colonna].median())
        else:
            valori_settore_qr.append(0.0)
            
        # Valore Azienda
        if not df_target.empty and colonna in df_target.columns:
            valori_azienda_qr.append(df_target.iloc[0][colonna])
        else:
            valori_azienda_qr.append(0.0)

    # 2. Costruzione Tabella 15
    sd_tab15 = doc.new_subdoc()
    t15 = sd_tab15.add_table(rows=3, cols=5)
    t15.style = 'Table Grid'
    
    # Intestazioni anni 
    headers_t15 = ['Quick Ratio', '2021', '2022', '2023', '2024']
    for i, h in enumerate(headers_t15):
        t15.cell(0, i).text = h
        t15.cell(0, i).paragraphs[0].runs[0].bold = True
        
    # Riga Azienda
    t15.cell(1, 0).text = nome_azienda
    t15.cell(1, 0).paragraphs[0].runs[0].bold = True
    for i, val in enumerate(valori_azienda_qr):
        t15.cell(1, i+1).text = f"{val:.2f}".replace('.', ',') if pd.notna(val) else "N.D."
        
    # Riga Settore (Mediana)
    t15.cell(2, 0).text = f"Settore {codice_nace_report} (Mediana)"
    t15.cell(2, 0).paragraphs[0].runs[0].bold = True
    for i, val in enumerate(valori_settore_qr):
        t15.cell(2, i+1).text = f"{val:.2f}".replace('.', ',') if pd.notna(val) else "N.D."
        
    context['tabella_15_dinamica'] = sd_tab15

    # 3. Generazione Immagine e Iniezione nel Word
    img_qr_buf = genera_grafico_trend(
        anni, valori_azienda_qr, valori_settore_qr, 
        nome_azienda, "Andamento Quick Ratio"
    )
    context['grafico_qr'] = InlineImage(doc, img_qr_buf, width=Mm(155))


    # =================================================================
    # 🏗️ COSTRUZIONE GRAFICO 17 (CONFRONTO DIRETTO QUICK RATIO 2024)
    # =================================================================
    
    # Prendiamo i dati del 2024 (l'ultimo valore della lista calcolata nel blocco precedente)
    val_az_qr_24 = valori_azienda_qr[-1] if len(valori_azienda_qr) > 0 else 0
    val_set_qr_24 = valori_settore_qr[-1] if len(valori_settore_qr) > 0 else 0

    # Generazione Immagine e Iniezione nel Word (riutilizziamo la funzione a barre)
    img_qr_24_buf = genera_grafico_confronto_singolo(
        val_az_qr_24, val_set_qr_24, nome_azienda, "Quick Ratio - 2024"
    )
    context['grafico_qr_2024'] = InlineImage(doc, img_qr_24_buf, width=Mm(125))


    # =================================================================
    # 🏗️ COSTRUZIONE TABELLA 16 E GRAFICO 18 (TREND ROTAZIONE CAP. INV.)
    # =================================================================
    
    # IMPORTANTE: Controlla che il nome di questa colonna corrisponda esattamente a quello del tuo Excel Orbis!
    col_base_rotazione = 'Indice di Rotazione del Capitale Investito (*)' 
    
    valori_settore_rotazione = []
    valori_azienda_rotazione = []
    
    # 1. Recupero Dati
    for anno in anni:
        colonna = f"{col_base_rotazione} {anno}"
        
        # Mediana Settore
        if colonna in df_orbis.columns:
            valori_settore_rotazione.append(df_orbis[colonna].median())
        else:
            valori_settore_rotazione.append(0.0)
            
        # Valore Azienda
        if not df_target.empty and colonna in df_target.columns:
            valori_azienda_rotazione.append(df_target.iloc[0][colonna])
        else:
            valori_azienda_rotazione.append(0.0)

    # 2. Costruzione Tabella 16
    sd_tab16 = doc.new_subdoc()
    t16 = sd_tab16.add_table(rows=3, cols=5)
    t16.style = 'Table Grid'
    
    # Intestazioni anni 
    headers_t16 = ['Indice Rotazione Cap.Inv.', '2021', '2022', '2023', '2024']
    for i, h in enumerate(headers_t16):
        t16.cell(0, i).text = h
        t16.cell(0, i).paragraphs[0].runs[0].bold = True
        
    # Riga Azienda
    t16.cell(1, 0).text = nome_azienda
    t16.cell(1, 0).paragraphs[0].runs[0].bold = True
    for i, val in enumerate(valori_azienda_rotazione):
        t16.cell(1, i+1).text = f"{val:.2f}".replace('.', ',') if pd.notna(val) else "N.D."
        
    # Riga Settore (Mediana)
    t16.cell(2, 0).text = f"Settore {codice_nace_report} (Mediana)"
    t16.cell(2, 0).paragraphs[0].runs[0].bold = True
    for i, val in enumerate(valori_settore_rotazione):
        t16.cell(2, i+1).text = f"{val:.2f}".replace('.', ',') if pd.notna(val) else "N.D."
        
    context['tabella_16_dinamica'] = sd_tab16

    # 3. Generazione Immagine e Iniezione nel Word
    img_rotazione_buf = genera_grafico_trend(
        anni, valori_azienda_rotazione, valori_settore_rotazione, 
        nome_azienda, "Andamento Rotazione Cap. Inv."
    )
    context['grafico_rotazione'] = InlineImage(doc, img_rotazione_buf, width=Mm(155))

    # =================================================================
    # 🏗️ COSTRUZIONE GRAFICO 19 (CONFRONTO DIRETTO ROTAZIONE CAP. INV. 2024)
    # =================================================================
    
    # Prendiamo i dati del 2024 (l'ultimo valore della lista calcolata nel blocco precedente)
    val_az_rotazione_24 = valori_azienda_rotazione[-1] if len(valori_azienda_rotazione) > 0 else 0
    val_set_rotazione_24 = valori_settore_rotazione[-1] if len(valori_settore_rotazione) > 0 else 0

    # Generazione Immagine e Iniezione nel Word (riutilizziamo la funzione a barre)
    img_rotazione_24_buf = genera_grafico_confronto_singolo(
        val_az_rotazione_24, val_set_rotazione_24, nome_azienda, "Rotazione Cap. Inv. - 2024"
    )
    context['grafico_rotazione_2024'] = InlineImage(doc, img_rotazione_24_buf, width=Mm(125))

    # =================================================================
    # 🏗️ NOTA METODOLOGICA: CALCOLO E COSTRUZIONE TABELLE TERZILI 2024
    # =================================================================

    # 1. Funzione intelligente per calcolare e formattare i terzili di una colonna
    # NOTA: la numerazione dei terzili segue il "terzile di merito" (1° = migliore, 3° = peggiore),
    # coerente con il motore di rating (punteggio_diretto/punteggio_inverso più sopra).
    # Per le variabili dirette (valore più alto = performance migliore) il 1° terzile è quindi il
    # terzo SUPERIORE della distribuzione (ordine decrescente). Il Gearing è l'unica variabile
    # inversa (valore più basso = performance migliore): per coerenza il suo 1° terzile resta il
    # terzo INFERIORE della distribuzione (ordine crescente).
    def calcola_terzili(df, col_name, inverso=False):
        if col_name not in df.columns:
            return ["N.D.", "N.D.", "N.D."]

        # Prendiamo solo i dati validi (senza NaN)
        s = df[col_name].dropna()
        if s.empty:
            return ["N.D.", "N.D.", "N.D."]

        vmin = s.min()
        v33 = s.quantile(0.3333)
        v66 = s.quantile(0.6666)
        vmax = s.max()

        # Funzione per formattare i numeri in stile italiano (1.234,56)
        def fmt(val):
            return f"{val:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

        fascia_bassa = f"[{fmt(vmin)} - {fmt(v33)})"
        fascia_media = f"[{fmt(v33)} - {fmt(v66)})"
        fascia_alta = f"[{fmt(v66)} - {fmt(vmax)}]"

        if inverso:
            # Gearing: 1° terzile (Best) = valori più bassi
            return [fascia_bassa, fascia_media, fascia_alta]
        else:
            # Variabili dirette: 1° terzile (Best) = valori più alti
            return [fascia_alta, fascia_media, fascia_bassa]

    # Funzione per costruire dinamicamente le tabelle Word dei terzili
    def crea_tabella_terzili(doc, headers, col_names_orbis, colonne_inverse=None):
        colonne_inverse = colonne_inverse or []
        sd = doc.new_subdoc()
        t = sd.add_table(rows=4, cols=4)
        t.style = 'Table Grid'

        # Intestazioni (Riga 0)
        t.cell(0, 0).text = "Terzili"
        t.cell(0, 0).paragraphs[0].runs[0].bold = True
        for i, h in enumerate(headers):
            t.cell(0, i+1).text = h
            t.cell(0, i+1).paragraphs[0].runs[0].bold = True

        # Nomi delle righe
        t.cell(1, 0).text = "1° (Best)"
        t.cell(2, 0).text = "2° (Average)"
        t.cell(3, 0).text = "3° (Worst)"
        for r in (1, 2, 3):
            t.cell(r, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Popolamento dati incrociati
        for i, col_name in enumerate(col_names_orbis):
            fasce = calcola_terzili(df_orbis, col_name, inverso=(col_name in colonne_inverse))
            t.cell(1, i+1).text = fasce[0]
            t.cell(2, i+1).text = fasce[1]
            t.cell(3, i+1).text = fasce[2]

        return sd

    # --- Creazione Tabella 1 (Profitto, EBITDA, EBIT) --- tutte variabili dirette
    headers_t1 = ['Margine di profitto', 'Margine EBITDA', 'Margine EBIT']
    colonne_t1 = [f"{col_base_profitto} 2024", f"{col_base_ebitda} 2024", f"{col_base_ebit} 2024"]
    context['tabella_terzili_1'] = crea_tabella_terzili(doc, headers_t1, colonne_t1)

    # --- Creazione Tabella 2 (Struttura 1°, Struttura 2°, Gearing) --- Gearing è inversa
    headers_t2 = ['Indice Struttura 1° Livello', 'Indice Struttura 2° Livello', 'Indice Gearing']
    colonne_t2 = [f"{col_base_strut1} 2024", f"{col_base_strut2} 2024", f"{col_base_gearing} 2024"]
    context['tabella_terzili_2'] = crea_tabella_terzili(doc, headers_t2, colonne_t2, colonne_inverse=[f"{col_base_gearing} 2024"])

    # --- Creazione Tabella 3 (Rotazione, Quick, Current) --- tutte variabili dirette
    headers_t3 = ['Indice Rotazione Cap.Inv.', 'Quick Ratio', 'Current Ratio']
    colonne_t3 = [f"{col_base_rotazione} 2024", f"{col_base_qr} 2024", f"{col_base_cr} 2024"]
    context['tabella_terzili_3'] = crea_tabella_terzili(doc, headers_t3, colonne_t3)

    # =================================================================
    # 🏗️ COSTRUZIONE TABELLE BENCHMARK (TOTALMENTE DINAMICHE DA APP.PY)
    # =================================================================

    def crea_tabella_benchmark(doc, titolo_colonna, dict_valori):
        sd = doc.new_subdoc()
        t = sd.add_table(rows=5, cols=3)
        t.style = 'Table Grid'
        
        totale_va = sum(dict_valori.values())
        
        t.cell(0, 0).text = titolo_colonna
        t.cell(0, 0).paragraphs[0].runs[0].bold = True
        t.cell(0, 1).text = "V.A."
        t.cell(0, 1).paragraphs[0].runs[0].bold = True
        t.cell(0, 2).text = "%"
        t.cell(0, 2).paragraphs[0].runs[0].bold = True
        
        for i, key in enumerate(['A', 'B', 'C']):
            va = dict_valori.get(key, 0)
            perc = (va / totale_va * 100) if totale_va > 0 else 0
            t.cell(i+1, 0).text = key
            t.cell(i+1, 1).text = f"{va:,}".replace(',', '.') 
            t.cell(i+1, 2).text = f"{perc:.2f}".replace('.', ',') 
            
        t.cell(4, 0).text = "Totale"
        t.cell(4, 0).paragraphs[0].runs[0].bold = True
        t.cell(4, 1).text = f"{totale_va:,}".replace(',', '.')
        t.cell(4, 1).paragraphs[0].runs[0].bold = True
        t.cell(4, 2).text = "100,00" if totale_va > 0 else "0,00"
        t.cell(4, 2).paragraphs[0].runs[0].bold = True
        
        return sd

    # ---------------------------------------------------------
    # 📊 MOTORE DI CALCOLO PANDAS (REPLICA FORMULE EXCEL APP.PY)
    # ---------------------------------------------------------
    df_b = df_orbis.copy()

    # Funzione per assegnare i punteggi 1, 2, 3 basata sui terzili
    def calcola_punti(serie_col, higher_is_better=True):
        s = pd.to_numeric(df_b[serie_col], errors='coerce')
        t1 = s.quantile(1/3)
        t2 = s.quantile(2/3)
        
        def assegna(x):
            if pd.isna(x): return 1 # Default se mancante
            if higher_is_better:
                return 3 if x >= t2 else (2 if x >= t1 else 1)
            else:
                return 3 if x <= t1 else (2 if x <= t2 else 1)
        return s.apply(assegna)

    # 1. Benchmark Economico (Più alto è meglio)
    pt_prof = calcola_punti(f"{col_base_profitto} 2024", True)
    pt_ebitda = calcola_punti(f"{col_base_ebitda} 2024", True)
    pt_ebit = calcola_punti(f"{col_base_ebit} 2024", True)
    df_b['Score_Eco'] = pt_prof + pt_ebitda + pt_ebit
    df_b['Bench_Eco'] = df_b['Score_Eco'].apply(lambda x: 'A' if x >= 8 else ('B' if x >= 5 else 'C'))

    # 2. Benchmark Finanziario (Rotazione: Più basso è meglio - Quick/Current: Più alto è meglio)
    pt_rot = calcola_punti(f"{col_base_rotazione} 2024", True) # in app.py cond_H <= T1
    pt_qr = calcola_punti(f"{col_base_qr} 2024", True)
    pt_cr = calcola_punti(f"{col_base_cr} 2024", True)
    df_b['Score_Fin'] = pt_rot + pt_qr + pt_cr
    df_b['Bench_Fin'] = df_b['Score_Fin'].apply(lambda x: 'A' if x >= 8 else ('B' if x >= 5 else 'C'))

    # 3. Benchmark Patrimoniale (Struttura: Più alto è meglio - Gearing: Più basso è meglio)
    pt_s1 = calcola_punti(f"{col_base_strut1} 2024", True)
    pt_s2 = calcola_punti(f"{col_base_strut2} 2024", True)
    pt_gear = calcola_punti(f"{col_base_gearing} 2024", False) # in app.py cond_N <= T1
    df_b['Score_Pat'] = pt_s1 + pt_s2 + pt_gear
    df_b['Bench_Pat'] = df_b['Score_Pat'].apply(lambda x: 'A' if x >= 8 else ('B' if x >= 5 else 'C'))

    # 4. Benchmark Totale
    def val_lettera(l):
        return 3 if l == 'A' else (2 if l == 'B' else 1)
    
    df_b['Score_Tot'] = df_b['Bench_Eco'].apply(val_lettera) + df_b['Bench_Fin'].apply(val_lettera) + df_b['Bench_Pat'].apply(val_lettera)
    df_b['Bench_Tot'] = df_b['Score_Tot'].apply(lambda x: 'A' if x >= 8 else ('B' if x >= 5 else 'C'))

    # Estrazione dei conteggi esatti
    conteggi_eco = df_b['Bench_Eco'].value_counts().to_dict()
    conteggi_fin = df_b['Bench_Fin'].value_counts().to_dict()
    conteggi_pat = df_b['Bench_Pat'].value_counts().to_dict()
    conteggi_tot = df_b['Bench_Tot'].value_counts().to_dict()

    # Iniezione nel Word
    context['tab_bench_eco'] = crea_tabella_benchmark(doc, "Benchmark Economico", conteggi_eco)
    context['tab_bench_pat'] = crea_tabella_benchmark(doc, "Benchmark Patrimoniale", conteggi_pat)
    context['tab_bench_fin'] = crea_tabella_benchmark(doc, "Benchmark Finanziario", conteggi_fin)
    context['tab_bench_tot'] = crea_tabella_benchmark(doc, "Benchmark Totale", conteggi_tot)



    # =================================================================
    # 🏗️ COSTRUZIONE TABELLE STATISTICHE DESCRITTIVE (PROFITTO, EBITDA, EBIT)
    # =================================================================

    # Funzione universale per creare la tabella delle statistiche
    def crea_tabella_statistiche(doc, titolo, col_base, anni, df):
        sd = doc.new_subdoc()
        t = sd.add_table(rows=6, cols=len(anni) + 1)
        t.style = 'Table Grid'
        
        # Intestazione Colonne
        t.cell(0, 0).text = titolo
        t.cell(0, 0).paragraphs[0].runs[0].bold = True
        for i, anno in enumerate(anni):
            t.cell(0, i+1).text = str(anno)
            t.cell(0, i+1).paragraphs[0].runs[0].bold = True
            
        # Intestazione Righe
        labels = ["Media", "Mediana", "Asimmetria", "Curtosi", "Deviazione standard"]
        for r, label in enumerate(labels):
            t.cell(r+1, 0).text = label
            t.cell(r+1, 0).paragraphs[0].runs[0].bold = True
            
        # Funzione di formattazione italiana (1.234,56)
        def fmt(val):
            if pd.isna(val): return "N.D."
            return f"{val:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

        # Calcolo dinamico per ogni anno
        for i, anno in enumerate(anni):
            col_name = f"{col_base} {anno}"
            if col_name in df.columns:
                # Prendiamo solo i numeri puliti (scartando i null/NaN)
                s = pd.to_numeric(df[col_name], errors='coerce').dropna()
                
                if not s.empty:
                    media = s.mean()
                    mediana = s.median()
                    asimmetria = s.skew()
                    curtosi = s.kurt()
                    dev_std = s.std()
                else:
                    media = mediana = asimmetria = curtosi = dev_std = float('nan')
            else:
                media = mediana = asimmetria = curtosi = dev_std = float('nan')
                
            # Scrittura nella griglia Word
            t.cell(1, i+1).text = fmt(media)
            t.cell(2, i+1).text = fmt(mediana)
            t.cell(3, i+1).text = fmt(asimmetria)
            t.cell(4, i+1).text = fmt(curtosi)
            t.cell(5, i+1).text = fmt(dev_std)
            
        return sd

    # ---------------------------------------------------------
    # Iniezione nel Word
    # Utilizziamo le variabili col_base_profitto, col_base_ebitda, col_base_ebit
    # definite in precedenza nei blocchi dei singoli indicatori
    # ---------------------------------------------------------
    
    context['tab_stat_profitto'] = crea_tabella_statistiche(doc, "Margine di Profitto", col_base_profitto, anni, df_orbis)
    context['tab_stat_ebitda'] = crea_tabella_statistiche(doc, "Margine EBITDA", col_base_ebitda, anni, df_orbis)
    context['tab_stat_ebit'] = crea_tabella_statistiche(doc, "Margine EBIT", col_base_ebit, anni, df_orbis)

    context['tab_stat_strut1'] = crea_tabella_statistiche(doc, "Indice Struttura 1° Liv.", col_base_strut1, anni, df_orbis)
    context['tab_stat_strut2'] = crea_tabella_statistiche(doc, "Indice Struttura 2° Liv.", col_base_strut2, anni, df_orbis)
    context['tab_stat_gearing'] = crea_tabella_statistiche(doc, "Gearing (%)", col_base_gearing, anni, df_orbis)

    context['tab_stat_cr'] = crea_tabella_statistiche(doc, "Current Ratio", col_base_cr, anni, df_orbis)
    context['tab_stat_qr'] = crea_tabella_statistiche(doc, "Quick Ratio", col_base_qr, anni, df_orbis)
    context['tab_stat_rotazione'] = crea_tabella_statistiche(doc, "Rotazione Cap. Inv.", col_base_rotazione, anni, df_orbis)

    
    # =================================================================
    # 🏗️ COSTRUZIONE TABELLE MEDIANE E VARIAZIONI YoY (Tab. 16 - 24)
    # =================================================================

    def crea_tabella_yoy(doc, titolo, col_base, anni, df):
        sd = doc.new_subdoc()
        # Calcoliamo il numero di colonne: Nome + Primo Anno + 2 colonne (V.m. e Delta) per ogni anno successivo
        num_cols = 2 + (len(anni) - 1) * 2 
        t = sd.add_table(rows=2, cols=num_cols)
        t.style = 'Table Grid'
        
        # --- Costruzione Intestazioni ---
        t.cell(0, 0).text = "Metrica"
        t.cell(0, 1).text = f"{anni[0]} V.m."
        col_idx = 2
        for anno in anni[1:]:
            t.cell(0, col_idx).text = f"{anno} V.m."
            t.cell(0, col_idx+1).text = "Δ %"
            col_idx += 2
            
        for i in range(num_cols):
            t.cell(0, i).paragraphs[0].runs[0].bold = True
            
        t.cell(1, 0).text = titolo
        t.cell(1, 0).paragraphs[0].runs[0].bold = True
        
        # Funzioni di formattazione
        def fmt_val(val):
            return f"{val:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') if pd.notna(val) else "N.D."
            
        prev_val = None
        
        # --- Inserimento Dati 1° Anno (es. 2021) ---
        col_name_0 = f"{col_base} {anni[0]}"
        if col_name_0 in df.columns:
            val_0 = df[col_name_0].median()
        else:
            val_0 = float('nan')
            
        t.cell(1, 1).text = fmt_val(val_0)
        prev_val = val_0
        
        # --- Inserimento Anni Successivi e Calcolo Delta % ---
        col_idx = 2
        for anno in anni[1:]:
            col_name = f"{col_base} {anno}"
            if col_name in df.columns:
                val = df[col_name].median()
            else:
                val = float('nan')
                
            t.cell(1, col_idx).text = fmt_val(val)
            
            # Calcolo formula Δ %
            if prev_val is not None and pd.notna(prev_val) and prev_val != 0 and pd.notna(val):
                delta = ((val - prev_val) / abs(prev_val)) * 100
                t.cell(1, col_idx+1).text = fmt_val(delta) # Usiamo la stessa formattazione
            else:
                t.cell(1, col_idx+1).text = "n.d."
                
            prev_val = val
            col_idx += 2
            
        return sd

    # ---------------------------------------------------------
    # Iniezione delle 9 tabelle finali nel Word
    # ---------------------------------------------------------
    
    # Equilibrio Economico
    context['tab_yoy_profitto'] = crea_tabella_yoy(doc, "Margine di Profitto", col_base_profitto, anni, df_orbis)
    context['tab_yoy_ebitda'] = crea_tabella_yoy(doc, "Margine EBITDA", col_base_ebitda, anni, df_orbis)
    context['tab_yoy_ebit'] = crea_tabella_yoy(doc, "Margine EBIT", col_base_ebit, anni, df_orbis)
    
    # Equilibrio Patrimoniale
    context['tab_yoy_strut1'] = crea_tabella_yoy(doc, "Indice Strut. 1° Liv.", col_base_strut1, anni, df_orbis)
    context['tab_yoy_strut2'] = crea_tabella_yoy(doc, "Indice Strut. 2° Liv.", col_base_strut2, anni, df_orbis)
    context['tab_yoy_gearing'] = crea_tabella_yoy(doc, "Gearing (%)", col_base_gearing, anni, df_orbis)
    
    # Equilibrio Finanziario
    context['tab_yoy_cr'] = crea_tabella_yoy(doc, "Current Ratio", col_base_cr, anni, df_orbis)
    context['tab_yoy_qr'] = crea_tabella_yoy(doc, "Quick Ratio", col_base_qr, anni, df_orbis)
    context['tab_yoy_rotazione'] = crea_tabella_yoy(doc, "Rotazione Cap. Inv.", col_base_rotazione, anni, df_orbis)

    #---

    # =================================================================
    # 🌟 VARIABILI SOTTO I GRAFICI (Ora calcolate nel punto giusto!)
    # =================================================================
    context['ind_gear'] = context.get('gearing', 'N.D.')
    context['med_mg_ebitda'] = format_euro(val_set_ebitda_24)
    context['med_mg_ebit'] = format_euro(val_set_ebit_24)
    context['med_mg_prof'] = format_euro(val_set_profitto_24)
    context['med_ind_str1'] = format_euro(val_set_strut1_24)
    context['med_ind_str2'] = format_euro(val_set_strut2_24)
    context['med_ind_gear'] = format_euro(val_set_gearing_24)
    context['med_ind_cr'] = format_euro(val_set_cr_24)
    context['med_ind_qr'] = format_euro(val_set_qr_24)
    context['med_ind_rot_cap'] = format_euro(val_set_rotazione_24)

    # --- INDENTAZIONE AUTOMATICA DEI TESTI DISCORSIVI ---
    chiavi_narrative = [
        'descr_rating_tot', 'descr_rating_eco', 'descr_rating_patr', 'descr_rating_fin', 'descr_sintesi',
        'intro_benchmark_eco', 'intro_benchmark_patr', 'intro_benchmark_fin', 'intro_margini', 'analisi_combinata',
        'intro_divario_strutturale', 'sintesi_quadriennio_patr', 'sintesi_modello_operativo_fin',
        'analisi_margini_operativi', 'analisi_margine_profitto', 'analisi_indici_struttura', 'analisi_gearing',
        'analisi_rotazione', 'analisi_current_ratio', 'analisi_quick_ratio', 'analisi_posizionamento_fin',
        'sintesi_profilo_integrato', 'sintesi_posizionamento_lungo_periodo', 'conclusione_patrimoniale',
        'conclusione_economica', 'conclusione_finanziaria_dettaglio', 'raccomandazione_finale', 'impatto_territoriale'
    ]

    # =================================================================
    # 📏 FIX SPAZIATURA E INDENTAZIONE NARRATIVA
    # =================================================================
    from docx.shared import Cm

    # 1. Spazzoliamo tutti i paragrafi del documento (ignorando l'interno delle tabelle)
    for p in doc.docx.paragraphs:
        if p.text.strip(): 
            # Interlinea 1.5
            p.paragraph_format.line_spacing = 1.0
            
            # Controlliamo se è un paragrafo che contiene le nostre variabili narrative
            # (Se il paragrafo inizia con uno dei tag che abbiamo inserito)
            for k in chiavi_narrative:
                if "{{" + k + "}}" in p.text:
                    # Indenta l'intero blocco di 1 cm a sinistra
                    p.paragraph_format.left_indent = Cm(1.0)
                    break

    # =================================================================
    # 🚑 FIX TABELLE FINALE E CENSURA TEASER
    # =================================================================
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.shared import RGBColor

    for chiave, elemento in context.items():
        if hasattr(elemento, 'add_paragraph') or type(elemento).__name__ == 'Subdoc':
            if hasattr(elemento, 'tables'):
                for t in elemento.tables:
                    # ---------------------------------------------------------
                    # 🔒 1. LOGICA DI CENSURA TABELLE
                    # ---------------------------------------------------------
                    if modalita_teaser and len(t.rows) > 0:
                        intestazioni = [cell.text.strip() for cell in t.rows[0].cells]
                        
                        colonne_da_censurare = []
                        for c_idx, testo in enumerate(intestazioni):
                            if '2024' in testo:
                                colonne_da_censurare.append(c_idx)
                                if c_idx + 1 < len(intestazioni) and 'Δ' in intestazioni[c_idx + 1]:
                                    colonne_da_censurare.append(c_idx + 1)
                        
                        if colonne_da_censurare:
                            for r_idx in range(1, len(t.rows)):
                                for c_idx in colonne_da_censurare:
                                    if c_idx < len(t.rows[r_idx].cells):
                                        t.rows[r_idx].cells[c_idx].text = " 🔒 PREMIUM "
                                        if len(t.rows[r_idx].cells[c_idx].paragraphs) > 0:
                                            run = t.rows[r_idx].cells[c_idx].paragraphs[0].runs[0]
                                            run.font.bold = True
                                            run.font.color.rgb = RGBColor(255, 0, 0)
                        
                        elif len(intestazioni) > 0 and intestazioni[0] == 'Terzili':
                            for riga in t.rows[1:]:
                                for c in range(1, len(riga.cells)):
                                    riga.cells[c].text = " 🔒 PREMIUM "
                                    if len(riga.cells[c].paragraphs) > 0:
                                        run = riga.cells[c].paragraphs[0].runs[0]
                                        run.font.bold = True
                                        run.font.color.rgb = RGBColor(255, 0, 0)

                        elif len(t.rows) > 2 and len(t.rows[2].cells) > 0 and 'Ranking' in t.rows[2].cells[0].text:
                            for r_idx in range(2, len(t.rows)):
                                for c_idx in range(1, len(t.rows[r_idx].cells)):
                                    t.rows[r_idx].cells[c_idx].text = " 🔒 PREMIUM "
                                    if len(t.rows[r_idx].cells[c_idx].paragraphs) > 0:
                                        run = t.rows[r_idx].cells[c_idx].paragraphs[0].runs[0]
                                        run.font.bold = True
                                        run.font.color.rgb = RGBColor(255, 0, 0)

                    # ---------------------------------------------------------
                    # 🎨 2. LOGICA DI LAYOUT E COLORI ORBIS
                    # ---------------------------------------------------------
                    try:
                        t.autofit = True
                        tbl_pr = t._tbl.tblPr
                        tbl_w = tbl_pr.xpath("./w:tblW")
                        if tbl_w:
                            tbl_w[0].set(qn('w:type'), 'auto')
                            tbl_w[0].set(qn('w:w'), '0')
                        t.alignment = WD_TABLE_ALIGNMENT.CENTER
                        
                        if len(t.rows) > 0:
                            for r_idx, row in enumerate(t.rows):
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    if r_idx == 0:
                                        shd = parse_xml(r'<w:shd %s w:val="clear" w:color="auto" w:fill="2B3A67"/>' % nsdecls('w'))
                                        cell._tc.get_or_add_tcPr().append(shd)
                                        for paragraph in cell.paragraphs:
                                            for run in paragraph.runs:
                                                run.font.color.rgb = RGBColor(255, 255, 255)
                    except:
                        pass

    # =================================================================
    # 🔒 3. CENSURA VARIABILI (Testo Puro = Zero Crash)
    # =================================================================
    if modalita_teaser:
        testo_censura_sicuro = " 🔒 PREMIUM "
        
        variabili_sensibili = [
            'mg_ebitda', 'mg_ebit', 'mg_prof',
            'ind_str1', 'ind_str2', 'gearing', 'ind_gear',
            'ind_rot_cap', 'ind_cr', 'ind_qr',
            'ricavi_mln', 'attivo_mln',
            'perc_ricavi_panel', 'perc_attivo_panel', 'perc_dip_area',
            'perc_imprese_regione', 'perc_ricavi_macroregione', 'tot_ricavi_macro_mln',
            'perc_ricavi_target_su_macro', 'perc_ricavi_categoria', 'perc_attivo_categoria',
            'rnk_naz_ebitda', 'rnk_reg_ebitda', 'rnk_naz_ebit', 'rnk_reg_ebit', 'rnk_naz_prof', 'rnk_reg_prof',
            'rnk_naz_strut1', 'rnk_reg_strut1', 'rnk_naz_strut2', 'rnk_reg_strut2', 'rnk_naz_gear', 'rnk_reg_gear',
            'rnk_naz_rot', 'rnk_reg_rot', 'rnk_naz_cr', 'rnk_reg_cr', 'rnk_naz_qr', 'rnk_reg_qr',
            'med_mg_ebitda', 'med_mg_ebit', 'med_mg_prof', 'med_ind_str1', 'med_ind_str2',
            'med_ind_gear', 'med_ind_cr', 'med_ind_qr', 'med_ind_rot_cap'
        ]
        
        valori_da_sostituire = []
        for k in variabili_sensibili:
            if k in context:
                valore_reale = str(context[k])
                context[k] = testo_censura_sicuro
                if valore_reale not in ["N.D.", "n.d.", "-", "", "0", "0,00", "0.0", "0.00"]:
                    valori_da_sostituire.append(valore_reale)
                
        # NOTA: la censura testuale non si limita più a un elenco fisso di 2 chiavi
        # ('analisi_trend_ebitda', 'impatto_territoriale'), che lasciava scoperte
        # decine di altre frasi narrative (bullet di benchmark, trend storici) che
        # riportano gli stessi valori sensibili in chiaro. Ora la sostituzione per
        # sottostringa viene applicata a QUALSIASI valore di testo residuo nel
        # context (le chiavi già svuotate sopra vengono escluse perché non serve
        # ripassarle), così da coprire anche eventuali nuove funzioni narrative
        # aggiunte in futuro senza dover aggiornare manualmente un elenco.
        chiavi_discorsive = [k for k in context if k not in variabili_sensibili and isinstance(context[k], str)]
        for k in chiavi_discorsive:
            testo_normale = context[k]
            for val in valori_da_sostituire:
                if len(val) >= 2:
                    testo_normale = testo_normale.replace(val, testo_censura_sicuro)
            context[k] = testo_normale

    # Ora genera il documento in totale sicurezza! (Renderizza il testo crudo)
    doc.render(context)

    # =================================================================
    # 🎨 4. COLORATORE NATIVO CHIRURGICO: Rende ROSSA SOLO la scritta premium
    # =================================================================
    if modalita_teaser:
        from docx.shared import RGBColor

        def colora_paragrafo_chirurgico(p):
            # Interveniamo solo se nel testo globale del paragrafo esiste la scure
            if "🔒 PREMIUM" in p.text:
                testo_completo = p.text
                
                # Ci salviamo lo stile del font del primo run per non perdere la formattazione del template
                font_name = p.runs[0].font.name if p.runs else None
                font_size = p.runs[0].font.size if p.runs else None
                
                # Svuotiamo i runs del paragrafo per ricostruirli in modo pulito
                p.text = "" 
                
                # Tagliamo il testo usando la scritta premium come separatore
                parti = re.split(r'(🔒 PREMIUM)', testo_completo)
                for parte in parti:
                    if not parte:
                        continue
                    
                    # Creiamo un run dedicato per questo frammento di testo
                    run = p.add_run(parte)
                    
                    # Ripristiniamo il font originale per non alterare il layout
                    if font_name: run.font.name = font_name
                    if font_size: run.font.size = font_size
                    
                    # Se questo specifico frammento è la scritta di censura, lo spariamo rosso e bold
                    if "🔒 PREMIUM" in parte:
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        run.font.bold = True

        # 1. Spazzoliamo tutti i paragrafi standard del documento (testi liberi, box, ecc.)
        for p in doc.docx.paragraphs:
            colora_paragrafo_chirurgico(p)

        # 2. Spazzoliamo tutti i paragrafi nascosti dentro le celle delle tabelle
        for table in doc.docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        colora_paragrafo_chirurgico(p)

    # =================================================================
    # 📏 FIX SPAZIATURA: Imposta interlinea 1.0 per tutto il testo narrativo
    # =================================================================
    for p in doc.docx.paragraphs:
        if p.text.strip():
            p.paragraph_format.line_spacing = 1.0

    output_word = io.BytesIO()
    doc.save(output_word)
    output_word.seek(0)

    # 🧹 Pulizia del file temporaneo
    try:
        os.remove(template_pulito)
    except:
        pass

    # =================================================================
    # 🔧 POST-PROCESSOR: Unisce i paragrafi frammentati dal rendering
    # =================================================================
    output_word = unisci_paragrafi_frammentati(output_word, ragione_sociale_pulita)

    # =================================================================
    # 🎨 POST-PROCESSOR: Migliora layout, indentazione e struttura
    # =================================================================
    output_word = migliora_layout(output_word)

    # =================================================================
    # 📐 POST-PROCESSOR: Corregge i riquadri fluttuanti degli indicatori
    # (Equilibrio Economico/Patrimoniale/Finanziario) che si accavallano
    # =================================================================
    output_word = correggi_riquadri_indicatori(output_word)

    # =================================================================
    # 📐 POST-PROCESSOR: Allinea i 3 box KPI dell'Executive Summary
    # (Totale Valore Produzione / Totale Attivo / Forza Lavoro)
    # =================================================================
    output_word = correggi_box_kpi_executive_summary(output_word)

    return output_word


def correggi_box_kpi_executive_summary(output_buffer):
    """
    Allinea i 3 box KPI dell'Executive Summary (Totale Valore Produzione,
    Totale Attivo, Forza Lavoro), ciascuno una tabella a una cella separata
    sovrapposta a un'unica immagine di sfondo fissa.

    Nel template il rientro (w:ind) del testo dentro le 3 celle era stato
    modificato a mano in modo scalare (641 -> 1379 -> 2099 twips), quindi il
    testo appariva sempre piu' spostato a destra passando da un box al
    successivo. Il box "Forza Lavoro", inoltre, teneva titolo e valore nello
    stesso paragrafo separati da un a-capo manuale (w:br) invece che in due
    paragrafi distinti come gli altri due: lo spazio sotto al titolo dipende
    percio' dall'interlinea invece che dalla spaziatura del paragrafo, e
    risultava diverso dagli altri due box.

    La funzione:
    - separa titolo e valore del box "Forza Lavoro" in due paragrafi come
      negli altri due box, cosi' la spaziatura sotto al titolo e' generata
      allo stesso modo in tutti e 3;
    - riporta il rientro del paragrafo-titolo (solo sinistro) e dei paragrafi
      di valore/descrizione (sinistro + destro, per contenere la larghezza
      del testo dentro il riquadro) alla stessa misura in tutti e 3 i box,
      usando come riferimento i valori del primo box (il meno alterato).
    """
    INDENT_LEFT = 641          # rientro sinistro del testo, uguale per tutti e 3 i box
    INDENT_RIGHT_VALORE = 4239  # rientro destro delle righe di valore/descrizione, per non sforare il riquadro

    def _separa_titolo_valore(paragrafo):
        """Se il paragrafo contiene un a-capo manuale (w:br), sposta tutto cio' che
        viene dopo in un nuovo paragrafo, cosi' titolo e valore hanno ciascuno la
        propria spaziatura di paragrafo invece di dipendere dall'interlinea."""
        p_el = paragrafo._p
        run_break = None
        for r in p_el.findall(qn('w:r')):
            if r.findall(qn('w:br')):
                run_break = r
                break
        if run_break is None:
            return

        parent = p_el.getparent()
        idx = list(parent).index(p_el)

        nuovo_p = OxmlElement('w:p')
        pPr_orig = p_el.find(qn('w:pPr'))
        if pPr_orig is not None:
            nuovo_p.append(copy.deepcopy(pPr_orig))

        tutti_i_run = p_el.findall(qn('w:r'))
        pos_break = tutti_i_run.index(run_break)
        for r in tutti_i_run[pos_break:]:
            p_el.remove(r)
            if r is not run_break:  # il run col solo <w:br> non serve piu': il confine ora e' il paragrafo
                nuovo_p.append(r)

        parent.insert(idx + 1, nuovo_p)

    output_buffer.seek(0)
    doc = docx.Document(output_buffer)

    prefissi = ('Totale Valore Produzione', 'Totale Attivo', 'Forza Lavoro')
    for t in doc.tables:
        primo_testo = t.rows[0].cells[0].text.strip()
        if not any(primo_testo.startswith(pref) for pref in prefissi):
            continue

        cell = t.rows[0].cells[0]
        _separa_titolo_valore(cell.paragraphs[0])

        for i, p in enumerate(cell.paragraphs):
            pPr = p._p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p._p.insert(0, pPr)
            ind = pPr.find(qn('w:ind'))
            if ind is None:
                ind = OxmlElement('w:ind')
                pPr.append(ind)
            ind.set(qn('w:left'), str(INDENT_LEFT))
            if i == 0:
                # titolo: solo rientro sinistro, larghezza piena per il testo del titolo
                if ind.get(qn('w:right')) is not None:
                    del ind.attrib[qn('w:right')]
            else:
                ind.set(qn('w:right'), str(INDENT_RIGHT_VALORE))

    result = io.BytesIO()
    doc.save(result)
    result.seek(0)
    return result


def correggi_riquadri_indicatori(output_buffer):
    """
    Corregge la spaziatura verticale dei gruppi di 3 riquadri fluttuanti
    (etichetta + descrizione) usati per gli indicatori di Equilibrio
    Economico/Patrimoniale/Finanziario. Nel template questi 3 riquadri sono
    impilati con appena ~6-7pt di margine l'uno dall'altro: il testo statico
    che contengono impagina anche di poco oltre l'altezza dichiarata (cy) e
    si accavallano fra loro. Il paragrafo che li ancora, inoltre, non
    riserva la propria altezza come farebbe un paragrafo normale: se cade a
    ridosso del fondo pagina il gruppo puo' "spaccarsi" fra due pagine invece
    di scendere in blocco alla pagina successiva.

    I gruppi si riconoscono genericamente (non per indice fisso) cercando
    paragrafi con 3+ forme ancorate "relativeFrom=paragraph" della stessa
    larghezza (le card indicatore condividono tutte cx ~6306820 EMU): questo
    esclude altri gruppi da 3 forme presenti nel documento (es. il diagramma
    dei terzili A/B/C) che hanno larghezze diverse fra loro e non c'entrano.

    Per ogni gruppo trovato:
    - allarga il gap fra i riquadri e aggiunge un margine di sicurezza alla
      loro altezza dichiarata, cosi' non si toccano piu';
    - se il paragrafo che li ancora porta anche testo reale (es. l'intro
      dell'Equilibrio Finanziario, che varia in lunghezza per azienda), le
      forme vengono spostate su un paragrafo dedicato SEMPRE vuoto inserito
      subito dopo: un paragrafo non fa scorrere le proprie righe attorno a
      una forma ancorata a se stesso, quindi se il testo dinamico va su piu'
      righe finisce sotto/dentro le forme invece di scorrervi attorno — le
      forme vanno ancorate a un paragrafo che non ha mai testo proprio;
    - riserva l'ingombro reale del gruppo con un'altezza di riga esatta sul
      paragrafo-ancora (ora sempre vuoto). Un paragrafo ad altezza esatta e'
      un blocco atomico per Word: se non c'e' spazio a sufficienza in fondo
      pagina scende TUTTO alla pagina successiva invece di spaccarsi a meta';
    - riduce lo spazio "before" riservato a mano sul paragrafo successivo
      quando e' spropositato (calibrato sull'altezza originale dei riquadri,
      ora ridondante dato che ci pensa la riserva appena creata) cosi' non
      resta piu' pagina con troppo spazio vuoto in eccesso.
    """
    CX_MIN, CX_MAX = 6000000, 6600000  # larghezza tipica delle card indicatore (~6.9in)
    GAP_EMU = 16 * 12700        # gap fra un riquadro e il successivo dello stesso gruppo
    MARGINE_EMU = 14 * 12700    # margine di sicurezza aggiunto all'altezza dichiarata di ciascun riquadro
    SOGLIA_BEFORE_HACK = 1000   # ventesimi di punto (=50pt): oltre questa soglia il "before" del
                                 # paragrafo successivo e' quasi certamente il vecchio spaziatore
                                 # calibrato a mano sull'altezza dei riquadri (ora ridondante)

    def _run_antenato(el):
        """Risale dalla forma ancorata al suo <w:r> contenitore, per poterlo spostare intero."""
        nodo = el
        while nodo is not None:
            if nodo.tag == qn('w:r'):
                return nodo
            nodo = nodo.getparent()
        return None

    output_buffer.seek(0)
    doc = docx.Document(output_buffer)

    i = 0
    while i < len(doc.paragraphs):
        paragrafi = doc.paragraphs
        p = paragrafi[i]

        gruppo = []
        for drawing in p._p.findall('.//' + qn('w:drawing')):
            anchor = drawing.find(qn('wp:anchor'))
            if anchor is None:
                continue
            posv = anchor.find(qn('wp:positionV'))
            extent = anchor.find(qn('wp:extent'))
            if posv is None or extent is None or posv.get('relativeFrom') != 'paragraph':
                continue
            cx = int(extent.get('cx'))
            if not (CX_MIN <= cx <= CX_MAX):
                continue
            offset_el = posv.find(qn('wp:posOffset'))
            if offset_el is None:
                continue
            gruppo.append({
                'offset_el': offset_el, 'extent_el': extent,
                'run_el': _run_antenato(drawing),
                'top': int(offset_el.text), 'cy': int(extent.get('cy')),
            })

        if len(gruppo) < 3:
            i += 1
            continue  # ci interessano solo i gruppi da 3 card indicatore (Economico/Patrimoniale/Finanziario)

        gruppo.sort(key=lambda g: g['top'])

        # Riimpila i riquadri con un gap piu' ampio e un margine di sicurezza sull'altezza
        running_top = gruppo[0]['top']
        for box in gruppo:
            box['new_top'] = running_top
            box['new_cy'] = box['cy'] + MARGINE_EMU
            running_top = box['new_top'] + box['new_cy'] + GAP_EMU
            box['offset_el'].text = str(box['new_top'])
            box['extent_el'].set('cy', str(box['new_cy']))

        # Ingombro reale del gruppo, da riservare come blocco atomico
        fondo_gruppo_emu = gruppo[-1]['new_top'] + gruppo[-1]['new_cy']
        altezza_riserva_ventesimi = int((fondo_gruppo_emu + 6 * 12700) / 635)

        if p.text.strip():
            # Il paragrafo-ancora porta anche testo reale (es. l'introduzione
            # dell'Equilibrio Finanziario, che varia in lunghezza per azienda): un
            # paragrafo non fa scorrere le proprie righe attorno a una forma ancorata
            # a se stesso, quindi se il testo va su piu' righe finisce sotto/dentro le
            # forme. Le spostiamo su un nuovo paragrafo dedicato, sempre vuoto, subito
            # dopo: i paragrafi successivi vi scorreranno attorno correttamente.
            nuovo_p = OxmlElement('w:p')
            pPr_new = OxmlElement('w:pPr')
            spacing_new = OxmlElement('w:spacing')
            spacing_new.set(qn('w:before'), '0')
            spacing_new.set(qn('w:after'), '0')
            spacing_new.set(qn('w:line'), str(altezza_riserva_ventesimi))
            spacing_new.set(qn('w:lineRule'), 'exact')
            pPr_new.append(spacing_new)
            nuovo_p.append(pPr_new)
            p._p.addnext(nuovo_p)
            for box in gruppo:
                if box['run_el'] is not None:
                    nuovo_p.append(box['run_el'])  # lxml sposta il nodo, non lo duplica
            idx_successivo = i + 2  # dopo il nuovo paragrafo-ancora
        else:
            pPr = p._p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p._p.insert(0, pPr)
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                pPr.append(spacing)
            spacing.set(qn('w:line'), str(altezza_riserva_ventesimi))
            spacing.set(qn('w:lineRule'), 'exact')
            idx_successivo = i + 1

        # Il paragrafo successivo aveva talvolta un "before" enorme, calibrato a mano
        # sull'altezza originale dei riquadri per non scriverci sopra: ora che la
        # riserva sopra e' corretta, quel valore e' ridondante e produce solo troppo
        # spazio vuoto in eccesso -> lo si riporta a una spaziatura normale.
        paragrafi = doc.paragraphs
        if idx_successivo < len(paragrafi):
            succ_pPr = paragrafi[idx_successivo]._p.find(qn('w:pPr'))
            if succ_pPr is not None:
                succ_spacing = succ_pPr.find(qn('w:spacing'))
                if succ_spacing is not None and succ_spacing.get(qn('w:before')):
                    if int(succ_spacing.get(qn('w:before'))) > SOGLIA_BEFORE_HACK:
                        succ_spacing.set(qn('w:before'), '120')

        i += 1

    result = io.BytesIO()
    doc.save(result)
    result.seek(0)
    return result


def unisci_paragrafi_frammentati(output_buffer, ragione_sociale):
    """
    Corregge la frammentazione dei paragrafi causata da variabili inline
    nel template docxtpl. Versione conservativa: non esegue mai merge
    tra stili di paragrafo incompatibili per evitare fusioni errate.
    """
    abbrevs_no_sentence = [
        'S.r.l.', 'S.p.A.', 'S.n.c.', 'S.a.s.', 'S.r.l.s.',
        'n.d.', 'ecc.', 'etc.', 'ca.', 'es.', '1\u00b0.', '2\u00b0.', 'Liv.'
    ]
    placeholder_vals = {'N.D.', 'n.d.', 'N/D', 'n/d', 'N.d.'}
    header_styles = {'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4'}

    def get_text(p):
        return p.text.strip()

    def is_section_header(p):
        style = p.style.name
        text = p.text.strip()
        if style in header_styles:
            return True
        if style == 'Normal' and len(text) < 60 and not any(c in text for c in ['.', ',', ';', ':']):
            return True
        if style == 'Body Text' and len(text) < 50 and not any(c in text for c in ['.', ',', ';', ':', '(', '"', '\'']):
            return True
        # Protect known section titles regardless of style
        known_titles = ['Benchmark Economico', 'Benchmark Patrimoniale', 'Benchmark Finanziario']
        if any(text == title for title in known_titles):
            return True
        return False

    def ends_sentence(text):
        if not text or text[-1] not in ['.', '!', '?']:
            return False
        for a in abbrevs_no_sentence:
            if text.endswith(a):
                return False
        words = text.split()
        if words and re.match(r'^[A-Za-z]{1,3}\.$', words[-1]):
            return False
        return True

    def copy_text_to(target_p, text, prefix=' '):
        r = OxmlElement('w:r')
        if target_p.runs:
            rpr = target_p.runs[0]._r.find(qn('w:rPr'))
            if rpr is not None:
                r.append(copy.deepcopy(rpr))
        t = OxmlElement('w:t')
        t.text = prefix + text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        target_p._p.append(r)

    def remove_p(p):
        """
        Rimuove il paragrafo se non ha forme ancorate.
        Se ha forme ancorate, lo svuota mantenendo la struttura per preservare il layout.
        """
        from docx.oxml.ns import qn as _qn
        wp_ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

        # Controlla se ha forme ancorate
        has_anchor = any(
            d.findall('.//' + ('{%s}anchor' % wp_ns))
            for d in p._p.findall('.//' + _qn('w:drawing'))
        )

        if has_anchor:
            # Ha forme ancorate: svuota senza rimuovere
            for child in list(p._p):
                if child.tag not in [_qn('w:pPr')]:
                    p._p.remove(child)
            pPr = p._p.find(_qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p._p.insert(0, pPr)
            old_spacing = pPr.find(_qn('w:spacing'))
            if old_spacing is not None:
                pPr.remove(old_spacing)
            spacing = OxmlElement('w:spacing')
            spacing.set(_qn('w:before'), '0')
            spacing.set(_qn('w:after'), '0')
            spacing.set(_qn('w:line'), '20')
            spacing.set(_qn('w:lineRule'), 'exact')
            pPr.append(spacing)
            old_rPr = pPr.find(_qn('w:rPr'))
            if old_rPr is not None:
                pPr.remove(old_rPr)
            rPr = OxmlElement('w:rPr')
            sz = OxmlElement('w:sz')
            sz.set(_qn('w:val'), '2')
            szCs = OxmlElement('w:szCs')
            szCs.set(_qn('w:val'), '2')
            rPr.append(sz)
            rPr.append(szCs)
            pPr.append(rPr)
        else:
            # Nessuna forma ancorata: rimuovi completamente
            p._p.getparent().remove(p._p)

    def has_anchor_shape(p):
        """True se il paragrafo ha forme flottanti ancorate — NON rimuovere."""
        wp_ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
        for drawing in p._p.findall('.//' + qn('w:drawing')):
            if drawing.findall('.//' + ('{%s}anchor' % wp_ns)):
                return True
        return False

    def is_numero(text):
        return bool(re.match(r'^-?\d+[.,]?\d*$', text)) and len(text) < 12

    def is_fragment(text):
        if not text:
            return False
        if text[0] in ['.', ',', '%', ';', ')']:
            return True
        if text[0].islower():
            return True
        if text.startswith('\u2013') or text.startswith('- '):
            return True
        continuators = [
            'si ', 'nel ', 'nella ', 'nelle ', 'una ', 'un ',
            'per ', 'con ', 'che ', 'di ', 'dell', 'nell',
            'al ', 'alla ', 'agli ', 'affronta', 'appare',
            'sono ', 'ha ', 'garantendo', 'riflettendo',
            'evidenziando', 'e la ', 'e il ', 'e i ',
            'risulti ', 'rappresenta', 'conferma', 'indica '
            '- ', 'Anno ', 'mediano ', '2021', '2022', '2023', '2024'
            'N.D. 20', 'N.D. 2021', 'N.D. 2022'
        ]
        return any(text.lower().startswith(c) for c in continuators)

    output_buffer.seek(0)
    doc = docx.Document(output_buffer)

    for _ in range(150):
        changed = False
        ps = list(doc.paragraphs)
        n = len(ps)

        for i in range(1, n - 1):
            curr_text = get_text(ps[i])
            if not curr_text:
                continue  # Skip already-cleared paragraphs

            # Find previous non-empty paragraph
            prev_idx = i - 1
            while prev_idx >= 0 and not get_text(ps[prev_idx]):
                prev_idx -= 1
            if prev_idx < 0:
                continue
            prev_text = get_text(ps[prev_idx])
            if not prev_text:
                continue

            prev_header = is_section_header(ps[prev_idx])

            # Mai processare paragrafi con forme ancorate
            if has_anchor_shape(ps[i]):
                continue

            # CASO 0: Valore placeholder (N.D.) — mai dentro intestazioni
            if curr_text in placeholder_vals and not ends_sentence(prev_text) and not prev_header:
                copy_text_to(ps[prev_idx], curr_text, ' ')
                remove_p(ps[i])
                changed = True
                break

            # CASO 1: Ragione sociale — mai dentro intestazioni
            elif curr_text == ragione_sociale and not prev_header:
                copy_text_to(ps[prev_idx], curr_text, ' ')
                remove_p(ps[i])
                changed = True
                break

            # CASO 2: Numero — solo tra Body Text
            elif (is_numero(curr_text) and not ends_sentence(prev_text) and
                  not prev_header and ps[prev_idx].style.name == 'Body Text'):
                copy_text_to(ps[prev_idx], curr_text, ' ')
                remove_p(ps[i])
                changed = True
                break

            # CASO 3: Punteggiatura — mai dentro intestazioni
            elif curr_text[0] in ['.', ',', '%'] and prev_text and not prev_header:
                if curr_text[0] == '.' and prev_text.endswith('.'):
                    remainder = curr_text[1:].strip()
                    if remainder:
                        copy_text_to(ps[prev_idx], remainder, ' ')
                    remove_p(ps[i])
                else:
                    copy_text_to(ps[prev_idx], curr_text, '')
                    remove_p(ps[i])
                changed = True
                break

            # CASO 4: Frammento — SOLO tra Body Text dello stesso stile
            elif (is_fragment(curr_text) and
                  not ends_sentence(prev_text) and
                  curr_text != ragione_sociale and
                  not prev_header and
                  ps[i].style.name == 'Body Text' and
                  ps[prev_idx].style.name == 'Body Text'):
                copy_text_to(ps[prev_idx], curr_text, ' ')
                remove_p(ps[i])
                changed = True
                break

        if not changed:
            break

    # =============================================================
    # Corregge i doppi punti tra runs adiacenti (es. "N.D." + ".")
    # =============================================================
    for p in doc.paragraphs:
        runs = [r for r in p.runs if r.text]
        for i in range(len(runs) - 1):
            curr_run = runs[i]
            next_run = runs[i + 1]
            # Se il run corrente termina con '.' e il prossimo inizia con '.'
            if curr_run.text.endswith('.') and next_run.text.startswith('.'):
                # Rimuovi il punto iniziale dal run successivo
                next_run.text = next_run.text[1:].lstrip()
                if not next_run.text:
                    next_run._r.getparent().remove(next_run._r)

    result = io.BytesIO()
    doc.save(result)
    result.seek(0)
    return result


def migliora_layout(output_buffer):
    """
    Migliora il layout del documento:
    1. Converte heading vuoti in Normal (pulisce TOC)
    2. Page break prima di ogni Heading 1 con contenuto
    3. KeepWithNext per heading
    4. Indentazione Body Text coerente
    5. Tabelle allineate con il testo, larghezza fissa, colonne uguali
    """
    output_buffer.seek(0)
    doc = docx.Document(output_buffer)

    BODY_INDENT = 412
    PAGE_W = 12240
    LEFT_MARGIN = 720
    RIGHT_MARGIN = 1080
    TEXT_WIDTH = PAGE_W - LEFT_MARGIN - RIGHT_MARGIN   # 10440
    TABLE_WIDTH = TEXT_WIDTH - BODY_INDENT              # 10028

    # ================================================================
    # 1. Heading vuoti → Normal + page break + keepWithNext
    # ================================================================
    for p in doc.paragraphs:
        style = p.style.name
        text = p.text.strip()
        pPr = p._p.find(qn('w:pPr'))

        if style in ['Heading 1', 'Heading 2', 'Heading 3'] and not text:
            try:
                p.style = doc.styles['Normal']
            except:
                pass
            continue

        if style == 'Heading 1' and text:
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p._p.insert(0, pPr)
            old_pb = pPr.find(qn('w:pageBreakBefore'))
            if old_pb is not None:
                pPr.remove(old_pb)
            pb = OxmlElement('w:pageBreakBefore')
            pb.set(qn('w:val'), 'true')
            pPr.insert(0, pb)
            kwn = pPr.find(qn('w:keepNext'))
            if kwn is None:
                kwn = OxmlElement('w:keepNext')
                pPr.append(kwn)

        if style in ['Heading 2', 'Heading 3'] and text:
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p._p.insert(0, pPr)
            kwn = pPr.find(qn('w:keepNext'))
            if kwn is None:
                kwn = OxmlElement('w:keepNext')
                pPr.append(kwn)

   

    # ================================================================
    # 3. Tabelle: larghezza fissa, indent con testo, colonne uguali
    # ================================================================
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    for tbl in doc.element.body.findall('.//' + qn('w:tbl')):
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        for tag in [qn('w:tblW'), qn('w:tblInd'), qn('w:jc'), qn('w:tblBorders')]:
            old = tblPr.find(tag)
            if old is not None:
                tblPr.remove(old)

        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'dxa')
        tblW.set(qn('w:w'), str(TABLE_WIDTH))
        tblPr.append(tblW)

        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:type'), 'dxa')
        tblInd.set(qn('w:w'), str(BODY_INDENT))
        tblPr.append(tblInd)

        tblBorders = parse_xml(
            r'<w:tblBorders %s>'
            r'<w:top w:val="single" w:sz="4" w:space="0" w:color="2B3A67"/>'
            r'<w:left w:val="single" w:sz="4" w:space="0" w:color="2B3A67"/>'
            r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="2B3A67"/>'
            r'<w:right w:val="single" w:sz="4" w:space="0" w:color="2B3A67"/>'
            r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="2B3A67"/>'
            r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="2B3A67"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr.append(tblBorders)

        rows = tbl.findall(qn('w:tr'))
        if not rows:
            continue

        # Evita che una singola riga venga spezzata tra due pagine, e che la tabella
        # venga divisa a metà tra due pagine: ogni riga (tranne l'ultima) resta
        # agganciata alla successiva, cosi' se non c'è spazio l'intera tabella
        # passa alla pagina seguente invece di spezzarsi a metà.
        for idx_row, row in enumerate(rows):
            trPr = row.find(qn('w:trPr'))
            if trPr is None:
                trPr = OxmlElement('w:trPr')
                row.insert(0, trPr)
            if trPr.find(qn('w:cantSplit')) is None:
                trPr.append(OxmlElement('w:cantSplit'))
            if idx_row < len(rows) - 1:
                for para in row.findall('.//' + qn('w:p')):
                    pPr = para.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        para.insert(0, pPr)
                    if pPr.find(qn('w:keepNext')) is None:
                        pPr.append(OxmlElement('w:keepNext'))

        num_cols = 0
        for cell in rows[0].findall(qn('w:tc')):
            tcPr = cell.find(qn('w:tcPr'))
            gs = tcPr.find(qn('w:gridSpan')) if tcPr is not None else None
            num_cols += int(gs.get(qn('w:val'), '1')) if gs is not None else 1

        if num_cols < 1:
            continue

        col_width = TABLE_WIDTH // num_cols

        for row in rows:
            for cell in row.findall(qn('w:tc')):
                tcPr = cell.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    cell.insert(0, tcPr)
                gs = tcPr.find(qn('w:gridSpan'))
                span = int(gs.get(qn('w:val'), '1')) if gs is not None else 1
                old_tcW = tcPr.find(qn('w:tcW'))
                if old_tcW is not None:
                    tcPr.remove(old_tcW)
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:type'), 'dxa')
                tcW.set(qn('w:w'), str(col_width * span))
                tcPr.insert(0, tcW)
    for p in doc.paragraphs:
        if p.style.name == 'Body Text' and p.text.strip():
             for run in p.runs:
                if run.text.strip() and run.font.size != Pt(10):
                    run.font.size = Pt(10)
             # Solo i paragrafi abbastanza lunghi da andare su 3+ righe vengono
             # giustificati: su paragrafi corti (1-2 righe) la giustificazione
             # tende a "stirare" vistosamente le parole dell'unica riga non finale.
             if len(p.text.strip()) >= 200 and p.paragraph_format.alignment is None:
                 p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # ================================================================
    # Normalizza gli spazi tra i contenuti: esattamente 2 paragrafi vuoti
    # tra un blocco e l'altro, 0 prima di ogni capitolo (Heading 1, che ha
    # gia' la sua interruzione di pagina). Copre anche i paragrafi vuoti
    # residui lasciati dai subdoc di tabelle/grafici dopo il rendering.
    # ================================================================
    def _is_blank(p):
        if p.text.strip():
            return False
        if p._p.findall('.//' + qn('w:drawing')):
            return False
        if p._p.findall('.//' + qn('w:tbl')):
            return False
        return True

    TOC_STYLES = {'toc 1', 'toc 2', 'toc 3', 'TOC 1', 'TOC 2', 'TOC 3'}
    paragrafi_norm = doc.paragraphs
    n_norm = len(paragrafi_norm)
    cover_page_end = 0
    appendice_start = n_norm
    exec_box_start = None
    exec_box_end = None
    for idx, pp in enumerate(paragrafi_norm):
        if pp.text.strip() == 'Sommario':
            cover_page_end = idx
        if pp.text.strip().startswith('Formalmente costituita') and exec_box_start is None:
            exec_box_start = idx
        if pp.text.strip().startswith('Con questo risultato') and exec_box_start is not None and exec_box_end is None:
            exec_box_end = idx
        if pp.text.strip() == 'Appendice' and pp.style.name == 'Heading 1':
            appendice_start = idx
            break

    # L'Appendice (formule) ha spaziature/rientri calibrati a mano per l'allineamento
    # visivo delle frazioni: non va normalizzata, altrimenti si sfasa.
    # Lo stesso vale per i 3 box "Totale Valore Produzione/Attivo/Forza Lavoro" in
    # Executive Summary: sono testo sovrapposto a un'unica immagine di sfondo fissa.
    i = cover_page_end + 1
    norm_actions = []
    while i < appendice_start:
        if _is_blank(paragrafi_norm[i]):
            start = i
            j = i
            while j < n_norm and _is_blank(paragrafi_norm[j]):
                j += 1
            end = j - 1
            before_p = paragrafi_norm[start - 1] if start - 1 >= 0 else None
            after_p = paragrafi_norm[end + 1] if end + 1 < n_norm else None
            before_style = before_p.style.name if before_p is not None else None
            after_style = after_p.style.name if after_p is not None else None
            after_text = after_p.text.strip() if after_p is not None else ''
            after_is_h1 = (after_p is not None and after_style == 'Heading 1' and after_text)
            # I riquadri annotati (icona/box grafici con offset verticale negativo) hanno
            # bisogno di uno spazio di sicurezza calibrato sulla loro altezza: non toccare
            # il gap subito prima o subito dopo un paragrafo con 2+ forme ancorate.
            before_has_multi_drawing = (
                before_p is not None
                and len(before_p._p.findall('.//' + qn('w:drawing'))) >= 2
            )
            after_has_multi_drawing = (
                after_p is not None
                and len(after_p._p.findall('.//' + qn('w:drawing'))) >= 2
            )
            in_exec_box_zone = (
                exec_box_start is not None and exec_box_end is not None
                and exec_box_start <= start <= exec_box_end
            )

            if before_style in TOC_STYLES or after_style in TOC_STYLES:
                target = None
            elif before_has_multi_drawing or after_has_multi_drawing:
                target = None
            elif in_exec_box_zone:
                target = None
            elif after_is_h1:
                target = 0
            else:
                target = 2
            norm_actions.append((start, end, target))
            i = j
        else:
            i += 1

    for start, end, target in reversed(norm_actions):
        if target is None:
            continue
        current_len = end - start + 1
        run_ps = [paragrafi_norm[k] for k in range(start, end + 1)]
        if target < current_len:
            for p_rm in run_ps[:current_len - target]:
                p_el = p_rm._p
                p_el.getparent().remove(p_el)
        elif target > current_len:
            anchor = run_ps[-1]._p if run_ps else None
            if anchor is not None:
                pPr_tpl = run_ps[0]._p.find(qn('w:pPr'))
                for _ in range(target - current_len):
                    new_p = OxmlElement('w:p')
                    if pPr_tpl is not None:
                        new_p.append(copy.deepcopy(pPr_tpl))
                    anchor.addnext(new_p)
                    anchor = new_p

    # ================================================================
    # Mantieni assieme le didascalie (Tabella/Figura) con tabelle e immagini,
    # cosi' non restano orfane su una pagina separata dal contenuto a cui si riferiscono
    # ================================================================
    def _set_keep_next(paragraph):
        pPr = paragraph._p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            paragraph._p.insert(0, pPr)
        if pPr.find(qn('w:keepNext')) is None:
            pPr.append(OxmlElement('w:keepNext'))

    paragrafi = doc.paragraphs
    n_par = len(paragrafi)
    for i, p in enumerate(paragrafi):
        testo = p.text.strip()
        if re.match(r'^Tabella\s+\d', testo):
            # La didascalia precede la tabella: resta assieme ai paragrafi vuoti
            # che la separano dalla tabella (senza oltrepassarla)
            _set_keep_next(p)
            j = i + 1
            while j < n_par and not paragrafi[j].text.strip():
                _set_keep_next(paragrafi[j])
                j += 1
        elif re.match(r'^Figura\s+\d', testo):
            # La didascalia segue l'immagine: risali fino al paragrafo con l'immagine
            j = i - 1
            while j >= 0:
                pj = paragrafi[j]
                _set_keep_next(pj)
                ha_immagine = bool(pj._p.findall('.//' + qn('w:drawing')))
                if pj.text.strip() or ha_immagine:
                    break
                j -= 1
  
    # Remove hyperlink character style from TOC entries
   # Remove hyperlinks from TOC by unwrapping w:hyperlink elements
    
    for p in doc.paragraphs:
        if p.style.name in ['TOC 1', 'TOC 2', 'TOC 3']:
            for hyperlink in p._p.findall('.//' + qn('w:hyperlink')):
                parent = hyperlink.getparent()
                idx = list(parent).index(hyperlink)
                for child in list(hyperlink):
                    parent.insert(idx, child)
                    idx += 1
                parent.remove(hyperlink)    
    # Fix inconsistent first line indentation
    for p in doc.paragraphs:
        if p.style.name == 'Body Text' and p.text.strip():
            pPr = p._p.find(qn('w:pPr'))
            if pPr is not None:
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    if ind.get(qn('w:firstLine')):
                        ind.attrib.pop(qn('w:firstLine'))
                    if ind.get(qn('w:hanging')):
                        ind.attrib.pop(qn('w:hanging'))            
    for tbl in doc.element.body.findall('.//' + qn('w:tbl')):
        rows = tbl.findall(qn('w:tr'))
        for row in rows:
            cells = row.findall(qn('w:tc'))
            for cell in cells:
                for para in cell.findall('.//' + qn('w:p')):
                    text = ''.join(t.text or '' for t in para.findall('.//' + qn('w:t'))).strip()
                    if re.match(r'^-?[\d\.,°%]+$', text):
                        pPr = para.find(qn('w:pPr'))
                        if pPr is None:
                            pPr = OxmlElement('w:pPr')
                            para.insert(0, pPr)
                        jc = pPr.find(qn('w:jc'))
                        if jc is None:
                            jc = OxmlElement('w:jc')
                            pPr.append(jc)
                        jc.set(qn('w:val'), 'right')
    result = io.BytesIO()
    doc.save(result)
    result.seek(0)
    return result
